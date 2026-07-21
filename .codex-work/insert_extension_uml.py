from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Pt


DOCX = Path(r"E:\swp391-rbl-project-se20a01_team3\docs\Group3_BeeAcademy_SDS_UML_UPDATED.docx")
IMAGES = Path(r"E:\swp391-rbl-project-se20a01_team3\.codex-work\extension_uml")
IMAGE_WIDTH = Emu(5715000)  # Match the existing SDS UML figures (6.25 in).


USE_CASES = [
    (
        "9.1. UC45 - Tích điểm từ kết quả bài kiểm tra",
        "UC45_Class_Diagram.png",
        "Figure UC45-1. Reward points derived from assessment scores",
        "UC45_Sequence_Diagram.png",
        "Figure UC45-2. Score-to-points and best-score delta flow",
        "Class diagram for UC45: reward points are derived from quiz and exam scores.",
        "Sequence diagram for UC45: first score, improved score, and no-increase alternatives.",
    ),
    (
        "9.2. UC46 - Đổi và sử dụng voucher giảm giá khóa học",
        "UC46_Class_Diagram.png",
        "Figure UC46-1. Voucher redemption and course-order discount classes",
        "UC46_Sequence_Diagram.png",
        "Figure UC46-2. Redeem, reserve, apply, use, and release voucher flow",
        "Class diagram for UC46: points are exchanged for vouchers used to discount course orders.",
        "Sequence diagram for UC46: voucher redemption and application during course checkout.",
    ),
    (
        "9.3. UC47 - Cấp và xác minh chứng chỉ",
        "UC47_Class_Diagram.png",
        "Figure UC47-1. Certificate eligibility, issuance, storage, and verification classes",
        "UC47_Sequence_Diagram.png",
        "Figure UC47-2. Certificate issuance and public verification flow",
        "Class diagram for UC47: certificate eligibility, issuance, storage, and verification.",
        "Sequence diagram for UC47: eligible issuance and public verification alternatives.",
    ),
]


def insert_before(target, paragraph):
    target._p.addprevious(paragraph._p)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_heading_before(doc, target, text, style, page_break=False):
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.paragraph_format.page_break_before = page_break
    set_keep_with_next(paragraph, True)
    insert_before(target, paragraph)
    return paragraph


def add_figure_before(doc, target, image_path, alt_text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    set_keep_with_next(paragraph, True)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=IMAGE_WIDTH)
    for prop in paragraph._p.xpath('.//wp:docPr'):
        prop.set("title", alt_text)
        prop.set("descr", alt_text)
    insert_before(target, paragraph)
    return paragraph


def add_caption_before(doc, target, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    insert_before(target, paragraph)
    return paragraph


def main():
    doc = Document(DOCX)
    if any(p.text.startswith("9.1. UC45") for p in doc.paragraphs):
        raise SystemExit("UC45-UC47 extension already exists; refusing to duplicate it.")

    target = next(
        p for p in doc.paragraphs
        if p.style and p.style.name == "Heading 1"
        and p.text.strip() == "III. Deployment and Verification Notes"
    )

    add_heading_before(
        doc,
        target,
        "9. Reward Points, Voucher and Certification Extensions",
        "Heading 2",
        page_break=True,
    )

    for index, (
        uc_heading,
        class_image,
        class_caption,
        sequence_image,
        sequence_caption,
        class_alt,
        sequence_alt,
    ) in enumerate(USE_CASES):
        add_heading_before(doc, target, uc_heading, "Heading 3", page_break=index > 0)
        add_heading_before(doc, target, "a. Class Diagram", "Heading 4")
        add_figure_before(doc, target, IMAGES / class_image, class_alt)
        add_caption_before(doc, target, class_caption)

        add_heading_before(doc, target, "b. Sequence Diagram", "Heading 4", page_break=True)
        add_figure_before(doc, target, IMAGES / sequence_image, sequence_alt)
        add_caption_before(doc, target, sequence_caption)

    # Ask Word to refresh fields (including the TOC) when the document is opened.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
        settings.append(update_fields)
    else:
        update_fields.set(qn("w:val"), "true")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
