from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


INPUT = Path(r"E:\swp391-rbl-project-se20a01_team3\docs\Group3_BeeAcademy_SDS_UML_UPDATED.docx")
OUT = Path(r"E:\swp391-rbl-project-se20a01_team3\.codex-work\sds_structure.json")


def paragraph_info(index, paragraph):
    drawings = []
    for drawing in paragraph._p.xpath('.//w:drawing'):
        embed_ids = drawing.xpath('.//a:blip/@r:embed')
        extents = drawing.xpath('.//wp:extent')
        extent = ({"cx": extents[0].get("cx"), "cy": extents[0].get("cy")}
                  if extents else None)
        for prop in drawing.xpath('.//wp:docPr'):
            drawings.append({
                "id": prop.get("id"),
                "name": prop.get("name"),
                "title": prop.get("title"),
                "descr": prop.get("descr"),
                "embed_ids": embed_ids,
                "extent": extent,
            })
    return {
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "drawings": drawings,
        "page_break_before": bool(paragraph.paragraph_format.page_break_before),
        "keep_with_next": bool(paragraph.paragraph_format.keep_with_next),
        "xml_has_page_break": bool(paragraph._p.xpath('.//w:br[@w:type="page"]')),
        "section_properties": bool(paragraph._p.xpath('./w:pPr/w:sectPr')),
        "alignment": str(paragraph.alignment),
        "space_before": paragraph.paragraph_format.space_before,
        "space_after": paragraph.paragraph_format.space_after,
        "runs": [
            {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "font_name": run.font.name,
                "font_size": run.font.size,
            }
            for run in paragraph.runs
        ],
    }


def table_info(index, table):
    rows = []
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    return {"index": index, "rows": rows}


def main():
    doc = Document(INPUT)
    payload = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "sections": [],
        "paragraphs": [paragraph_info(i, p) for i, p in enumerate(doc.paragraphs)],
        "tables": [table_info(i, t) for i, t in enumerate(doc.tables)],
    }
    for paragraph in payload["paragraphs"]:
        for drawing in paragraph["drawings"]:
            drawing["targets"] = [
                str(doc.part.rels[embed_id].target_ref)
                for embed_id in drawing["embed_ids"]
                if embed_id in doc.part.rels
            ]
    for i, section in enumerate(doc.sections):
        payload["sections"].append({
            "index": i,
            "width": section.page_width,
            "height": section.page_height,
            "orientation": str(section.orientation),
            "top_margin": section.top_margin,
            "bottom_margin": section.bottom_margin,
            "left_margin": section.left_margin,
            "right_margin": section.right_margin,
        })

    with zipfile.ZipFile(INPUT) as zf:
        payload["media"] = [
            {"name": info.filename, "size": info.file_size}
            for info in zf.infolist()
            if info.filename.startswith("word/media/")
        ]
    OUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({k: payload[k] for k in (
        "paragraph_count", "table_count", "inline_shape_count", "section_count"
    )}, ensure_ascii=False))
    print(f"media_count={len(payload['media'])}")
    print(OUT)


if __name__ == "__main__":
    main()
