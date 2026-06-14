# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def ct(cell, text, bold=False, center=False, sz=11, color=None, italic=False, wrap=True):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))

def th(table, headers, bg='1F3864', sz=11):
    """Format header row of a table."""
    row = table.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        set_cell_bg(c, bg)
        ct(c, h, bold=True, center=True, sz=sz, color='FFFFFF')

def td(table, data, alt_bg='EBF3FB', sz=11):
    """Fill data rows (starting row 1)."""
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.cell(r + 1, c)
            if r % 2 == 1:
                set_cell_bg(cell, alt_bg)
            ct(cell, str(val) if val is not None else '', sz=sz)

def set_col_widths(table, widths_cm):
    for i, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(15)
    run.font.bold = True
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True
    return p

def body(doc, text, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p

def bul(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'• {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    return p

def small_space(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

def mktbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    return t

# ═════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════
doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2.5)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(13)

# ── TRANG BÌA ────────────────────────────────────────────────────────────────
for _ in range(2): doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TRƯỜNG ĐẠI HỌC FPT – FPT UNIVERSITY')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Khoa Công nghệ Thông tin')
r.font.name = 'Times New Roman'; r.font.size = Pt(13)

for _ in range(5): doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BÁO CÁO THỰC HÀNH')
r.font.name = 'Times New Roman'; r.font.size = Pt(20); r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LAB 4 – TẠO TEST PLAN VỚI SỰ HỖ TRỢ CỦA AI')
r.font.name = 'Times New Roman'; r.font.size = Pt(17); r.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Chuẩn áp dụng: IEEE 829-2008  |  Môn học: SWT301 – Software Testing')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.italic = True

for _ in range(2): doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DỰ ÁN: BEE ACADEMY')
r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Nền Tảng Học Trực Tuyến Dành Cho Học Sinh THCS (Lớp 6–9)')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.italic = True

for _ in range(3): doc.add_paragraph()

# Cover info table
ctbl = mktbl(doc, 7, 2)
cover_info = [
    ('Môn học', 'SWT301 – Software Testing'),
    ('Chủ đề', 'Topic 5 – Test Management'),
    ('CLO', 'CLO5 – Xây dựng test plan toàn diện có tích hợp AI'),
    ('Chuẩn áp dụng', 'IEEE 829-2008'),
    ('Dự án thực hành', 'Bee Academy – Online Learning Platform v1.0.0'),
    ('Sinh viên thực hiện', 'Võ Văn Thành Đạt'),
    ('Ngày thực hiện', '12/06/2026'),
]
for i, (k, v) in enumerate(cover_info):
    rw = ctbl.rows[i]
    set_cell_bg(rw.cells[0], 'D6E4F7')
    ct(rw.cells[0], k, bold=True, sz=12)
    ct(rw.cells[1], v, sz=12)
ctbl.columns[0].width = Cm(5.5)
ctbl.columns[1].width = Cm(9)

doc.add_page_break()

# ── MỤC LỤC ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MỤC LỤC')
r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.font.bold = True

doc.add_paragraph()

toc = [
    (False, 'PHẦN A – TEST PLAN THEO IEEE 829-2008'),
    (True,  'Mục 1  –  Test Plan Identifier'),
    (True,  'Mục 2  –  Introduction'),
    (True,  'Mục 3  –  Test Items'),
    (True,  'Mục 4  –  Features to be Tested'),
    (True,  'Mục 5  –  Features Not to be Tested'),
    (True,  'Mục 6  –  Approach (Test Strategy)'),
    (True,  'Mục 7  –  Item Pass/Fail Criteria'),
    (True,  'Mục 8  –  Suspension & Resumption Requirements'),
    (True,  'Mục 9  –  Test Deliverables'),
    (True,  'Mục 10 –  Testing Tasks'),
    (True,  'Mục 11 –  Environmental Needs'),
    (True,  'Mục 12 –  Responsibilities'),
    (True,  'Mục 13 –  Staffing and Training Needs'),
    (True,  'Mục 14 –  Schedule'),
    (True,  'Mục 15 –  Risks and Contingencies'),
    (True,  'Mục 16 –  Approvals'),
    (False, 'PHẦN B – AI INTERACTION LOG'),
    (False, 'PHẦN C – CÂU HỎI THẢO LUẬN'),
]
for is_sub, label in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if is_sub:
        p.paragraph_format.left_indent = Cm(1.2)
    r = p.add_run(label)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.bold = not is_sub

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# DỰ ÁN THỰC HÀNH – BEE ACADEMY (tương tự phần EduTrack trong đề bài)
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
r = p.add_run('Dự án thực hành – Bee Academy')
r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.font.bold = True
# border bottom
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F3864')
pBdr.append(bot); pPr.append(pBdr)

h2(doc, 'Mô tả hệ thống')
body(doc,
    'Bee Academy là nền tảng học trực tuyến dành cho học sinh THCS (lớp 6–9), '
    'cung cấp các khóa học video theo môn học. Hệ thống gồm 4 nhóm người dùng:')
for item in [
    'Học sinh: Xem và mua khóa học, làm quiz theo chương, theo dõi tiến độ, nhận chứng chỉ hoàn thành',
    'Giáo viên: Tạo và quản lý khóa học, upload video bài giảng, cấu hình quiz, xem doanh thu',
    'Phụ huynh: Liên kết tài khoản với học sinh, theo dõi tiến độ học tập, liên hệ giáo viên',
    'Admin: Duyệt khóa học, quản lý tài khoản người dùng, xác nhận payout GV, tạo báo cáo',
]:
    bul(doc, item)

h2(doc, 'Thông tin phiên bản')
doc.add_paragraph()
t = mktbl(doc, 7, 2)
th(t, ['Thông tin', 'Chi tiết'])
td(t, [
    ('Tên hệ thống',        'Bee Academy – Online Learning Platform'),
    ('Phiên bản kiểm thử',  'v1.0.0'),
    ('Môi trường',          'Web application (React 19 + Vite 6 frontend / Spring Boot 3.2 backend / PostgreSQL – Supabase)'),
    ('Ngày bắt đầu test',   '12/06/2026'),
    ('Ngày kết thúc test',  '15/07/2026'),
    ('Team size',           '3 tester + 1 test lead'),
])
set_col_widths(t, [4.5, 10])

h2(doc, 'Các module chính')
modules = [
    ('Module 1: Authentication & Authorization',
     '- Đăng nhập / đăng xuất (email/password + Google OAuth)\n'
     '- Đăng ký tài khoản với OTP email\n'
     '- JWT ES256 (ECDSA P-256) qua Supabase GoTrue\n'
     '- Phân quyền theo role: STUDENT / TEACHER / PARENT / ADMIN'),
    ('Module 2: Course Catalog & Enrollment',
     '- Duyệt và tìm kiếm khóa học, filter theo danh mục / cấp lớp / giá\n'
     '- Xem chi tiết khóa học, xem bài học thử miễn phí (isFree=true)\n'
     '- Đăng ký khóa học sau khi thanh toán thành công'),
    ('Module 3: Payment (VNPay / MoMo)',
     '- Thêm vào giỏ hàng (yêu cầu đăng nhập)\n'
     '- Thanh toán qua VNPay / MoMo – tiền về TK công ty\n'
     '- Hệ thống ghi revenue_splits; xem lịch sử đơn hàng'),
    ('Module 4: Teacher Portal',
     '- CRUD khóa học / chương / bài học\n'
     '- Upload video bài giảng (Supabase private bucket)\n'
     '- Submit khóa học để Admin duyệt (DRAFT → PENDING)'),
    ('Module 5: Quiz & Assessment',
     '- Ngân hàng câu hỏi (CRUD, phân loại Easy/Medium/Hard)\n'
     '- Cấu hình quiz theo chương (số câu, thời gian, điểm đạt)\n'
     '- Học sinh làm bài – chấm điểm theo JSONB snapshot'),
    ('Module 6: Admin Portal',
     '- Duyệt / từ chối / yêu cầu sửa khóa học\n'
     '- Quản lý tài khoản người dùng\n'
     '- Xuất Excel payout và xác nhận chuyển khoản GV'),
    ('Module 7: Parent Portal',
     '- Gửi lời mời liên kết qua email học sinh\n'
     '- Theo dõi tiến độ học (chỉ khi link ACTIVE)\n'
     '- Liên hệ giáo viên, xem lịch sử thanh toán cho con'),
    ('Module 8: Student Dashboard',
     '- Hồ sơ cá nhân, đổi mật khẩu, upload avatar\n'
     '- Danh sách yêu thích, lịch sử đơn hàng\n'
     '- Xử lý lời mời liên kết từ phụ huynh'),
]
for title, detail in modules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.bold = True
    # detail lines
    for line in detail.split('\n'):
        line = line.strip()
        if line.startswith('-'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.2)
            p2.paragraph_format.space_after = Pt(1)
            r2 = p2.add_run(line)
            r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PHẦN A
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẦN A – TEST PLAN THEO IEEE 829-2008')
r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.font.bold = True
doc.add_paragraph()

# ─── MỤC 1 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 1 – TEST PLAN IDENTIFIER')

body(doc, 'Tài liệu này định danh Test Plan chính thức (Master Test Plan) cho dự án Bee Academy '
     'phiên bản v1.0.0, được lập theo chuẩn IEEE 829-2008.')
doc.add_paragraph()

t = mktbl(doc, 9, 2)
th(t, ['Trường', 'Giá trị'])
td(t, [
    ('Document ID',        'TP-BEEACADEMY-2026-001'),
    ('Project Name',       'Bee Academy – Online Learning Platform'),
    ('System Version',     'v1.0.0'),
    ('Test Plan Version',  '1.0'),
    ('Author',             'Võ Văn Thành Đạt'),
    ('Created Date',       '12/06/2026'),
    ('Review Date',        '19/06/2026'),
    ('Status',             'Draft'),
])
set_col_widths(t, [5, 9.5])

# ─── MỤC 2 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 2 – INTRODUCTION')

h2(doc, '2.1  Mục đích (Purpose)')
body(doc,
    'Tài liệu Test Plan này xác định phạm vi, phương pháp tiếp cận, nguồn lực và lịch '
    'trình kiểm thử cho hệ thống Bee Academy v1.0.0 – nền tảng học trực tuyến dành cho '
    'học sinh THCS (lớp 6–9) tại Việt Nam. Mục tiêu là đảm bảo hệ thống đáp ứng yêu cầu '
    'chức năng và phi chức năng trước khi phát hành chính thức.')

h2(doc, '2.2  Phạm vi áp dụng (Scope)')
body(doc,
    'Test Plan này áp dụng cho giai đoạn System Testing và UAT của Bee Academy v1.0.0, '
    'bao gồm 8 module chính: Authentication & Authorization, Course Catalog, Enrollment & '
    'Payment, Student Dashboard, Teacher Portal, Quiz & Assessment, Admin Portal, và Parent '
    'Portal. Môi trường kiểm thử là React 19 + Vite 6 (frontend) kết hợp Spring Boot 3.2 '
    '(backend) với cơ sở dữ liệu PostgreSQL thông qua Supabase.')

h2(doc, '2.3  Đối tượng đọc (Intended Audience)')
for item in ['Test Lead và QA Engineers – thực thi các hoạt động kiểm thử',
             'Development Team – fix defects và hỗ trợ môi trường',
             'Project Manager – theo dõi tiến độ và ra quyết định',
             'Product Owner / Stakeholders – phê duyệt UAT và sign-off release']:
    bul(doc, item)

h2(doc, '2.4  Tài liệu tham chiếu (References)')
refs = [
    'IEEE 829-2008 – IEEE Standard for Software and System Test Documentation',
    'CLAUDE.md – Bee Academy Architecture & Project Documentation',
    'BEE ACADEMY.md – UseCase v6.5 (48 Use Cases, 9 modules, 7 actors)',
    'ISTQB Foundation Level Syllabus – Chapter 5: Test Management',
    'Spring Boot 3.2 Security Documentation – JWT ES256 / ECDSA P-256',
    'Supabase Documentation – Storage Buckets & RLS Policies',
]
for r in refs:
    bul(doc, r)

# ─── MỤC 3 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 3 – TEST ITEMS')

body(doc, 'Dưới đây là danh sách các module/component của Bee Academy v1.0.0 được đưa vào '
     'phạm vi kiểm thử, kèm theo phiên bản và trạng thái sẵn sàng.')
doc.add_paragraph()

t = mktbl(doc, 9, 5)
th(t, ['Item ID', 'Module', 'Version', 'Mô tả ngắn', 'Trạng thái'])
td(t, [
    ('TI-01', 'Authentication & Authorization', 'v1.0.0',
     'Đăng nhập email/pw, OTP, Google OAuth (JWT ES256 + Supabase GoTrue)', 'Ready'),
    ('TI-02', 'Course Catalog & Detail',         'v1.0.0',
     'Duyệt, tìm kiếm, lọc khóa học; xem chi tiết và bài học thử', 'Ready'),
    ('TI-03', 'Enrollment & Payment',            'v1.0.0',
     'Giỏ hàng, thanh toán VNPay/MoMo, xem lịch sử đơn hàng', 'In Dev'),
    ('TI-04', 'Student Dashboard',               'v1.0.0',
     'Hồ sơ cá nhân, tài khoản, avatar upload, yêu thích', 'Ready'),
    ('TI-05', 'Teacher Portal',                  'v1.0.0',
     'CRUD khóa học/chương/bài; upload video (private) + tài liệu; submit duyệt', 'Partial'),
    ('TI-06', 'Quiz & Assessment System',        'v1.0.0',
     'Config quiz chương, ngân hàng câu hỏi, làm bài + chấm điểm JSONB snapshot', 'Ready'),
    ('TI-07', 'Admin Portal',                    'v1.0.0',
     'Duyệt khóa học, quản lý người dùng, báo cáo doanh thu, xác nhận payout', 'Partial'),
    ('TI-08', 'Parent Portal',                   'v1.0.0',
     'Liên kết học sinh, theo dõi tiến độ, liên hệ giáo viên', 'Ready'),
])
set_col_widths(t, [1.5, 3.2, 1.8, 6.5, 1.5])

# ─── MỤC 4 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 4 – FEATURES TO BE TESTED')

body(doc, 'Bảng dưới liệt kê tối thiểu 18 tính năng nằm trong phạm vi kiểm thử, '
     'phân bổ đều trên 8 module chính.')
doc.add_paragraph()

t = mktbl(doc, 19, 5)
th(t, ['Feature ID', 'Tên tính năng', 'Module', 'Priority', 'Mô tả ngắn'])
feat_data = [
    ('F-01', 'Đăng nhập email/password',                  'Auth',        'High',   'Nhập đúng/sai credentials; kiểm tra JWT trả về'),
    ('F-02', 'Đăng ký tài khoản với OTP email',           'Auth',        'High',   'Flow gửi OTP → xác minh → tạo profile STUDENT'),
    ('F-03', 'Đăng nhập Google OAuth (ES256 JWT)',        'Auth',        'High',   'Google consent → Supabase callback → sync profile'),
    ('F-04', 'Phân quyền theo role (RBAC)',               'Auth',        'High',   'Student/Teacher/Parent/Admin đúng route guards'),
    ('F-05', 'Refresh token & session timeout',           'Auth',        'Medium', 'Token hết hạn tự refresh; logout revoke token'),
    ('F-06', 'Duyệt danh sách khóa học + bộ lọc',        'Course',      'High',   'Filter category, giá, cấp lớp; phân trang'),
    ('F-07', 'Xem chi tiết và bài học thử miễn phí',     'Course',      'High',   'isFree=true cho phép xem; private video cần enroll'),
    ('F-08', 'Thêm vào giỏ – yêu cầu đăng nhập',         'Payment',     'High',   'Chưa login → redirect /login với state.from'),
    ('F-09', 'Thanh toán VNPay/MoMo (sandbox)',           'Payment',     'High',   'Tạo order → redirect payment → webhook → enroll'),
    ('F-10', 'Xem lịch sử đơn hàng',                     'Dashboard',   'Medium', 'Danh sách orders, trạng thái, ngày mua'),
    ('F-11', 'Cập nhật hồ sơ và upload avatar',          'Dashboard',   'Medium', 'Đổi tên, bio; upload ảnh lên Supabase Storage'),
    ('F-12', 'Giáo viên tạo/sửa/xóa khóa học',          'Teacher',     'High',   'CRUD với chapter/lesson; trạng thái DRAFT→PENDING'),
    ('F-13', 'Upload video bài giảng (private bucket)',   'Teacher',     'High',   'Tải lên Supabase private; lưu storagePath'),
    ('F-14', 'Submit khóa học để Admin duyệt',           'Teacher',     'High',   'Đổi trạng thái DRAFT→PENDING; admin nhận notification'),
    ('F-15', 'Config quiz chương + ngân hàng câu hỏi',   'Quiz',        'High',   'CRUD câu hỏi; set số câu, thời gian, passing score'),
    ('F-16', 'Học sinh làm quiz + xem kết quả',          'Quiz',        'High',   'Random pick từ ngân hàng; JSONB snapshot; chấm điểm'),
    ('F-17', 'Admin duyệt / từ chối / yêu cầu sửa',     'Admin',       'High',   'Review content; approve→PUBLISHED; reject kèm lý do'),
    ('F-18', 'Phụ huynh liên kết và theo dõi tiến độ',  'Parent',      'Medium', 'Gửi lời mời email; sau khi ACTIVE xem progress'),
]
td(t, feat_data)
set_col_widths(t, [1.4, 4, 2, 1.8, 5.3])

# ─── MỤC 5 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 5 – FEATURES NOT TO BE TESTED')

body(doc, 'Các tính năng sau được loại trừ khỏi phạm vi kiểm thử trong v1.0.0, '
     'kèm lý do cụ thể:')
doc.add_paragraph()

t = mktbl(doc, 9, 2)
th(t, ['Tính năng loại trừ', 'Lý do loại trừ'])
td(t, [
    ('Penetration Testing / Security Audit',
     'Yêu cầu chuyên gia bảo mật riêng; sẽ thực hiện trước go-live production'),
    ('Load Testing / Performance Testing',
     'Cần môi trường staging cấu hình giống production; không khả thi trong môi trường dev local'),
    ('Cross-browser (Safari, Edge)',
     'Scope v1.0 chỉ hỗ trợ Chrome + Firefox; Safari/Edge được test ở sprint tiếp theo'),
    ('Mobile responsiveness (< 768px)',
     'Version mobile sẽ được test riêng trong Sprint 3 khi hoàn thiện responsive design'),
    ('Tích hợp MoMo thực (production API)',
     'Chỉ có VNPay sandbox; MoMo production API cần ký hợp đồng và được test riêng'),
    ('Chức năng Chứng chỉ (UC42-43)',
     'Chưa được phát triển trong v1.0.0; dự kiến Sprint 4'),
    ('Tin nhắn / Chat AI với GV (UC20-21)',
     'Module chat chưa phát triển; cần tích hợp thêm WebSocket hoặc third-party'),
    ('Chức năng TK Ngân hàng GV (UC45-46)',
     'Chưa hoàn thiện backend; Admin payout manual sẽ test khi feature complete'),
])
set_col_widths(t, [5.5, 9])

# ─── MỤC 6 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 6 – APPROACH (TEST STRATEGY)')

body(doc, 'Chiến lược kiểm thử của Bee Academy v1.0.0 áp dụng kiểm thử đa tầng '
     '(multi-level), kết hợp kỹ thuật black-box và white-box, hướng đến đảm bảo '
     'chất lượng toàn diện trước khi phát hành.')

h2(doc, '6.1  Test Levels (Cấp độ kiểm thử)')
doc.add_paragraph()
t = mktbl(doc, 5, 4)
th(t, ['Cấp độ', 'Phụ trách', 'Kỹ thuật', 'Công cụ'])
td(t, [
    ('Unit Testing',        'Dev Team',       'White-box, Statement Coverage (≥80%)', 'JUnit 5 / Jest'),
    ('Integration Testing', 'Dev + QA',       'API Testing, Black-box, Contract Test', 'Postman / REST Assured'),
    ('System Testing',      'QA Team',        'Black-box: EP, BVA; Exploratory', 'Playwright / JIRA'),
    ('Acceptance Testing',  'PO + Stakeholders', 'Scenario-based, User Acceptance', 'Manual / TestRail'),
])
set_col_widths(t, [3.5, 3, 4, 4])

h2(doc, '6.2  Test Types (Loại hình kiểm thử)')
doc.add_paragraph()
t = mktbl(doc, 5, 3)
th(t, ['Loại hình', 'Mục tiêu', 'Phạm vi áp dụng'])
td(t, [
    ('Functional Testing',     'Kiểm tra logic nghiệp vụ đúng theo UseCase v6.5',
     'Tất cả 18 features trong Mục 4'),
    ('Security Testing',       'Kiểm tra JWT authentication, authorization bypass, SQL injection',
     'Module Auth, Teacher Portal, Admin Portal'),
    ('Usability Testing',      'Kiểm tra UX/UI trực quan, responsive desktop',
     'Landing Page, Course Detail, Student Dashboard'),
    ('Regression Testing',     'Đảm bảo bug fix không phá vỡ tính năng hiện có',
     'Sau mỗi lần fix defect Critical/Major'),
])
set_col_widths(t, [3.5, 6, 5])

h2(doc, '6.3  Test Design Techniques')
for item in [
    'Equivalence Partitioning (EP) – phân hoạch tương đương cho input validation (email, password, giá tiền)',
    'Boundary Value Analysis (BVA) – kiểm tra biên cho price range, grade filter, file size upload',
    'Decision Table – cho flow thanh toán (VNPay callback: success/fail/cancel/timeout)',
    'State Transition – cho vòng đời khóa học (DRAFT→PENDING→PUBLISHED/REJECTED)',
    'Error Guessing – JWT expired, invalid OTP, Supabase signed URL hết hạn',
    'Exploratory Testing – tìm lỗi không có trong test case bằng session-based exploration',
]:
    bul(doc, item)

h2(doc, '6.4  Entry & Exit Criteria')
doc.add_paragraph()
t = mktbl(doc, 3, 3)
th(t, ['Cấp độ', 'Entry Criteria (Điều kiện bắt đầu)', 'Exit Criteria (Điều kiện kết thúc)'])
td(t, [
    ('System Testing',
     'Build deploy thành công lên test env; Unit test pass ≥80%; Test cases & data đã chuẩn bị',
     'Test case pass rate ≥95%; Không có Critical defect còn mở; Test Summary Report được tạo'),
    ('Acceptance Testing',
     'System Testing đã pass Exit Criteria; Regression Testing hoàn thành; Môi trường UAT sẵn sàng',
     'Tất cả UAT scenarios PASS; PO ký sign-off; Release checklist hoàn chỉnh'),
])
set_col_widths(t, [2.5, 6.5, 5.5])

# ─── MỤC 7 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 7 – ITEM PASS/FAIL CRITERIA')

h2(doc, '7.1  Pass Criteria (Tiêu chí PASS)')
for item in [
    'Tỷ lệ test case PASS ≥ 95% trên tổng số test cases đã thực thi',
    'Không có defect mức Critical hoặc Blocker còn ở trạng thái Open',
    'Không có defect mức Major còn Open sau ngày D-3 trước release',
    'Tất cả defect Major đã có workaround được document trong JIRA',
    'Module Authentication: không có bất kỳ lỗi Security nào (JWT bypass, SQL injection)',
    'Module Payment: toàn bộ luồng thanh toán VNPay sandbox hoàn thành thành công',
    'Test Summary Report được Test Lead phê duyệt',
]:
    bul(doc, item)

h2(doc, '7.2  Fail Criteria (Tiêu chí FAIL – dừng release)')
for item in [
    'Tỷ lệ test case PASS < 90%',
    'Tồn tại ít nhất 1 defect Critical/Blocker chưa được fix',
    'Module Authentication có bất kỳ lỗi security nào còn Open',
    'Luồng thanh toán VNPay không hoàn thành được (payment flow broken)',
    'Dữ liệu user bị mất hoặc bị ghi đè sai sau khi thực hiện giao dịch',
    'Hệ thống crash hoặc không khởi động được trong môi trường test',
]:
    bul(doc, item)

h2(doc, '7.3  Defect Severity Definition')
doc.add_paragraph()
t = mktbl(doc, 5, 3)
th(t, ['Mức độ', 'Định nghĩa', 'Ví dụ trong Bee Academy'])
td(t, [
    ('Critical / Blocker', 'Hệ thống không sử dụng được; không có workaround',
     'Không thể đăng nhập; payment crash; data loss'),
    ('Major',              'Chức năng quan trọng bị lỗi; có workaround tạm thời',
     'Quiz không chấm điểm đúng; video không phát được'),
    ('Minor',              'Chức năng hoạt động nhưng có lỗi nhỏ; ít ảnh hưởng',
     'Text sai chính tả; UI misalignment nhỏ'),
    ('Cosmetic',           'Lỗi giao diện không ảnh hưởng chức năng',
     'Icon hiển thị sai màu; spacing không đều'),
])
set_col_widths(t, [2.5, 5.5, 6.5])

# ─── MỤC 8 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 8 – SUSPENSION CRITERIA AND RESUMPTION REQUIREMENTS')

body(doc, 'Bảng dưới mô tả các tình huống phải tạm dừng kiểm thử và điều kiện cần '
     'đáp ứng để tiếp tục.')
doc.add_paragraph()

t = mktbl(doc, 5, 3)
th(t, ['Tình huống', 'Điều kiện tạm dừng', 'Điều kiện tiếp tục'])
td(t, [
    ('Build/Deploy thất bại',
     'Môi trường test không deploy được phiên bản mới; error rate > 50%',
     'Build pass; smoke test 5 tính năng cốt lõi thành công; Dev xác nhận'),
    ('Critical Blocker xuất hiện',
     'Phát hiện ≥1 defect Critical ảnh hưởng module đang test',
     'Dev fix và merge; QA verify fix; Regression test module đó pass'),
    ('Môi trường test down',
     'Backend hoặc Supabase unreachable > 30 phút liên tục',
     'Môi trường được restore; health check pass; Test Lead xác nhận'),
    ('Test data bị corrupt',
     'Dữ liệu test bị sai, thiếu hoặc bị ghi đè do lỗi script/migration',
     'Restore lại test data từ backup; QA xác nhận data integrity'),
])
set_col_widths(t, [3.5, 5.5, 5.5])

body(doc,
    'Khi tạm dừng, Test Lead phải ghi nhận nguyên nhân vào Daily Report, '
    'cập nhật trạng thái trong JIRA và thông báo đến Project Manager trong vòng 1 giờ.')

# ─── MỤC 9 ──────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 9 – TEST DELIVERABLES')

body(doc, 'Danh sách tài liệu đầu ra (deliverables) được tạo ra trong suốt quá trình kiểm thử:')
doc.add_paragraph()

t = mktbl(doc, 11, 4)
th(t, ['Giai đoạn', 'Tài liệu', 'Người chịu trách nhiệm', 'Deadline'])
td(t, [
    ('Pre-testing',  'Master Test Plan (tài liệu này)',              'Test Lead',   '12/06/2026'),
    ('Pre-testing',  'Test Cases (Excel / TestRail)',                'QA Team',     '17/06/2026'),
    ('Pre-testing',  'Test Data Script (SQL + seed data)',           'QA Team',     '16/06/2026'),
    ('Pre-testing',  'Test Environment Setup Guide',                 'Dev + QA',    '15/06/2026'),
    ('During testing', 'Daily Test Progress Report',                 'Test Lead',   'Hàng ngày'),
    ('During testing', 'Defect Report (JIRA tickets)',               'QA Team',     'Khi phát hiện'),
    ('During testing', 'Test Execution Log',                         'QA Team',     'Hàng ngày'),
    ('During testing', 'Change Request Log (nếu req thay đổi)',      'Test Lead',   'Khi phát sinh'),
    ('Post-testing', 'Test Summary Report',                          'Test Lead',   '14/07/2026'),
    ('Post-testing', 'Defect Metrics & Quality Report',              'Test Lead',   '14/07/2026'),
])
set_col_widths(t, [2.8, 5, 3.5, 3.2])

# ─── MỤC 10 ─────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 10 – TESTING TASKS')

body(doc, 'Bảng dưới phân rã các hoạt động kiểm thử thành các task cụ thể, có thứ tự '
     'thực hiện và dependency rõ ràng.')
doc.add_paragraph()

t = mktbl(doc, 11, 7)
th(t, ['Task ID', 'Tên Task', 'Predecessor', 'Effort', 'Role', 'Start', 'End'])
td(t, [
    ('T01', 'Phân tích requirements & UseCase v6.5',    '–',              '8h',  'Test Lead',       '12/06', '13/06'),
    ('T02', 'Thiết kế test cases (18 features)',         'T01',            '24h', 'QA Team',         '14/06', '17/06'),
    ('T03', 'Chuẩn bị test data & seed scripts',        'T01',            '8h',  'QA Team',         '14/06', '16/06'),
    ('T04', 'Cài đặt và cấu hình môi trường test',     'T01',            '4h',  'Dev + QA',        '12/06', '14/06'),
    ('T05', 'System Test – Auth + Course + Dashboard',  'T02,T03,T04',    '20h', 'QA Team',         '18/06', '23/06'),
    ('T06', 'System Test – Payment + Teacher Portal',   'T05',            '20h', 'QA Team',         '24/06', '30/06'),
    ('T07', 'System Test – Quiz + Admin + Parent',      'T06',            '20h', 'QA Team',         '01/07', '07/07'),
    ('T08', 'Regression Testing (sau khi fix defects)', 'T07',            '16h', 'QA Team',         '08/07', '10/07'),
    ('T09', 'User Acceptance Testing (UAT)',            'T08',            '16h', 'PO + Stakeholders','11/07', '13/07'),
    ('T10', 'Viết Test Summary Report & Sign-off',      'T09',            '4h',  'Test Lead',        '14/07', '15/07'),
])
set_col_widths(t, [1.3, 5, 2.5, 1.5, 2.5, 1.5, 1.5])

# ─── MỤC 11 ─────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 11 – ENVIRONMENTAL NEEDS')

h2(doc, '11.1  Yêu cầu phần cứng (Hardware Requirements)')
doc.add_paragraph()
t = mktbl(doc, 4, 2)
th(t, ['Thành phần', 'Cấu hình tối thiểu'])
td(t, [
    ('Test Server (Backend)',     'CPU: 4 cores, RAM: 8GB, Storage: 100GB SSD; hoặc Cloud: Supabase + Railway'),
    ('Tester Workstation',        'CPU: 2 cores, RAM: 8GB, OS: Windows 11 / Ubuntu 22.04'),
    ('CI/CD Pipeline',            'GitHub Actions (2 workers, 4GB RAM per runner)'),
])
set_col_widths(t, [4.5, 10])

h2(doc, '11.2  Yêu cầu phần mềm (Software Requirements)')
doc.add_paragraph()
t = mktbl(doc, 12, 3)
th(t, ['Phần mềm', 'Phiên bản', 'Mục đích'])
td(t, [
    ('OS',                  'Windows 11 / Ubuntu 22.04',  'Môi trường chạy test'),
    ('Node.js',             '20 LTS',                     'Build và chạy React 19 + Vite frontend'),
    ('Java JDK',            '17',                         'Runtime Spring Boot 3.2 backend'),
    ('Maven',               '3.9+',                       'Build tool backend'),
    ('PostgreSQL (Supabase)','15',                         'Cơ sở dữ liệu; sử dụng Supabase test project'),
    ('Chrome',              'Latest',                     'Primary browser (System Test)'),
    ('Firefox',             'Latest',                     'Secondary browser (Regression)'),
    ('Postman',             'v10+',                       'API testing – Auth endpoints, JWT validation'),
    ('Playwright',          '1.40+',                      'UI automation – E2E test scenarios'),
    ('JIRA',                'Cloud',                      'Defect tracking và test management'),
    ('Git / GitHub',        'Latest',                     'Version control; CI/CD integration'),
])
set_col_widths(t, [3.5, 3.5, 7.5])

h2(doc, '11.3  Cấu hình môi trường test (Test Environment Configuration)')
for item in [
    'Frontend: http://localhost:3000 (Vite dev server, port đã cố định trong vite.config.ts)',
    'Backend: http://localhost:8080 (Spring Boot, cổng cấu hình trong backend/.env)',
    'Database: Supabase PostgreSQL (dùng project test riêng, tách biệt production)',
    'Storage: Supabase Storage – 2 buckets: course-videos (private) + course-docs (public)',
    'JWT: ES256 (ECDSA P-256) verify qua Supabase JWKS endpoint',
    'CORS: cho phép origin http://localhost:3000',
    'DEV_MODE=true: OTP log ra console thay vì gửi email thật',
]:
    bul(doc, item)

h2(doc, '11.4  Test Data Requirements')
for item in [
    '5 tài khoản Student (valid credentials + edge cases: mật khẩu sai, tài khoản bị khóa)',
    '3 tài khoản Teacher (1 có khóa học đã duyệt, 1 đang pending, 1 mới tạo)',
    '2 tài khoản Parent (1 đã link với student, 1 chưa link)',
    '1 tài khoản Admin',
    '10 khóa học mẫu (DRAFT: 3, PENDING: 3, PUBLISHED: 4, REJECTED: 1)',
    'Test payment cards VNPay sandbox: success + fail + timeout scenarios',
    '50 câu hỏi ngân hàng mẫu (10 câu/module, phân bổ Easy/Medium/Hard)',
    'File video mẫu (≤100MB, định dạng MP4) và tài liệu PDF (≤10MB) cho upload test',
]:
    bul(doc, item)

# ─── MỤC 12 ─────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 12 – RESPONSIBILITIES')

body(doc, 'Ma trận RACI phân công trách nhiệm cho từng hoạt động kiểm thử '
     '(R=Responsible, A=Accountable, C=Consulted, I=Informed):')
doc.add_paragraph()

t = mktbl(doc, 9, 6)
th(t, ['Hoạt động', 'Test Lead', 'QA Engineer', 'Developer', 'PM', 'PO'])
td(t, [
    ('Viết Test Plan',          'R/A', 'C',   'C',   'I',   'I'),
    ('Thiết kế Test Cases',     'A',   'R',   'C',   'I',   'C'),
    ('Chuẩn bị Test Data',      'A',   'R',   'C',   'I',   '–'),
    ('Cài đặt môi trường',      'A',   'R',   'R/A', 'I',   '–'),
    ('Thực thi Testing',        'A',   'R',   'I',   'I',   'I'),
    ('Bug Report & Tracking',   'A',   'R',   'I',   'I',   'I'),
    ('Fix Defects',             'I',   'C',   'R/A', 'I',   '–'),
    ('Test Summary Report',     'R/A', 'C',   'I',   'I',   'I'),
])
set_col_widths(t, [4.5, 2, 2.2, 2.2, 1.5, 1.5])

# ─── MỤC 13 ─────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 13 – STAFFING AND TRAINING NEEDS')

h2(doc, '13.1  Thành phần nhân lực')
doc.add_paragraph()
t = mktbl(doc, 4, 5)
th(t, ['Vai trò', 'Số lượng', 'Allocation', 'Kỹ năng yêu cầu', 'Đào tạo cần thiết'])
td(t, [
    ('Test Lead',     '1', '100%',
     'IEEE 829-2008; ISTQB Foundation; API testing; JIRA',
     'Không cần thêm'),
    ('QA Engineer',   '2', '100%',
     'Manual testing; SQL cơ bản; Chrome DevTools; Postman',
     'Playwright workshop (2 ngày); Supabase API basics (1 ngày)'),
    ('Dev (support)', '2', '20%',
     'Unit test JUnit 5; code review; environment setup',
     'Không cần thêm'),
])
set_col_widths(t, [2.5, 1.8, 2, 5, 3.2])

h2(doc, '13.2  Kế hoạch đào tạo (Training Plan)')
doc.add_paragraph()
t = mktbl(doc, 3, 4)
th(t, ['Khóa đào tạo', 'Đối tượng', 'Thời lượng', 'Thời điểm'])
td(t, [
    ('Playwright E2E Automation Workshop',     'QA Engineers', '2 ngày',   'Tuần 1 (12-13/06/2026)'),
    ('Supabase Storage & JWT ES256 Overview',  'QA Team',      '3 tiếng',  'Ngày 14/06/2026'),
])
set_col_widths(t, [5.5, 2.5, 2.5, 4])

# ─── MỤC 14 ─────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 14 – SCHEDULE')

h2(doc, '14.1  Milestones')
doc.add_paragraph()
t = mktbl(doc, 7, 3)
th(t, ['Milestone', 'Ngày', 'Mô tả'])
td(t, [
    ('MS-01', '12/06/2026', 'Kickoff, Test Plan được phê duyệt'),
    ('MS-02', '17/06/2026', 'Test Cases & Test Data ready; môi trường test sẵn sàng'),
    ('MS-03', '07/07/2026', 'System Testing hoàn thành (cả 8 modules)'),
    ('MS-04', '10/07/2026', 'Regression Testing hoàn thành'),
    ('MS-05', '13/07/2026', 'UAT hoàn thành; PO ký sign-off'),
    ('MS-06', '15/07/2026', 'Test Summary Report được nộp'),
])
set_col_widths(t, [2.5, 3.5, 8.5])

h2(doc, '14.2  Lịch tuần (Gantt Chart)')
doc.add_paragraph()
t = mktbl(doc, 8, 6)
th(t, ['Phase / Activity', 'T1 (12-16/6)', 'T2 (17-23/6)', 'T3 (24-30/6)', 'T4 (1-10/7)', 'T5 (11-15/7)'])
phase_data = [
    ('Planning & Test Design',       'XX', 'XX', '',   '',   ''),
    ('Environment Setup',            'XX', '',   '',   '',   ''),
    ('System Test – Auth & Course',  '',   'XX', '',   '',   ''),
    ('System Test – Pay & Teacher',  '',   '',   'XX', '',   ''),
    ('System Test – Quiz,Admin,PH',  '',   '',   '',   'XX', ''),
    ('Regression Testing',           '',   '',   '',   'XX', ''),
    ('UAT & Test Closure',           '',   '',   '',   '',   'XX'),
]
td(t, phase_data)
set_col_widths(t, [4.5, 2.2, 2.2, 2.2, 2.2, 2.2])

# ─── MỤC 15 ─────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 15 – RISKS AND CONTINGENCIES')

body(doc, 'Bảng dưới xác định các rủi ro kiểm thử, đánh giá mức độ (L×I) và '
     'kế hoạch giảm thiểu / dự phòng. '
     '(L=Likelihood: H/M/L  |  I=Impact: H/M/L)')
doc.add_paragraph()

t = mktbl(doc, 10, 7)
th(t, ['Risk ID', 'Rủi ro', 'L', 'I', 'Level', 'Mitigation', 'Contingency'])
td(t, [
    ('R01', 'Môi trường test không ổn định (Supabase local dev, CORS errors)', 'M','H','High',
     'Chuẩn bị môi trường staging riêng; docker-compose backup',
     'Dùng Supabase hosted test project thay localhost'),
    ('R02', 'JWT ES256 key rotation làm tất cả user bị logout bất ngờ', 'L','H','Medium',
     'Không rotate JWKS key trong quá trình test; ghi chú trong env setup',
     'Restart backend và re-login; ghi log incident'),
    ('R03', 'VNPay sandbox không phản ánh đúng behavior production', 'H','M','High',
     'Dùng cả test card success + fail; document sai lệch nếu có',
     'Test payment flow với mock server nếu sandbox không ổn định'),
    ('R04', 'Requirements thay đổi trong khi test (UseCase v6.5 chưa frozen)', 'M','H','High',
     'Freeze requirements từ T1; mọi thay đổi qua Change Request',
     'Re-scope test cases; điều chỉnh schedule nếu cần'),
    ('R05', 'Thiếu nhân lực do nghỉ ốm hoặc lý do cá nhân', 'L','H','Medium',
     'Cross-training giữa QA Engineers; tài liệu test cases rõ ràng',
     'Test Lead tạm thời hỗ trợ thực thi; ưu tiên test features High'),
    ('R06', 'Build chậm trễ từ Dev Team ảnh hưởng schedule', 'M','M','Medium',
     'Daily stand-up; CI/CD tự động; early integration merge',
     'Ưu tiên test các module đã Ready trước; delay test module In Dev'),
    ('R07', 'Test data bị corrupt sau khi test payment flow', 'M','H','High',
     'Dùng transaction rollback hoặc reset script sau mỗi payment test session',
     'Restore test DB từ seed script; ghi nhận incident vào log'),
    ('R08', 'Supabase Storage signed URL hết hạn (TTL 1h) trong mid-test session', 'L','M','Low',
     'Regenerate signed URL trước mỗi video access test; note TTL trong test case',
     'Tăng TTL trong test env hoặc test video trong session < 1h'),
    ('R09', 'Private video bucket bị misconfigured → public access', 'L','H','Medium',
     'Verify RLS policy trên Supabase dashboard trước khi test',
     'Fix policy ngay; security incident report; re-test toàn bộ video access'),
])
set_col_widths(t, [1.3, 4.7, 0.7, 0.7, 1.5, 3.8, 3.8])

# ─── MỤC 16 ─────────────────────────────────────────────────────────────────
h1(doc, 'MỤC 16 – APPROVALS')

body(doc, 'Tài liệu này có hiệu lực khi được ký phê duyệt bởi tất cả các bên liên quan dưới đây:')
doc.add_paragraph()

t = mktbl(doc, 5, 4)
th(t, ['Vai trò', 'Họ và tên', 'Chữ ký', 'Ngày phê duyệt'])
td(t, [
    ('Test Lead',        'Võ Văn Thành Đạt',  '', ''),
    ('Project Manager',  'Nguyễn Văn A',       '', ''),
    ('Product Owner',    'Trần Thị B',          '', ''),
    ('Dev Lead',         'Lê Văn C',            '', ''),
])
set_col_widths(t, [3.5, 3.5, 4, 3.5])

doc.add_paragraph()
body(doc,
    'Lưu ý: Tài liệu này ở trạng thái Draft. Chữ ký sẽ được thu thập sau khi '
    'hoàn thành vòng review lần 1 (Review Date: 19/06/2026).')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PHẦN B – AI INTERACTION LOG
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẦN B – AI INTERACTION LOG')
r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.font.bold = True
doc.add_paragraph()

body(doc,
    'Phần này ghi lại 3 lần tương tác với AI (Claude Sonnet 4.6) khi soạn thảo Test Plan, '
    'kèm đánh giá chất lượng kết quả và những điểm đã được chỉnh sửa.')

# --- Interaction 1 ---
h2(doc, 'Tương tác 1 – Mục 1 & 2: Test Plan Identifier + Introduction')
doc.add_paragraph()

t = mktbl(doc, 2, 2)
th(t, ['Thông tin', 'Chi tiết'], bg='2E4057')
td(t, [('Prompt đã dùng',
    'Create Test Plan Identifier and Introduction sections for Bee Academy - an online learning '
    'platform for Vietnamese THCS students (grade 6-9). Tech stack: React 19 + Spring Boot 3.2 '
    '+ PostgreSQL via Supabase (JWT ES256/ECDSA P-256). Users: students, teachers, parents, admins. '
    'Payment: VNPay/MoMo. Follow IEEE 829-2008 format.')])
set_col_widths(t, [3.5, 11])

doc.add_paragraph()
body(doc, '[Ảnh chụp màn hình cuộc hội thoại với AI – đính kèm bên dưới]')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[ Screenshot 1: Prompt và Response từ AI cho Mục 1 & 2 ]')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.italic = True
r.font.color.rgb = RGBColor(0x80,0x80,0x80)

doc.add_paragraph()
h3(doc, 'Bảng đánh giá kết quả AI:')
t = mktbl(doc, 6, 3)
th(t, ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'])
td(t, [
    ('Đúng chuẩn IEEE 829-2008',            '4/5',
     'AI tạo đúng các trường Document ID, Author, Status nhưng thiếu Review Date và document history'),
    ('Phù hợp với context Bee Academy',     '3/5',
     'AI không biết về JWT ES256/ECDSA P-256 của Supabase, đề xuất HS256 generic'),
    ('Đủ chi tiết, có thể dùng ngay',       '3/5',
     'Introduction quá ngắn, thiếu phần tài liệu tham chiếu đặc thù dự án'),
    ('Cần chỉnh sửa nhiều hay ít',          '3/5',
     'Phải chỉnh sửa vừa phải: thêm JWT specifics, Supabase endpoints, VNPay context'),
    ('Tổng',                                 '13/20', ''),
])
set_col_widths(t, [5, 2, 7.5])

h3(doc, 'Những gì đã chỉnh sửa so với output AI:')
for item in [
    'Thêm thông tin JWT ES256 (ECDSA P-256) thay vì HS256 mà AI đề xuất',
    'Bổ sung Supabase JWKS endpoint cụ thể vào References',
    'Thêm UseCase v6.5 vào danh sách tài liệu tham chiếu',
    'Điều chỉnh scope: loại bỏ Next.js (AI tưởng dự án dùng Next.js) → Vite 6',
]:
    bul(doc, item)

# --- Interaction 2 ---
h2(doc, 'Tương tác 2 – Mục 6: Approach (Test Strategy)')
doc.add_paragraph()

t = mktbl(doc, 2, 2)
th(t, ['Thông tin', 'Chi tiết'], bg='2E4057')
td(t, [('Prompt đã dùng',
    'Write the Test Approach section for Bee Academy following IEEE 829-2008. The system includes: '
    'Auth (JWT ES256 ECDSA, Google OAuth, OTP), Course management with video private Supabase bucket, '
    'VNPay payment integration, Teacher portal with video upload, Quiz system with JSONB snapshot grading. '
    'Include test levels (unit/integration/system/acceptance), test types, design techniques '
    '(EP, BVA, decision table, state transition), entry/exit criteria for each level, and tools.')])
set_col_widths(t, [3.5, 11])

doc.add_paragraph()
body(doc, '[Ảnh chụp màn hình cuộc hội thoại với AI – đính kèm bên dưới]')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[ Screenshot 2: Prompt và Response từ AI cho Mục 6 ]')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.italic = True
r.font.color.rgb = RGBColor(0x80,0x80,0x80)

doc.add_paragraph()
h3(doc, 'Bảng đánh giá kết quả AI:')
t = mktbl(doc, 6, 3)
th(t, ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'])
td(t, [
    ('Đúng chuẩn IEEE 829-2008',            '5/5',
     'Sau khi được cung cấp context đầy đủ, AI tạo đúng cấu trúc 4 test levels với entry/exit criteria'),
    ('Phù hợp với context Bee Academy',     '4/5',
     'AI áp dụng EP và BVA đúng cho các trường hợp của Bee Academy (giá tiền VND, cấp lớp 6-9)'),
    ('Đủ chi tiết, có thể dùng ngay',       '4/5',
     'Khá đầy đủ; entry/exit criteria cụ thể, test types phân loại rõ ràng'),
    ('Cần chỉnh sửa nhiều hay ít',          '4/5',
     'Chỉ cần chỉnh sửa nhỏ: thay Selenium → Playwright, thêm State Transition cho course lifecycle'),
    ('Tổng',                                 '17/20', ''),
])
set_col_widths(t, [5, 2, 7.5])

h3(doc, 'Những gì đã chỉnh sửa so với output AI:')
for item in [
    'Thay Selenium 4.x bằng Playwright 1.40+ (phù hợp hơn với React SPA)',
    'Thêm State Transition Testing cho vòng đời khóa học: DRAFT→PENDING→PUBLISHED/REJECTED',
    'Bổ sung test type Security Testing riêng cho module Auth (AI gộp chung vào Functional)',
    'Điều chỉnh pass rate threshold: từ 90% (AI đề xuất) lên 95% theo requirement của project',
]:
    bul(doc, item)

# --- Interaction 3 ---
h2(doc, 'Tương tác 3 – Mục 15: Risks and Contingencies')
doc.add_paragraph()

t = mktbl(doc, 2, 2)
th(t, ['Thông tin', 'Chi tiết'], bg='2E4057')
td(t, [('Prompt đã dùng',
    'Identify and analyze testing risks for Bee Academy online learning platform (React 19 + Spring Boot 3.2 '
    '+ Supabase PostgreSQL). Specific technical concerns: VNPay sandbox vs production behavioral differences, '
    'Supabase private bucket signed URL TTL 1 hour, JWT ES256 key dependencies, Supabase free tier rate limits '
    'during testing. Follow IEEE 829-2008 format with Risk ID, description, category, likelihood (H/M/L), '
    'impact (H/M/L), risk level, mitigation strategy, and contingency plan.')])
set_col_widths(t, [3.5, 11])

doc.add_paragraph()
body(doc, '[Ảnh chụp màn hình cuộc hội thoại với AI – đính kèm bên dưới]')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[ Screenshot 3: Prompt và Response từ AI cho Mục 15 ]')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.italic = True
r.font.color.rgb = RGBColor(0x80,0x80,0x80)

doc.add_paragraph()
h3(doc, 'Bảng đánh giá kết quả AI:')
t = mktbl(doc, 6, 3)
th(t, ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'])
td(t, [
    ('Đúng chuẩn IEEE 829-2008',            '4/5',
     'Format Risk ID, L, I, Level, Mitigation, Contingency đúng chuẩn; thiếu cột Risk Category'),
    ('Phù hợp với context Bee Academy',     '4/5',
     'AI nhận diện được rủi ro kỹ thuật đặc thù (signed URL, VNPay sandbox) sau khi được cung cấp context'),
    ('Đủ chi tiết, có thể dùng ngay',       '3/5',
     'Thiếu rủi ro về business (legal compliance payment, PO sign-off delay) và rủi ro test data'),
    ('Cần chỉnh sửa nhiều hay ít',          '3/5',
     'Chỉnh sửa vừa phải: thêm 3 rủi ro mới, xóa 2 rủi ro generic không phù hợp'),
    ('Tổng',                                 '14/20', ''),
])
set_col_widths(t, [5, 2, 7.5])

h3(doc, 'Những gì đã chỉnh sửa so với output AI:')
for item in [
    'Thêm R07 (Test data corrupt sau payment flow) – AI bỏ sót hoàn toàn rủi ro này',
    'Thêm R08 (Signed URL TTL expire mid-session) – AI đề cập nhưng không có contingency plan',
    'Thêm R09 (Supabase RLS misconfiguration → video public) – rủi ro security quan trọng',
    'Xóa "Database backup failure" – không phù hợp với context dùng Supabase managed service',
    'Điều chỉnh mức L của R03 (VNPay sandbox) từ M→H vì đã gặp issue thực tế khi dev',
]:
    bul(doc, item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PHẦN C – CÂU HỎI THẢO LUẬN
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẦN C – CÂU HỎI THẢO LUẬN')
r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.font.bold = True
doc.add_paragraph()

# --- Q1 ---
h1(doc, 'CÂU HỎI 1')
body(doc,
    'Khi sử dụng AI để tạo Test Plan, AI thường mạnh ở những mục nào và yếu ở những mục nào '
    'trong chuẩn IEEE 829-2008? Giải thích tại sao có sự khác biệt đó. Một tester chuyên nghiệp '
    'cần bổ sung gì mà AI không thể tự làm được?', justify=False)
p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()

h2(doc, 'Trả lời:')

body(doc,
    'Qua trải nghiệm thực tế khi làm bài, AI (Claude Sonnet 4.6) thể hiện rõ điểm mạnh và '
    'điểm yếu khác nhau tùy từng mục của IEEE 829-2008.')

h3(doc, 'Những mục AI làm tốt:')
body(doc,
    'AI mạnh nhất ở các mục có cấu trúc cố định, theo template rõ ràng: Mục 1 (Test Plan '
    'Identifier), Mục 9 (Test Deliverables), Mục 12 (Responsibilities – RACI matrix), Mục 14 '
    '(Schedule – Gantt chart), Mục 16 (Approvals). Lý do là những mục này về cơ bản chỉ cần '
    'điền thông tin vào template có sẵn; AI được huấn luyện trên hàng nghìn tài liệu IEEE, '
    'nên nó nắm vững cấu trúc bắt buộc. Ví dụ, khi yêu cầu tạo RACI matrix, AI trả về '
    'đúng format R/A/C/I với 5 roles chuẩn ngay lần đầu.')

body(doc,
    'AI cũng làm tốt Mục 6 (Approach) khi được cung cấp context đầy đủ về tech stack và '
    'business domain. Sau khi tôi mô tả chi tiết về JWT ES256, Supabase Storage, VNPay payment '
    'flow, AI tạo được test levels và entry/exit criteria khá chính xác, chỉ cần điều chỉnh nhỏ.')

h3(doc, 'Những mục AI còn yếu:')
body(doc,
    'AI yếu nhất ở Mục 4 (Features to be Tested) và Mục 7 (Pass/Fail Criteria). Đối với '
    'Mục 4, AI có xu hướng liệt kê các feature theo tên module một cách generic (Login, '
    'Register, View Course...) mà không hiểu các ràng buộc business thực tế. Ví dụ, AI không '
    'biết rằng nút "Thêm vào giỏ hàng" trong Bee Academy chỉ khả dụng khi đã login, hoặc '
    'việc giáo viên submit khóa học phải qua flow DRAFT→PENDING→Admin duyệt. Những business '
    'rule này chỉ có trong CLAUDE.md và UseCase v6.5 – tài liệu nội bộ mà AI không có access.')

body(doc,
    'Mục 15 (Risks) là điểm yếu đáng chú ý nhất. AI thường đề xuất rủi ro generic '
    '(build failure, resource shortage, environment instability) – những rủi ro mà mọi '
    'dự án web đều gặp. AI bỏ sót hoàn toàn các rủi ro đặc thù của Bee Academy: JWT ES256 '
    'key rotation có thể làm tất cả user bị logout đồng loạt, Supabase signed URL TTL 1 giờ '
    'có thể expire mid-test-session, hay việc private video bucket bị misconfigure RLS sẽ '
    'gây lộ nội dung có phí. Đây là những rủi ro "chỉ người trong dự án mới biết".')

h3(doc, 'Điều tester chuyên nghiệp cần bổ sung:')
for item in [
    'Hiểu domain & business context: AI không thể đọc giữa các dòng của UseCase, biết '
     'stakeholder nào quan tâm điều gì, hay đâu là "killer feature" cần test kỹ nhất.',
    'Judgment call về priority và risk: Quyết định pass rate threshold 95% hay 90%, '
     'hay việc một lỗi Minor trên payment page thực chất là Critical – AI không có '
     'context để ra quyết định này.',
    'Thương lượng với stakeholders: Test Plan cần phản ánh thỏa thuận thực sự giữa '
     'QA, Dev và Business. AI tạo ra tài liệu một chiều, không qua negotiation.',
    'Verify feasibility: AI đề xuất lịch test 1 tháng với 3 tester nhưng không biết '
     'team còn đang làm việc khác, hay môi trường Supabase free tier có rate limit.',
    'Cập nhật liên tục: Test Plan phải thay đổi theo project evolution. AI chỉ tạo '
     'snapshot tại một thời điểm, không tự cập nhật khi requirement thay đổi.',
]:
    bul(doc, item)

# --- Q2 ---
h1(doc, 'CÂU HỎI 2')
body(doc,
    'Mục 15 (Risks & Contingencies) yêu cầu tester phải suy nghĩ về những gì có thể sai. '
    'So sánh danh sách rủi ro do AI đề xuất với những rủi ro bạn tự nghĩ ra dựa trên kiến '
    'thức môn học. AI có bỏ sót loại rủi ro nào không? Từ đó, rút ra bài học về giới hạn '
    'của AI trong Test Management.', justify=False)

doc.add_paragraph()

h2(doc, 'Trả lời:')

h3(doc, 'So sánh rủi ro AI đề xuất vs. rủi ro tự phân tích:')

body(doc,
    'Khi tôi yêu cầu AI liệt kê rủi ro cho Bee Academy, danh sách ban đầu AI trả về '
    'bao gồm: (1) Unstable test environment, (2) Resource shortage, (3) Requirements '
    'change, (4) Build delays, (5) Missing test data. Đây đều là những rủi ro kinh điển '
    'trong mọi dự án phần mềm – đúng nhưng chưa đủ.')

body(doc,
    'Khi tự phân tích dựa trên kiến thức môn SWT301 và kiến trúc thực tế của Bee Academy '
    'từ CLAUDE.md, tôi phát hiện thêm các loại rủi ro mà AI đã bỏ sót hoàn toàn:')

h3(doc, 'Loại rủi ro AI bỏ sót – Technical Risk đặc thù:')
for item in [
    'R02 – JWT ES256 Key Dependency: Bee Academy dùng ECDSA P-256, khác với HS256 thông thường. '
     'Nếu JWKS endpoint của Supabase fetch thất bại khi backend khởi động, ES256 verifier = null và '
     'toàn bộ JWT sẽ fail authentication. Đây là rủi ro duy nhất của Supabase JWT implementation, '
     'không có trong bất kỳ checklist generic nào.',
    'R08 – Signed URL TTL Expiry: Video trong private Supabase bucket được truy cập qua signed URL '
     'TTL 1 giờ. Nếu test session kéo dài > 1 giờ, URL hết hạn và tester nhầm tưởng đây là bug '
     'video playback. AI không biết điều này vì không có access vào code của Lesson.videoStoragePath.',
    'R09 – Supabase RLS Misconfiguration: Row Level Security của Supabase nếu không được set đúng '
     'cho bucket course-videos (private), video có thể bị public access mà không cần auth. '
     'Đây là security risk nghiêm trọng mà chỉ tester hiểu storage architecture mới nhận ra.',
]:
    bul(doc, item)

h3(doc, 'Loại rủi ro AI bỏ sót – Business Risk:')
for item in [
    'Risk về VNPay behavioral difference: Qua thực tế dev, VNPay webhook không reliable trong '
     'local environment (ngay cả với ngrok tunneling). AI không biết điều này vì chỉ biết '
     '"payment integration risk" theo nghĩa chung chung.',
    'Risk về DEV_MODE=true trong production: CLAUDE.md ghi rõ DEV_MODE=true trong backend/.env '
     '– nếu quên đổi khi deploy, OTP sẽ log ra console thay vì gửi email thật, đây là '
     'security vulnerability nghiêm trọng. AI không phân tích được config file cụ thể này.',
]:
    bul(doc, item)

h3(doc, 'Bài học rút ra về giới hạn của AI trong Test Management:')

body(doc,
    'AI hoạt động dựa trên pattern matching từ training data. Nó giỏi tạo ra những gì "thường '
    'gặp" trong tài liệu kiểm thử, nhưng không thể suy luận từ code base, architecture '
    'decision, hay business constraint đặc thù của từng project.')

body(doc,
    'Giới hạn cốt lõi của AI trong Test Management là: AI không có "project context" – '
    'nó không đọc được CLAUDE.md, UseCase v6.5, hay biết rằng DEV_MODE=true trong .env '
    'là một time bomb chờ phát nổ khi go-live. Một Test Manager giỏi phải kết hợp '
    'kiến thức framework (IEEE 829-2008, ISTQB) với deep understanding về hệ thống '
    'cụ thể đang test. AI có thể hỗ trợ phần framework, nhưng phần domain knowledge '
    'và critical thinking vẫn phải do con người đảm nhận.')

body(doc,
    'Kết luận thực tiễn: Workflow hiệu quả nhất là dùng AI để tạo "80% foundation" '
    '(structure, template, generic risks, standard criteria), sau đó tester chuyên nghiệp '
    'review, bổ sung project-specific knowledge, và validate feasibility. Không nên '
    'copy-paste output AI trực tiếp vào deliverable mà không qua critical review – '
    'đây là điều Lab 4 muốn chúng ta học được.')

# ─── References ──────────────────────────────────────────────────────────────
doc.add_page_break()
h1(doc, 'TÀI LIỆU THAM KHẢO')
refs_list = [
    'IEEE 829-2008 – IEEE Standard for Software and System Test Documentation',
    'ISTQB Foundation Level Syllabus v4.0 – Chapter 5: Test Management',
    'Foundations of Software Testing (Craig & Jaskiel) – Chapter 9',
    'ISTQB Glossary – https://glossary.istqb.org',
    'Bee Academy CLAUDE.md – Bee Academy Architecture & Project Documentation',
    'Bee Academy UseCase v6.5 – BEE ACADEMY.md (48 use cases, 9 modules)',
    'Supabase Documentation – Storage Buckets, RLS, JWT Authentication',
    'Spring Boot Security Documentation – JWT ES256 / ECDSA P-256 Configuration',
    'VNPay Developer Documentation – Sandbox Testing Guide',
]
for i, r in enumerate(refs_list, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'[{i}] {r}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = r'd:\Vo_Van_Thanh_Dat\HocAI\Lab4_BeeAcademy_TestPlan.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
