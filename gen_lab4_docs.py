# -*- coding: utf-8 -*-
"""Sinh 3 file deliverable cho Lab 4 (SWT301) - dự án Bee Academy.
1. Test Plan tiếng Anh (Phần A - 16 mục IEEE 829-2008)
2. Test Plan tiếng Việt (Phần A - 16 mục)
3. Phần B + C tiếng Việt (AI Interaction Log + Câu hỏi thảo luận)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r'd:\Vo_Van_Thanh_Dat\HocAI'
PRIMARY = RGBColor(0xF5, 0xA6, 0x23)   # bee-yellow
DARK = RGBColor(0x1A, 0x1A, 0x2E)      # bee-dark
HEADER_FILL = '1A1A2E'
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def style_doc(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')


def add_cover(doc, lines, meta_rows):
    for txt, size, bold, color, space in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(space)
    doc.add_paragraph()
    add_table(doc, meta_rows, header=False, widths=[2.0, 4.2])


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = DARK
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'F5A623')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x44)
    return p


def add_p(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, rows, header=True, widths=None):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(10)
            if header and i == 0:
                run.bold = True
                run.font.color.rgb = HEADER_TEXT
                set_cell_bg(cell, HEADER_FILL)
            if widths:
                cell.width = Inches(widths[j])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def render(doc, blocks):
    for b in blocks:
        kind = b[0]
        if kind == 'h1':
            add_h1(doc, b[1])
        elif kind == 'h2':
            add_h2(doc, b[1])
        elif kind == 'p':
            add_p(doc, b[1])
        elif kind == 'bullet':
            add_bullet(doc, b[1])
        elif kind == 'table':
            add_table(doc, b[1])
        elif kind == 'spacer':
            doc.add_paragraph()
        elif kind == 'pagebreak':
            doc.add_page_break()


# =====================================================================
# ============== PHẦN A - TIẾNG VIỆT (TEST PLAN) ======================
# =====================================================================
VI_A = [
    ('h1', 'GIỚI THIỆU DỰ ÁN – BEE ACADEMY'),
    ('p', 'Bee Academy là nền tảng học trực tuyến dành cho học sinh THCS (lớp 6–9), cung cấp các khóa học video theo môn học. Hệ thống gồm 4 nhóm người dùng:'),
    ('bullet', 'Học sinh: Xem và mua khóa học, làm quiz theo chương, theo dõi tiến độ, nhận chứng chỉ hoàn thành'),
    ('bullet', 'Giáo viên: Tạo và quản lý khóa học, upload video bài giảng, cấu hình quiz, xem doanh thu'),
    ('bullet', 'Phụ huynh: Liên kết tài khoản với học sinh, theo dõi tiến độ học tập, liên hệ giáo viên'),
    ('bullet', 'Admin: Duyệt khóa học, quản lý tài khoản người dùng, xác nhận payout GV, tạo báo cáo'),
    ('h2', 'Thông tin phiên bản'),
    ('table', [
        ['Thông tin', 'Chi tiết'],
        ['Tên hệ thống', 'Bee Academy – Online Learning Platform'],
        ['Phiên bản kiểm thử', 'v1.0.0'],
        ['Môi trường', 'Web application (React 19 + Vite 6 frontend / Spring Boot 3.2 backend / PostgreSQL – Supabase)'],
        ['Ngày bắt đầu test', '12/06/2026'],
        ['Ngày kết thúc test', '15/07/2026'],
        ['Team size', '3 tester + 1 test lead'],
    ]),

    ('h1', 'MỤC 1 – TEST PLAN IDENTIFIER'),
    ('p', 'Tài liệu này định danh Test Plan chính thức (Master Test Plan) cho dự án Bee Academy phiên bản v1.0.0, được lập theo chuẩn IEEE 829-2008.'),
    ('table', [
        ['Trường', 'Giá trị'],
        ['Document ID', 'TP-BEEACADEMY-2026-001'],
        ['Project Name', 'Bee Academy – Online Learning Platform'],
        ['System Version', 'v1.0.0'],
        ['Test Plan Version', '1.0'],
        ['Author', 'Võ Văn Thành Đạt'],
        ['Created Date', '12/06/2026'],
        ['Review Date', '19/06/2026'],
        ['Status', 'Draft'],
    ]),

    ('h1', 'MỤC 2 – INTRODUCTION'),
    ('h2', '2.1  Mục đích (Purpose)'),
    ('p', 'Tài liệu Test Plan này xác định phạm vi, phương pháp tiếp cận, nguồn lực và lịch trình kiểm thử cho hệ thống Bee Academy v1.0.0 – nền tảng học trực tuyến dành cho học sinh THCS (lớp 6–9) tại Việt Nam. Mục tiêu là đảm bảo hệ thống đáp ứng yêu cầu chức năng và phi chức năng trước khi phát hành chính thức.'),
    ('h2', '2.2  Phạm vi áp dụng (Scope)'),
    ('p', 'Test Plan này áp dụng cho giai đoạn System Testing và UAT của Bee Academy v1.0.0, bao gồm 8 module chính: Authentication & Authorization, Course Catalog, Enrollment & Payment, Student Dashboard, Teacher Portal, Quiz & Assessment, Admin Portal, và Parent Portal. Môi trường kiểm thử là React 19 + Vite 6 (frontend) kết hợp Spring Boot 3.2 (backend) với cơ sở dữ liệu PostgreSQL thông qua Supabase.'),
    ('h2', '2.3  Đối tượng đọc (Intended Audience)'),
    ('bullet', 'Test Lead và QA Engineers – thực thi các hoạt động kiểm thử'),
    ('bullet', 'Development Team – fix defects và hỗ trợ môi trường'),
    ('bullet', 'Project Manager – theo dõi tiến độ và ra quyết định'),
    ('bullet', 'Product Owner / Stakeholders – phê duyệt UAT và sign-off release'),
    ('h2', '2.4  Tài liệu tham chiếu (References)'),
    ('bullet', 'IEEE 829-2008 – IEEE Standard for Software and System Test Documentation'),
    ('bullet', 'CLAUDE.md – Bee Academy Architecture & Project Documentation'),
    ('bullet', 'BEE ACADEMY.md – UseCase v6.5 (48 Use Cases, 9 modules, 7 actors)'),
    ('bullet', 'ISTQB Foundation Level Syllabus – Chapter 5: Test Management'),
    ('bullet', 'Spring Boot 3.2 Security Documentation – JWT ES256 / ECDSA P-256'),
    ('bullet', 'Supabase Documentation – Storage Buckets & RLS Policies'),

    ('h1', 'MỤC 3 – TEST ITEMS'),
    ('p', 'Dưới đây là danh sách các module/component của Bee Academy v1.0.0 được đưa vào phạm vi kiểm thử, kèm theo phiên bản và trạng thái sẵn sàng.'),
    ('table', [
        ['Item ID', 'Module', 'Version', 'Mô tả ngắn', 'Trạng thái'],
        ['TI-01', 'Authentication & Authorization', 'v1.0.0', 'Đăng nhập email/pw, OTP, Google OAuth (JWT ES256 + Supabase GoTrue)', 'Ready'],
        ['TI-02', 'Course Catalog & Detail', 'v1.0.0', 'Duyệt, tìm kiếm, lọc khóa học; xem chi tiết và bài học thử', 'Ready'],
        ['TI-03', 'Enrollment & Payment', 'v1.0.0', 'Giỏ hàng, thanh toán VNPay/MoMo, xem lịch sử đơn hàng', 'In Dev'],
        ['TI-04', 'Student Dashboard', 'v1.0.0', 'Hồ sơ cá nhân, tài khoản, avatar upload, yêu thích', 'Ready'],
        ['TI-05', 'Teacher Portal', 'v1.0.0', 'CRUD khóa học/chương/bài; upload video (private) + tài liệu; submit duyệt', 'Partial'],
        ['TI-06', 'Quiz & Assessment System', 'v1.0.0', 'Config quiz chương, ngân hàng câu hỏi, làm bài + chấm điểm JSONB snapshot', 'Ready'],
        ['TI-07', 'Admin Portal', 'v1.0.0', 'Duyệt khóa học, quản lý người dùng, báo cáo doanh thu, xác nhận payout', 'Partial'],
        ['TI-08', 'Parent Portal', 'v1.0.0', 'Liên kết học sinh, theo dõi tiến độ, liên hệ giáo viên', 'Ready'],
    ]),

    ('h1', 'MỤC 4 – FEATURES TO BE TESTED'),
    ('p', 'Bảng dưới liệt kê tối thiểu 18 tính năng nằm trong phạm vi kiểm thử, phân bổ đều trên 8 module chính.'),
    ('table', [
        ['Feature ID', 'Tên tính năng', 'Module', 'Priority', 'Mô tả ngắn'],
        ['F-01', 'Đăng nhập email/password', 'Auth', 'High', 'Nhập đúng/sai credentials; kiểm tra JWT trả về'],
        ['F-02', 'Đăng ký tài khoản với OTP email', 'Auth', 'High', 'Flow gửi OTP → xác minh → tạo profile STUDENT'],
        ['F-03', 'Đăng nhập Google OAuth (ES256 JWT)', 'Auth', 'High', 'Google consent → Supabase callback → sync profile'],
        ['F-04', 'Phân quyền theo role (RBAC)', 'Auth', 'High', 'Student/Teacher/Parent/Admin đúng route guards'],
        ['F-05', 'Refresh token & session timeout', 'Auth', 'Medium', 'Token hết hạn tự refresh; logout revoke token'],
        ['F-06', 'Duyệt danh sách khóa học + bộ lọc', 'Course', 'High', 'Filter category, giá, cấp lớp; phân trang'],
        ['F-07', 'Xem chi tiết và bài học thử miễn phí', 'Course', 'High', 'isFree=true cho phép xem; private video cần enroll'],
        ['F-08', 'Thêm vào giỏ – yêu cầu đăng nhập', 'Payment', 'High', 'Chưa login → redirect /login với state.from'],
        ['F-09', 'Thanh toán VNPay/MoMo (sandbox)', 'Payment', 'High', 'Tạo order → redirect payment → webhook → enroll'],
        ['F-10', 'Xem lịch sử đơn hàng', 'Dashboard', 'Medium', 'Danh sách orders, trạng thái, ngày mua'],
        ['F-11', 'Cập nhật hồ sơ và upload avatar', 'Dashboard', 'Medium', 'Đổi tên, bio; upload ảnh lên Supabase Storage'],
        ['F-12', 'Giáo viên tạo/sửa/xóa khóa học', 'Teacher', 'High', 'CRUD với chapter/lesson; trạng thái DRAFT→PENDING'],
        ['F-13', 'Upload video bài giảng (private bucket)', 'Teacher', 'High', 'Tải lên Supabase private; lưu storagePath'],
        ['F-14', 'Submit khóa học để Admin duyệt', 'Teacher', 'High', 'Đổi trạng thái DRAFT→PENDING; admin nhận notification'],
        ['F-15', 'Config quiz chương + ngân hàng câu hỏi', 'Quiz', 'High', 'CRUD câu hỏi; set số câu, thời gian, passing score'],
        ['F-16', 'Học sinh làm quiz + xem kết quả', 'Quiz', 'High', 'Random pick từ ngân hàng; JSONB snapshot; chấm điểm'],
        ['F-17', 'Admin duyệt / từ chối / yêu cầu sửa', 'Admin', 'High', 'Review content; approve→PUBLISHED; reject kèm lý do'],
        ['F-18', 'Phụ huynh liên kết và theo dõi tiến độ', 'Parent', 'Medium', 'Gửi lời mời email; sau khi ACTIVE xem progress'],
    ]),

    ('h1', 'MỤC 5 – FEATURES NOT TO BE TESTED'),
    ('p', 'Các tính năng sau được loại trừ khỏi phạm vi kiểm thử trong v1.0.0, kèm lý do cụ thể:'),
    ('table', [
        ['Tính năng loại trừ', 'Lý do loại trừ'],
        ['Penetration Testing / Security Audit', 'Yêu cầu chuyên gia bảo mật riêng; sẽ thực hiện trước go-live production'],
        ['Load Testing / Performance Testing', 'Cần môi trường staging cấu hình giống production; không khả thi trong môi trường dev local'],
        ['Cross-browser (Safari, Edge)', 'Scope v1.0 chỉ hỗ trợ Chrome + Firefox; Safari/Edge được test ở sprint tiếp theo'],
        ['Mobile responsiveness (< 768px)', 'Version mobile sẽ được test riêng trong Sprint 3 khi hoàn thiện responsive design'],
        ['Tích hợp MoMo thực (production API)', 'Chỉ có VNPay sandbox; MoMo production API cần ký hợp đồng và được test riêng'],
        ['Chức năng Chứng chỉ (UC42-43)', 'Chưa được phát triển trong v1.0.0; dự kiến Sprint 4'],
        ['Tin nhắn / Chat AI với GV (UC20-21)', 'Module chat chưa phát triển; cần tích hợp thêm WebSocket hoặc third-party'],
        ['Chức năng TK Ngân hàng GV (UC45-46)', 'Chưa hoàn thiện backend; Admin payout manual sẽ test khi feature complete'],
    ]),

    ('h1', 'MỤC 6 – APPROACH (TEST STRATEGY)'),
    ('p', 'Chiến lược kiểm thử của Bee Academy v1.0.0 áp dụng kiểm thử đa tầng (multi-level), kết hợp kỹ thuật black-box và white-box, hướng đến đảm bảo chất lượng toàn diện trước khi phát hành.'),
    ('h2', '6.1  Test Levels (Cấp độ kiểm thử)'),
    ('table', [
        ['Cấp độ', 'Phụ trách', 'Kỹ thuật', 'Công cụ'],
        ['Unit Testing', 'Dev Team', 'White-box, Statement Coverage (≥80%)', 'JUnit 5 / Jest'],
        ['Integration Testing', 'Dev + QA', 'API Testing, Black-box, Contract Test', 'Postman / REST Assured'],
        ['System Testing', 'QA Team', 'Black-box: EP, BVA; Exploratory', 'Playwright / JIRA'],
        ['Acceptance Testing', 'PO + Stakeholders', 'Scenario-based, User Acceptance', 'Manual / TestRail'],
    ]),
    ('h2', '6.2  Test Types (Loại hình kiểm thử)'),
    ('table', [
        ['Loại hình', 'Mục tiêu', 'Phạm vi áp dụng'],
        ['Functional Testing', 'Kiểm tra logic nghiệp vụ đúng theo UseCase v6.5', 'Tất cả 18 features trong Mục 4'],
        ['Security Testing', 'Kiểm tra JWT authentication, authorization bypass, SQL injection', 'Module Auth, Teacher Portal, Admin Portal'],
        ['Usability Testing', 'Kiểm tra UX/UI trực quan, responsive desktop', 'Landing Page, Course Detail, Student Dashboard'],
        ['Regression Testing', 'Đảm bảo bug fix không phá vỡ tính năng hiện có', 'Sau mỗi lần fix defect Critical/Major'],
    ]),
    ('h2', '6.3  Test Design Techniques'),
    ('bullet', 'Equivalence Partitioning (EP) – phân hoạch tương đương cho input validation (email, password, giá tiền)'),
    ('bullet', 'Boundary Value Analysis (BVA) – kiểm tra biên cho price range, grade filter, file size upload'),
    ('bullet', 'Decision Table – cho flow thanh toán (VNPay callback: success/fail/cancel/timeout)'),
    ('bullet', 'State Transition – cho vòng đời khóa học (DRAFT→PENDING→PUBLISHED/REJECTED)'),
    ('bullet', 'Error Guessing – JWT expired, invalid OTP, Supabase signed URL hết hạn'),
    ('bullet', 'Exploratory Testing – tìm lỗi không có trong test case bằng session-based exploration'),
    ('h2', '6.4  Entry & Exit Criteria'),
    ('table', [
        ['Cấp độ', 'Entry Criteria (Điều kiện bắt đầu)', 'Exit Criteria (Điều kiện kết thúc)'],
        ['System Testing', 'Build deploy thành công lên test env; Unit test pass ≥80%; Test cases & data đã chuẩn bị', 'Test case pass rate ≥95%; Không có Critical defect còn mở; Test Summary Report được tạo'],
        ['Acceptance Testing', 'System Testing đã pass Exit Criteria; Regression Testing hoàn thành; Môi trường UAT sẵn sàng', 'Tất cả UAT scenarios PASS; PO ký sign-off; Release checklist hoàn chỉnh'],
    ]),

    ('h1', 'MỤC 7 – ITEM PASS/FAIL CRITERIA'),
    ('h2', '7.1  Pass Criteria (Tiêu chí PASS)'),
    ('bullet', 'Tỷ lệ test case PASS ≥ 95% trên tổng số test cases đã thực thi'),
    ('bullet', 'Không có defect mức Critical hoặc Blocker còn ở trạng thái Open'),
    ('bullet', 'Không có defect mức Major còn Open sau ngày D-3 trước release'),
    ('bullet', 'Tất cả defect Major đã có workaround được document trong JIRA'),
    ('bullet', 'Module Authentication: không có bất kỳ lỗi Security nào (JWT bypass, SQL injection)'),
    ('bullet', 'Module Payment: toàn bộ luồng thanh toán VNPay sandbox hoàn thành thành công'),
    ('bullet', 'Test Summary Report được Test Lead phê duyệt'),
    ('h2', '7.2  Fail Criteria (Tiêu chí FAIL – dừng release)'),
    ('bullet', 'Tỷ lệ test case PASS < 90%'),
    ('bullet', 'Tồn tại ít nhất 1 defect Critical/Blocker chưa được fix'),
    ('bullet', 'Module Authentication có bất kỳ lỗi security nào còn Open'),
    ('bullet', 'Luồng thanh toán VNPay không hoàn thành được (payment flow broken)'),
    ('bullet', 'Dữ liệu user bị mất hoặc bị ghi đè sai sau khi thực hiện giao dịch'),
    ('bullet', 'Hệ thống crash hoặc không khởi động được trong môi trường test'),
    ('h2', '7.3  Defect Severity Definition'),
    ('table', [
        ['Mức độ', 'Định nghĩa', 'Ví dụ trong Bee Academy'],
        ['Critical / Blocker', 'Hệ thống không sử dụng được; không có workaround', 'Không thể đăng nhập; payment crash; data loss'],
        ['Major', 'Chức năng quan trọng bị lỗi; có workaround tạm thời', 'Quiz không chấm điểm đúng; video không phát được'],
        ['Minor', 'Chức năng hoạt động nhưng có lỗi nhỏ; ít ảnh hưởng', 'Text sai chính tả; UI misalignment nhỏ'],
        ['Cosmetic', 'Lỗi giao diện không ảnh hưởng chức năng', 'Icon hiển thị sai màu; spacing không đều'],
    ]),

    ('h1', 'MỤC 8 – SUSPENSION CRITERIA AND RESUMPTION REQUIREMENTS'),
    ('p', 'Bảng dưới mô tả các tình huống phải tạm dừng kiểm thử và điều kiện cần đáp ứng để tiếp tục.'),
    ('table', [
        ['Tình huống', 'Điều kiện tạm dừng', 'Điều kiện tiếp tục'],
        ['Build/Deploy thất bại', 'Môi trường test không deploy được phiên bản mới; error rate > 50%', 'Build pass; smoke test 5 tính năng cốt lõi thành công; Dev xác nhận'],
        ['Critical Blocker xuất hiện', 'Phát hiện ≥1 defect Critical ảnh hưởng module đang test', 'Dev fix và merge; QA verify fix; Regression test module đó pass'],
        ['Môi trường test down', 'Backend hoặc Supabase unreachable > 30 phút liên tục', 'Môi trường được restore; health check pass; Test Lead xác nhận'],
        ['Test data bị corrupt', 'Dữ liệu test bị sai, thiếu hoặc bị ghi đè do lỗi script/migration', 'Restore lại test data từ backup; QA xác nhận data integrity'],
    ]),
    ('p', 'Khi tạm dừng, Test Lead phải ghi nhận nguyên nhân vào Daily Report, cập nhật trạng thái trong JIRA và thông báo đến Project Manager trong vòng 1 giờ.'),

    ('h1', 'MỤC 9 – TEST DELIVERABLES'),
    ('p', 'Danh sách tài liệu đầu ra (deliverables) được tạo ra trong suốt quá trình kiểm thử:'),
    ('table', [
        ['Giai đoạn', 'Tài liệu', 'Người chịu trách nhiệm', 'Deadline'],
        ['Pre-testing', 'Master Test Plan (tài liệu này)', 'Test Lead', '12/06/2026'],
        ['Pre-testing', 'Test Cases (Excel / TestRail)', 'QA Team', '17/06/2026'],
        ['Pre-testing', 'Test Data Script (SQL + seed data)', 'QA Team', '16/06/2026'],
        ['Pre-testing', 'Test Environment Setup Guide', 'Dev + QA', '15/06/2026'],
        ['During testing', 'Daily Test Progress Report', 'Test Lead', 'Hàng ngày'],
        ['During testing', 'Defect Report (JIRA tickets)', 'QA Team', 'Khi phát hiện'],
        ['During testing', 'Test Execution Log', 'QA Team', 'Hàng ngày'],
        ['During testing', 'Change Request Log (nếu req thay đổi)', 'Test Lead', 'Khi phát sinh'],
        ['Post-testing', 'Test Summary Report', 'Test Lead', '14/07/2026'],
        ['Post-testing', 'Defect Metrics & Quality Report', 'Test Lead', '14/07/2026'],
    ]),

    ('h1', 'MỤC 10 – TESTING TASKS'),
    ('p', 'Bảng dưới phân rã các hoạt động kiểm thử thành các task cụ thể, có thứ tự thực hiện và dependency rõ ràng.'),
    ('table', [
        ['Task ID', 'Tên Task', 'Predecessor', 'Effort', 'Role', 'Start', 'End'],
        ['T01', 'Phân tích requirements & UseCase v6.5', '–', '8h', 'Test Lead', '12/06', '13/06'],
        ['T02', 'Thiết kế test cases (18 features)', 'T01', '24h', 'QA Team', '14/06', '17/06'],
        ['T03', 'Chuẩn bị test data & seed scripts', 'T01', '8h', 'QA Team', '14/06', '16/06'],
        ['T04', 'Cài đặt và cấu hình môi trường test', 'T01', '4h', 'Dev + QA', '12/06', '14/06'],
        ['T05', 'System Test – Auth + Course + Dashboard', 'T02,T03,T04', '20h', 'QA Team', '18/06', '23/06'],
        ['T06', 'System Test – Payment + Teacher Portal', 'T05', '20h', 'QA Team', '24/06', '30/06'],
        ['T07', 'System Test – Quiz + Admin + Parent', 'T06', '20h', 'QA Team', '01/07', '07/07'],
        ['T08', 'Regression Testing (sau khi fix defects)', 'T07', '16h', 'QA Team', '08/07', '10/07'],
        ['T09', 'User Acceptance Testing (UAT)', 'T08', '16h', 'PO + Stakeholders', '11/07', '13/07'],
        ['T10', 'Viết Test Summary Report & Sign-off', 'T09', '4h', 'Test Lead', '14/07', '15/07'],
    ]),

    ('h1', 'MỤC 11 – ENVIRONMENTAL NEEDS'),
    ('h2', '11.1  Yêu cầu phần cứng (Hardware Requirements)'),
    ('table', [
        ['Thành phần', 'Cấu hình tối thiểu'],
        ['Test Server (Backend)', 'CPU: 4 cores, RAM: 8GB, Storage: 100GB SSD; hoặc Cloud: Supabase + Railway'],
        ['Tester Workstation', 'CPU: 2 cores, RAM: 8GB, OS: Windows 11 / Ubuntu 22.04'],
        ['CI/CD Pipeline', 'GitHub Actions (2 workers, 4GB RAM per runner)'],
    ]),
    ('h2', '11.2  Yêu cầu phần mềm (Software Requirements)'),
    ('table', [
        ['Phần mềm', 'Phiên bản', 'Mục đích'],
        ['OS', 'Windows 11 / Ubuntu 22.04', 'Môi trường chạy test'],
        ['Node.js', '20 LTS', 'Build và chạy React 19 + Vite frontend'],
        ['Java JDK', '17', 'Runtime Spring Boot 3.2 backend'],
        ['Maven', '3.9+', 'Build tool backend'],
        ['PostgreSQL (Supabase)', '15', 'Cơ sở dữ liệu; sử dụng Supabase test project'],
        ['Chrome', 'Latest', 'Primary browser (System Test)'],
        ['Firefox', 'Latest', 'Secondary browser (Regression)'],
        ['Postman', 'v10+', 'API testing – Auth endpoints, JWT validation'],
        ['Playwright', '1.40+', 'UI automation – E2E test scenarios'],
        ['JIRA', 'Cloud', 'Defect tracking và test management'],
        ['Git / GitHub', 'Latest', 'Version control; CI/CD integration'],
    ]),
    ('h2', '11.3  Cấu hình môi trường test (Test Environment Configuration)'),
    ('bullet', 'Frontend: http://localhost:3000 (Vite dev server, port đã cố định trong vite.config.ts)'),
    ('bullet', 'Backend: http://localhost:8080 (Spring Boot, cổng cấu hình trong backend/.env)'),
    ('bullet', 'Database: Supabase PostgreSQL (dùng project test riêng, tách biệt production)'),
    ('bullet', 'Storage: Supabase Storage – 2 buckets: course-videos (private) + course-docs (public)'),
    ('bullet', 'JWT: ES256 (ECDSA P-256) verify qua Supabase JWKS endpoint'),
    ('bullet', 'CORS: cho phép origin http://localhost:3000'),
    ('bullet', 'DEV_MODE=true: OTP log ra console thay vì gửi email thật'),
    ('h2', '11.4  Test Data Requirements'),
    ('bullet', '5 tài khoản Student (valid credentials + edge cases: mật khẩu sai, tài khoản bị khóa)'),
    ('bullet', '3 tài khoản Teacher (1 có khóa học đã duyệt, 1 đang pending, 1 mới tạo)'),
    ('bullet', '2 tài khoản Parent (1 đã link với student, 1 chưa link)'),
    ('bullet', '1 tài khoản Admin'),
    ('bullet', '10 khóa học mẫu (DRAFT: 3, PENDING: 3, PUBLISHED: 4, REJECTED: 1)'),
    ('bullet', 'Test payment cards VNPay sandbox: success + fail + timeout scenarios'),
    ('bullet', '50 câu hỏi ngân hàng mẫu (10 câu/module, phân bổ Easy/Medium/Hard)'),
    ('bullet', 'File video mẫu (≤100MB, định dạng MP4) và tài liệu PDF (≤10MB) cho upload test'),

    ('h1', 'MỤC 12 – RESPONSIBILITIES'),
    ('p', 'Ma trận RACI phân công trách nhiệm cho từng hoạt động kiểm thử (R=Responsible, A=Accountable, C=Consulted, I=Informed):'),
    ('table', [
        ['Hoạt động', 'Test Lead', 'QA Engineer', 'Developer', 'PM', 'PO'],
        ['Viết Test Plan', 'R/A', 'C', 'C', 'I', 'I'],
        ['Thiết kế Test Cases', 'A', 'R', 'C', 'I', 'C'],
        ['Chuẩn bị Test Data', 'A', 'R', 'C', 'I', '–'],
        ['Cài đặt môi trường', 'A', 'R', 'R/A', 'I', '–'],
        ['Thực thi Testing', 'A', 'R', 'I', 'I', 'I'],
        ['Bug Report & Tracking', 'A', 'R', 'I', 'I', 'I'],
        ['Fix Defects', 'I', 'C', 'R/A', 'I', '–'],
        ['Test Summary Report', 'R/A', 'C', 'I', 'I', 'I'],
    ]),

    ('h1', 'MỤC 13 – STAFFING AND TRAINING NEEDS'),
    ('h2', '13.1  Thành phần nhân lực'),
    ('table', [
        ['Vai trò', 'Số lượng', 'Allocation', 'Kỹ năng yêu cầu', 'Đào tạo cần thiết'],
        ['Test Lead', '1', '100%', 'IEEE 829-2008; ISTQB Foundation; API testing; JIRA', 'Không cần thêm'],
        ['QA Engineer', '2', '100%', 'Manual testing; SQL cơ bản; Chrome DevTools; Postman', 'Playwright workshop (2 ngày); Supabase API basics (1 ngày)'],
        ['Dev (support)', '2', '20%', 'Unit test JUnit 5; code review; environment setup', 'Không cần thêm'],
    ]),
    ('h2', '13.2  Kế hoạch đào tạo (Training Plan)'),
    ('table', [
        ['Khóa đào tạo', 'Đối tượng', 'Thời lượng', 'Thời điểm'],
        ['Playwright E2E Automation Workshop', 'QA Engineers', '2 ngày', 'Tuần 1 (12-13/06/2026)'],
        ['Supabase Storage & JWT ES256 Overview', 'QA Team', '3 tiếng', 'Ngày 14/06/2026'],
    ]),

    ('h1', 'MỤC 14 – SCHEDULE'),
    ('h2', '14.1  Milestones'),
    ('table', [
        ['Milestone', 'Ngày', 'Mô tả'],
        ['MS-01', '12/06/2026', 'Kickoff, Test Plan được phê duyệt'],
        ['MS-02', '17/06/2026', 'Test Cases & Test Data ready; môi trường test sẵn sàng'],
        ['MS-03', '07/07/2026', 'System Testing hoàn thành (cả 8 modules)'],
        ['MS-04', '10/07/2026', 'Regression Testing hoàn thành'],
        ['MS-05', '13/07/2026', 'UAT hoàn thành; PO ký sign-off'],
        ['MS-06', '15/07/2026', 'Test Summary Report được nộp'],
    ]),
    ('h2', '14.2  Lịch tuần (Gantt Chart)'),
    ('table', [
        ['Phase / Activity', 'T1 (12-16/6)', 'T2 (17-23/6)', 'T3 (24-30/6)', 'T4 (1-10/7)', 'T5 (11-15/7)'],
        ['Planning & Test Design', 'XX', 'XX', '', '', ''],
        ['Environment Setup', 'XX', '', '', '', ''],
        ['System Test – Auth & Course', '', 'XX', '', '', ''],
        ['System Test – Pay & Teacher', '', '', 'XX', '', ''],
        ['System Test – Quiz,Admin,PH', '', '', '', 'XX', ''],
        ['Regression Testing', '', '', '', 'XX', ''],
        ['UAT & Test Closure', '', '', '', '', 'XX'],
    ]),

    ('h1', 'MỤC 15 – RISKS AND CONTINGENCIES'),
    ('p', 'Bảng dưới xác định các rủi ro kiểm thử, đánh giá mức độ (L×I) và kế hoạch giảm thiểu / dự phòng. (L=Likelihood: H/M/L  |  I=Impact: H/M/L)'),
    ('table', [
        ['Risk ID', 'Rủi ro', 'L', 'I', 'Level', 'Mitigation', 'Contingency'],
        ['R01', 'Môi trường test không ổn định (Supabase local dev, CORS errors)', 'M', 'H', 'High', 'Chuẩn bị môi trường staging riêng; docker-compose backup', 'Dùng Supabase hosted test project thay localhost'],
        ['R02', 'JWT ES256 key rotation làm tất cả user bị logout bất ngờ', 'L', 'H', 'Medium', 'Không rotate JWKS key trong quá trình test; ghi chú trong env setup', 'Restart backend và re-login; ghi log incident'],
        ['R03', 'VNPay sandbox không phản ánh đúng behavior production', 'H', 'M', 'High', 'Dùng cả test card success + fail; document sai lệch nếu có', 'Test payment flow với mock server nếu sandbox không ổn định'],
        ['R04', 'Requirements thay đổi trong khi test (UseCase v6.5 chưa frozen)', 'M', 'H', 'High', 'Freeze requirements từ T1; mọi thay đổi qua Change Request', 'Re-scope test cases; điều chỉnh schedule nếu cần'],
        ['R05', 'Thiếu nhân lực do nghỉ ốm hoặc lý do cá nhân', 'L', 'H', 'Medium', 'Cross-training giữa QA Engineers; tài liệu test cases rõ ràng', 'Test Lead tạm thời hỗ trợ thực thi; ưu tiên test features High'],
        ['R06', 'Build chậm trễ từ Dev Team ảnh hưởng schedule', 'M', 'M', 'Medium', 'Daily stand-up; CI/CD tự động; early integration merge', 'Ưu tiên test các module đã Ready trước; delay test module In Dev'],
        ['R07', 'Test data bị corrupt sau khi test payment flow', 'M', 'H', 'High', 'Dùng transaction rollback hoặc reset script sau mỗi payment test session', 'Restore test DB từ seed script; ghi nhận incident vào log'],
        ['R08', 'Supabase Storage signed URL hết hạn (TTL 1h) trong mid-test session', 'L', 'M', 'Low', 'Regenerate signed URL trước mỗi video access test; note TTL trong test case', 'Tăng TTL trong test env hoặc test video trong session < 1h'],
        ['R09', 'Private video bucket bị misconfigured → public access', 'L', 'H', 'Medium', 'Verify RLS policy trên Supabase dashboard trước khi test', 'Fix policy ngay; security incident report; re-test toàn bộ video access'],
    ]),

    ('h1', 'MỤC 16 – APPROVALS'),
    ('p', 'Tài liệu này có hiệu lực khi được ký phê duyệt bởi tất cả các bên liên quan dưới đây:'),
    ('table', [
        ['Vai trò', 'Họ và tên', 'Chữ ký', 'Ngày phê duyệt'],
        ['Test Lead', 'Võ Văn Thành Đạt', '', ''],
        ['Project Manager', 'Nguyễn Văn A', '', ''],
        ['Product Owner', 'Trần Thị B', '', ''],
        ['Dev Lead', 'Lê Văn C', '', ''],
    ]),
    ('p', 'Lưu ý: Tài liệu này ở trạng thái Draft. Chữ ký sẽ được thu thập sau khi hoàn thành vòng review lần 1 (Review Date: 19/06/2026).'),
]


# =====================================================================
# ============== PART A - ENGLISH (TEST PLAN) =========================
# =====================================================================
EN_A = [
    ('h1', 'PROJECT OVERVIEW – BEE ACADEMY'),
    ('p', 'Bee Academy is an online learning platform for Vietnamese lower-secondary students (grades 6–9), offering video courses by subject. The system serves four user groups:'),
    ('bullet', 'Student: Browse and purchase courses, take per-chapter quizzes, track progress, and receive completion certificates'),
    ('bullet', 'Teacher: Create and manage courses, upload lecture videos, configure quizzes, and view revenue'),
    ('bullet', 'Parent: Link to a student account, monitor learning progress, and contact teachers'),
    ('bullet', 'Admin: Approve courses, manage user accounts, confirm teacher payouts, and generate reports'),
    ('h2', 'Version Information'),
    ('table', [
        ['Information', 'Detail'],
        ['System Name', 'Bee Academy – Online Learning Platform'],
        ['Test Version', 'v1.0.0'],
        ['Environment', 'Web application (React 19 + Vite 6 frontend / Spring Boot 3.2 backend / PostgreSQL – Supabase)'],
        ['Test Start Date', '2026-06-12'],
        ['Test End Date', '2026-07-15'],
        ['Team Size', '3 testers + 1 test lead'],
    ]),

    ('h1', 'SECTION 1 – TEST PLAN IDENTIFIER'),
    ('p', 'This document identifies the official Master Test Plan (MTP) for the Bee Academy project, version v1.0.0, prepared in accordance with the IEEE 829-2008 standard.'),
    ('table', [
        ['Field', 'Value'],
        ['Document ID', 'TP-BEEACADEMY-2026-001'],
        ['Project Name', 'Bee Academy – Online Learning Platform'],
        ['System Version', 'v1.0.0'],
        ['Test Plan Version', '1.0'],
        ['Author', 'Vo Van Thanh Dat'],
        ['Created Date', '2026-06-12'],
        ['Review Date', '2026-06-19'],
        ['Status', 'Draft'],
    ]),

    ('h1', 'SECTION 2 – INTRODUCTION'),
    ('h2', '2.1  Purpose'),
    ('p', 'This Test Plan defines the scope, approach, resources, and schedule of the testing activities for Bee Academy v1.0.0 – an online learning platform for Vietnamese lower-secondary students (grades 6–9). The objective is to ensure the system meets its functional and non-functional requirements before official release.'),
    ('h2', '2.2  Scope'),
    ('p', 'This Test Plan applies to the System Testing and UAT phases of Bee Academy v1.0.0, covering eight core modules: Authentication & Authorization, Course Catalog, Enrollment & Payment, Student Dashboard, Teacher Portal, Quiz & Assessment, Admin Portal, and Parent Portal. The test environment uses React 19 + Vite 6 (frontend) together with Spring Boot 3.2 (backend) and a PostgreSQL database provided through Supabase.'),
    ('h2', '2.3  Intended Audience'),
    ('bullet', 'Test Lead and QA Engineers – execute the testing activities'),
    ('bullet', 'Development Team – fix defects and support the environment'),
    ('bullet', 'Project Manager – track progress and make decisions'),
    ('bullet', 'Product Owner / Stakeholders – approve UAT and sign off the release'),
    ('h2', '2.4  References'),
    ('bullet', 'IEEE 829-2008 – IEEE Standard for Software and System Test Documentation'),
    ('bullet', 'CLAUDE.md – Bee Academy Architecture & Project Documentation'),
    ('bullet', 'BEE ACADEMY.md – Use Case v6.5 (48 use cases, 9 modules, 7 actors)'),
    ('bullet', 'ISTQB Foundation Level Syllabus – Chapter 5: Test Management'),
    ('bullet', 'Spring Boot 3.2 Security Documentation – JWT ES256 / ECDSA P-256'),
    ('bullet', 'Supabase Documentation – Storage Buckets & RLS Policies'),

    ('h1', 'SECTION 3 – TEST ITEMS'),
    ('p', 'The following table lists the modules/components of Bee Academy v1.0.0 included in the testing scope, together with their version and readiness status.'),
    ('table', [
        ['Item ID', 'Module', 'Version', 'Brief Description', 'Status'],
        ['TI-01', 'Authentication & Authorization', 'v1.0.0', 'Email/password login, OTP, Google OAuth (JWT ES256 + Supabase GoTrue)', 'Ready'],
        ['TI-02', 'Course Catalog & Detail', 'v1.0.0', 'Browse, search, filter courses; view detail and preview lessons', 'Ready'],
        ['TI-03', 'Enrollment & Payment', 'v1.0.0', 'Cart, VNPay/MoMo payment, view order history', 'In Dev'],
        ['TI-04', 'Student Dashboard', 'v1.0.0', 'Profile, account, avatar upload, favorites', 'Ready'],
        ['TI-05', 'Teacher Portal', 'v1.0.0', 'CRUD courses/chapters/lessons; upload video (private) + docs; submit for review', 'Partial'],
        ['TI-06', 'Quiz & Assessment System', 'v1.0.0', 'Per-chapter quiz config, question bank, attempt + grading via JSONB snapshot', 'Ready'],
        ['TI-07', 'Admin Portal', 'v1.0.0', 'Course approval, user management, revenue reports, payout confirmation', 'Partial'],
        ['TI-08', 'Parent Portal', 'v1.0.0', 'Link to student, monitor progress, contact teacher', 'Ready'],
    ]),

    ('h1', 'SECTION 4 – FEATURES TO BE TESTED'),
    ('p', 'The table below lists at least 18 features within the testing scope, evenly distributed across the 8 core modules.'),
    ('table', [
        ['Feature ID', 'Feature Name', 'Module', 'Priority', 'Brief Description'],
        ['F-01', 'Email/password login', 'Auth', 'High', 'Valid/invalid credentials; verify the returned JWT'],
        ['F-02', 'Account registration with email OTP', 'Auth', 'High', 'Flow: send OTP → verify → create STUDENT profile'],
        ['F-03', 'Google OAuth login (ES256 JWT)', 'Auth', 'High', 'Google consent → Supabase callback → sync profile'],
        ['F-04', 'Role-based access control (RBAC)', 'Auth', 'High', 'Student/Teacher/Parent/Admin correct route guards'],
        ['F-05', 'Refresh token & session timeout', 'Auth', 'Medium', 'Expired token auto-refresh; logout revokes token'],
        ['F-06', 'Browse course list + filters', 'Course', 'High', 'Filter by category, price, grade; pagination'],
        ['F-07', 'View detail and free preview lesson', 'Course', 'High', 'isFree=true allows viewing; private video requires enrollment'],
        ['F-08', 'Add to cart – login required', 'Payment', 'High', 'Not logged in → redirect to /login with state.from'],
        ['F-09', 'VNPay/MoMo payment (sandbox)', 'Payment', 'High', 'Create order → redirect payment → webhook → enroll'],
        ['F-10', 'View order history', 'Dashboard', 'Medium', 'Order list, status, purchase date'],
        ['F-11', 'Update profile and upload avatar', 'Dashboard', 'Medium', 'Edit name, bio; upload image to Supabase Storage'],
        ['F-12', 'Teacher create/edit/delete course', 'Teacher', 'High', 'CRUD with chapter/lesson; status DRAFT→PENDING'],
        ['F-13', 'Upload lecture video (private bucket)', 'Teacher', 'High', 'Upload to Supabase private; store storagePath'],
        ['F-14', 'Submit course for admin review', 'Teacher', 'High', 'Change status DRAFT→PENDING; admin receives notification'],
        ['F-15', 'Configure chapter quiz + question bank', 'Quiz', 'High', 'CRUD questions; set count, time, passing score'],
        ['F-16', 'Student takes quiz + views result', 'Quiz', 'High', 'Random pick from bank; JSONB snapshot; auto-grading'],
        ['F-17', 'Admin approve / reject / request revision', 'Admin', 'High', 'Review content; approve→PUBLISHED; reject with reason'],
        ['F-18', 'Parent links and monitors progress', 'Parent', 'Medium', 'Send email invite; after ACTIVE, view progress'],
    ]),

    ('h1', 'SECTION 5 – FEATURES NOT TO BE TESTED'),
    ('p', 'The following features are excluded from the v1.0.0 testing scope, with specific justifications:'),
    ('table', [
        ['Excluded Feature', 'Justification'],
        ['Penetration Testing / Security Audit', 'Requires a dedicated security specialist; to be performed before production go-live'],
        ['Load Testing / Performance Testing', 'Requires a staging environment configured like production; not feasible in local dev'],
        ['Cross-browser (Safari, Edge)', 'v1.0 scope supports only Chrome + Firefox; Safari/Edge tested in a later sprint'],
        ['Mobile responsiveness (< 768px)', 'Mobile version will be tested separately in Sprint 3 once responsive design is complete'],
        ['Real MoMo integration (production API)', 'Only VNPay sandbox is available; MoMo production API requires a contract and separate testing'],
        ['Certificate feature (UC42-43)', 'Not developed in v1.0.0; planned for Sprint 4'],
        ['Messaging / AI chat with teacher (UC20-21)', 'Chat module not developed; requires WebSocket or third-party integration'],
        ['Teacher bank account feature (UC45-46)', 'Backend not complete; manual admin payout will be tested when the feature is complete'],
    ]),

    ('h1', 'SECTION 6 – APPROACH (TEST STRATEGY)'),
    ('p', 'The testing strategy for Bee Academy v1.0.0 applies multi-level testing, combining black-box and white-box techniques to ensure comprehensive quality assurance before release.'),
    ('h2', '6.1  Test Levels'),
    ('table', [
        ['Level', 'Owner', 'Technique', 'Tools'],
        ['Unit Testing', 'Dev Team', 'White-box, Statement Coverage (≥80%)', 'JUnit 5 / Jest'],
        ['Integration Testing', 'Dev + QA', 'API Testing, Black-box, Contract Test', 'Postman / REST Assured'],
        ['System Testing', 'QA Team', 'Black-box: EP, BVA; Exploratory', 'Playwright / JIRA'],
        ['Acceptance Testing', 'PO + Stakeholders', 'Scenario-based, User Acceptance', 'Manual / TestRail'],
    ]),
    ('h2', '6.2  Test Types'),
    ('table', [
        ['Type', 'Objective', 'Applied Scope'],
        ['Functional Testing', 'Verify business logic against Use Case v6.5', 'All 18 features in Section 4'],
        ['Security Testing', 'Test JWT authentication, authorization bypass, SQL injection', 'Auth, Teacher Portal, Admin Portal modules'],
        ['Usability Testing', 'Test intuitive UX/UI, desktop responsiveness', 'Landing Page, Course Detail, Student Dashboard'],
        ['Regression Testing', 'Ensure bug fixes do not break existing features', 'After each Critical/Major defect fix'],
    ]),
    ('h2', '6.3  Test Design Techniques'),
    ('bullet', 'Equivalence Partitioning (EP) – for input validation (email, password, price)'),
    ('bullet', 'Boundary Value Analysis (BVA) – for price range, grade filter, upload file size'),
    ('bullet', 'Decision Table – for the payment flow (VNPay callback: success/fail/cancel/timeout)'),
    ('bullet', 'State Transition – for the course lifecycle (DRAFT→PENDING→PUBLISHED/REJECTED)'),
    ('bullet', 'Error Guessing – expired JWT, invalid OTP, expired Supabase signed URL'),
    ('bullet', 'Exploratory Testing – find defects not covered by test cases via session-based exploration'),
    ('h2', '6.4  Entry & Exit Criteria'),
    ('table', [
        ['Level', 'Entry Criteria', 'Exit Criteria'],
        ['System Testing', 'Build deployed successfully to test env; unit test pass ≥80%; test cases & data prepared', 'Test case pass rate ≥95%; no open Critical defect; Test Summary Report produced'],
        ['Acceptance Testing', 'System Testing passed its exit criteria; regression testing complete; UAT environment ready', 'All UAT scenarios PASS; PO signs off; release checklist complete'],
    ]),

    ('h1', 'SECTION 7 – ITEM PASS/FAIL CRITERIA'),
    ('h2', '7.1  Pass Criteria'),
    ('bullet', 'Test case PASS rate ≥ 95% of all executed test cases'),
    ('bullet', 'No Critical or Blocker defect in Open status'),
    ('bullet', 'No Major defect remaining Open after D-3 before release'),
    ('bullet', 'All Major defects have a documented workaround in JIRA'),
    ('bullet', 'Authentication module: no Security defect of any kind (JWT bypass, SQL injection)'),
    ('bullet', 'Payment module: the entire VNPay sandbox payment flow completes successfully'),
    ('bullet', 'Test Summary Report approved by the Test Lead'),
    ('h2', '7.2  Fail Criteria (stop release)'),
    ('bullet', 'Test case PASS rate < 90%'),
    ('bullet', 'At least one unfixed Critical/Blocker defect exists'),
    ('bullet', 'The Authentication module has any open security defect'),
    ('bullet', 'The VNPay payment flow cannot complete (payment flow broken)'),
    ('bullet', 'User data is lost or incorrectly overwritten after a transaction'),
    ('bullet', 'The system crashes or fails to start in the test environment'),
    ('h2', '7.3  Defect Severity Definition'),
    ('table', [
        ['Severity', 'Definition', 'Example in Bee Academy'],
        ['Critical / Blocker', 'System unusable; no workaround', 'Cannot log in; payment crash; data loss'],
        ['Major', 'Important function broken; temporary workaround exists', 'Quiz grades incorrectly; video fails to play'],
        ['Minor', 'Function works but has a small defect; low impact', 'Typo in text; minor UI misalignment'],
        ['Cosmetic', 'UI defect not affecting functionality', 'Icon shows wrong color; uneven spacing'],
    ]),

    ('h1', 'SECTION 8 – SUSPENSION CRITERIA AND RESUMPTION REQUIREMENTS'),
    ('p', 'The table below describes situations that require suspending testing and the conditions required to resume.'),
    ('table', [
        ['Situation', 'Suspension Condition', 'Resumption Condition'],
        ['Build/Deploy failure', 'Test env cannot deploy the new version; error rate > 50%', 'Build passes; smoke test of 5 core features succeeds; Dev confirms'],
        ['Critical Blocker appears', '≥1 Critical defect found affecting the module under test', 'Dev fixes and merges; QA verifies the fix; regression test of that module passes'],
        ['Test environment down', 'Backend or Supabase unreachable for > 30 continuous minutes', 'Environment restored; health check passes; Test Lead confirms'],
        ['Test data corrupted', 'Test data is wrong, missing, or overwritten due to a script/migration error', 'Restore test data from backup; QA confirms data integrity'],
    ]),
    ('p', 'When suspended, the Test Lead must record the cause in the Daily Report, update the status in JIRA, and notify the Project Manager within 1 hour.'),

    ('h1', 'SECTION 9 – TEST DELIVERABLES'),
    ('p', 'List of deliverables produced throughout the testing process:'),
    ('table', [
        ['Phase', 'Deliverable', 'Responsible', 'Deadline'],
        ['Pre-testing', 'Master Test Plan (this document)', 'Test Lead', '2026-06-12'],
        ['Pre-testing', 'Test Cases (Excel / TestRail)', 'QA Team', '2026-06-17'],
        ['Pre-testing', 'Test Data Script (SQL + seed data)', 'QA Team', '2026-06-16'],
        ['Pre-testing', 'Test Environment Setup Guide', 'Dev + QA', '2026-06-15'],
        ['During testing', 'Daily Test Progress Report', 'Test Lead', 'Daily'],
        ['During testing', 'Defect Report (JIRA tickets)', 'QA Team', 'When found'],
        ['During testing', 'Test Execution Log', 'QA Team', 'Daily'],
        ['During testing', 'Change Request Log (if requirements change)', 'Test Lead', 'As needed'],
        ['Post-testing', 'Test Summary Report', 'Test Lead', '2026-07-14'],
        ['Post-testing', 'Defect Metrics & Quality Report', 'Test Lead', '2026-07-14'],
    ]),

    ('h1', 'SECTION 10 – TESTING TASKS'),
    ('p', 'The table below decomposes the testing activities into specific tasks, with execution order and clear dependencies.'),
    ('table', [
        ['Task ID', 'Task Name', 'Predecessor', 'Effort', 'Role', 'Start', 'End'],
        ['T01', 'Analyze requirements & Use Case v6.5', '–', '8h', 'Test Lead', '06/12', '06/13'],
        ['T02', 'Design test cases (18 features)', 'T01', '24h', 'QA Team', '06/14', '06/17'],
        ['T03', 'Prepare test data & seed scripts', 'T01', '8h', 'QA Team', '06/14', '06/16'],
        ['T04', 'Set up and configure test environment', 'T01', '4h', 'Dev + QA', '06/12', '06/14'],
        ['T05', 'System Test – Auth + Course + Dashboard', 'T02,T03,T04', '20h', 'QA Team', '06/18', '06/23'],
        ['T06', 'System Test – Payment + Teacher Portal', 'T05', '20h', 'QA Team', '06/24', '06/30'],
        ['T07', 'System Test – Quiz + Admin + Parent', 'T06', '20h', 'QA Team', '07/01', '07/07'],
        ['T08', 'Regression Testing (after fixing defects)', 'T07', '16h', 'QA Team', '07/08', '07/10'],
        ['T09', 'User Acceptance Testing (UAT)', 'T08', '16h', 'PO + Stakeholders', '07/11', '07/13'],
        ['T10', 'Write Test Summary Report & Sign-off', 'T09', '4h', 'Test Lead', '07/14', '07/15'],
    ]),

    ('h1', 'SECTION 11 – ENVIRONMENTAL NEEDS'),
    ('h2', '11.1  Hardware Requirements'),
    ('table', [
        ['Component', 'Minimum Configuration'],
        ['Test Server (Backend)', 'CPU: 4 cores, RAM: 8GB, Storage: 100GB SSD; or Cloud: Supabase + Railway'],
        ['Tester Workstation', 'CPU: 2 cores, RAM: 8GB, OS: Windows 11 / Ubuntu 22.04'],
        ['CI/CD Pipeline', 'GitHub Actions (2 workers, 4GB RAM per runner)'],
    ]),
    ('h2', '11.2  Software Requirements'),
    ('table', [
        ['Software', 'Version', 'Purpose'],
        ['OS', 'Windows 11 / Ubuntu 22.04', 'Test execution environment'],
        ['Node.js', '20 LTS', 'Build and run React 19 + Vite frontend'],
        ['Java JDK', '17', 'Spring Boot 3.2 backend runtime'],
        ['Maven', '3.9+', 'Backend build tool'],
        ['PostgreSQL (Supabase)', '15', 'Database; uses a Supabase test project'],
        ['Chrome', 'Latest', 'Primary browser (System Test)'],
        ['Firefox', 'Latest', 'Secondary browser (Regression)'],
        ['Postman', 'v10+', 'API testing – Auth endpoints, JWT validation'],
        ['Playwright', '1.40+', 'UI automation – E2E test scenarios'],
        ['JIRA', 'Cloud', 'Defect tracking and test management'],
        ['Git / GitHub', 'Latest', 'Version control; CI/CD integration'],
    ]),
    ('h2', '11.3  Test Environment Configuration'),
    ('bullet', 'Frontend: http://localhost:3000 (Vite dev server, port fixed in vite.config.ts)'),
    ('bullet', 'Backend: http://localhost:8080 (Spring Boot, port configured in backend/.env)'),
    ('bullet', 'Database: Supabase PostgreSQL (separate test project, isolated from production)'),
    ('bullet', 'Storage: Supabase Storage – 2 buckets: course-videos (private) + course-docs (public)'),
    ('bullet', 'JWT: ES256 (ECDSA P-256) verified via the Supabase JWKS endpoint'),
    ('bullet', 'CORS: allow origin http://localhost:3000'),
    ('bullet', 'DEV_MODE=true: OTP logged to console instead of sending a real email'),
    ('h2', '11.4  Test Data Requirements'),
    ('bullet', '5 Student accounts (valid credentials + edge cases: wrong password, locked account)'),
    ('bullet', '3 Teacher accounts (1 with an approved course, 1 pending, 1 newly created)'),
    ('bullet', '2 Parent accounts (1 linked to a student, 1 not linked)'),
    ('bullet', '1 Admin account'),
    ('bullet', '10 sample courses (DRAFT: 3, PENDING: 3, PUBLISHED: 4, REJECTED: 1)'),
    ('bullet', 'VNPay sandbox test payment cards: success + fail + timeout scenarios'),
    ('bullet', '50 sample question-bank items (10 per module, distributed Easy/Medium/Hard)'),
    ('bullet', 'Sample video file (≤100MB, MP4) and PDF document (≤10MB) for upload testing'),

    ('h1', 'SECTION 12 – RESPONSIBILITIES'),
    ('p', 'RACI matrix assigning responsibility for each testing activity (R=Responsible, A=Accountable, C=Consulted, I=Informed):'),
    ('table', [
        ['Activity', 'Test Lead', 'QA Engineer', 'Developer', 'PM', 'PO'],
        ['Write Test Plan', 'R/A', 'C', 'C', 'I', 'I'],
        ['Design Test Cases', 'A', 'R', 'C', 'I', 'C'],
        ['Prepare Test Data', 'A', 'R', 'C', 'I', '–'],
        ['Set up environment', 'A', 'R', 'R/A', 'I', '–'],
        ['Execute Testing', 'A', 'R', 'I', 'I', 'I'],
        ['Bug Report & Tracking', 'A', 'R', 'I', 'I', 'I'],
        ['Fix Defects', 'I', 'C', 'R/A', 'I', '–'],
        ['Test Summary Report', 'R/A', 'C', 'I', 'I', 'I'],
    ]),

    ('h1', 'SECTION 13 – STAFFING AND TRAINING NEEDS'),
    ('h2', '13.1  Team Composition'),
    ('table', [
        ['Role', 'Headcount', 'Allocation', 'Required Skills', 'Training Needed'],
        ['Test Lead', '1', '100%', 'IEEE 829-2008; ISTQB Foundation; API testing; JIRA', 'None'],
        ['QA Engineer', '2', '100%', 'Manual testing; basic SQL; Chrome DevTools; Postman', 'Playwright workshop (2 days); Supabase API basics (1 day)'],
        ['Dev (support)', '2', '20%', 'JUnit 5 unit test; code review; environment setup', 'None'],
    ]),
    ('h2', '13.2  Training Plan'),
    ('table', [
        ['Training Course', 'Audience', 'Duration', 'Timing'],
        ['Playwright E2E Automation Workshop', 'QA Engineers', '2 days', 'Week 1 (2026-06-12 to 06-13)'],
        ['Supabase Storage & JWT ES256 Overview', 'QA Team', '3 hours', '2026-06-14'],
    ]),

    ('h1', 'SECTION 14 – SCHEDULE'),
    ('h2', '14.1  Milestones'),
    ('table', [
        ['Milestone', 'Date', 'Description'],
        ['MS-01', '2026-06-12', 'Kickoff, Test Plan approved'],
        ['MS-02', '2026-06-17', 'Test Cases & Test Data ready; test environment ready'],
        ['MS-03', '2026-07-07', 'System Testing complete (all 8 modules)'],
        ['MS-04', '2026-07-10', 'Regression Testing complete'],
        ['MS-05', '2026-07-13', 'UAT complete; PO sign-off'],
        ['MS-06', '2026-07-15', 'Test Summary Report delivered'],
    ]),
    ('h2', '14.2  Weekly Schedule (Gantt Chart)'),
    ('table', [
        ['Phase / Activity', 'W1 (6/12-16)', 'W2 (6/17-23)', 'W3 (6/24-30)', 'W4 (7/1-10)', 'W5 (7/11-15)'],
        ['Planning & Test Design', 'XX', 'XX', '', '', ''],
        ['Environment Setup', 'XX', '', '', '', ''],
        ['System Test – Auth & Course', '', 'XX', '', '', ''],
        ['System Test – Pay & Teacher', '', '', 'XX', '', ''],
        ['System Test – Quiz,Admin,Parent', '', '', '', 'XX', ''],
        ['Regression Testing', '', '', '', 'XX', ''],
        ['UAT & Test Closure', '', '', '', '', 'XX'],
    ]),

    ('h1', 'SECTION 15 – RISKS AND CONTINGENCIES'),
    ('p', 'The table below identifies testing risks, rates their level (L×I), and defines mitigation / contingency plans. (L=Likelihood: H/M/L  |  I=Impact: H/M/L)'),
    ('table', [
        ['Risk ID', 'Risk', 'L', 'I', 'Level', 'Mitigation', 'Contingency'],
        ['R01', 'Unstable test environment (Supabase local dev, CORS errors)', 'M', 'H', 'High', 'Prepare a separate staging environment; docker-compose backup', 'Use a hosted Supabase test project instead of localhost'],
        ['R02', 'JWT ES256 key rotation logs out all users unexpectedly', 'L', 'H', 'Medium', 'Do not rotate the JWKS key during testing; note it in env setup', 'Restart backend and re-login; log the incident'],
        ['R03', 'VNPay sandbox does not reflect production behavior', 'H', 'M', 'High', 'Use both success + fail test cards; document any discrepancies', 'Test the payment flow with a mock server if the sandbox is unstable'],
        ['R04', 'Requirements change during testing (Use Case v6.5 not frozen)', 'M', 'H', 'High', 'Freeze requirements from W1; all changes go through a Change Request', 'Re-scope test cases; adjust the schedule if needed'],
        ['R05', 'Staff shortage due to sick leave or personal reasons', 'L', 'H', 'Medium', 'Cross-training among QA Engineers; clear test case documentation', 'Test Lead temporarily supports execution; prioritize High features'],
        ['R06', 'Build delays from the Dev Team affect the schedule', 'M', 'M', 'Medium', 'Daily stand-up; automated CI/CD; early integration merge', 'Test Ready modules first; delay testing In-Dev modules'],
        ['R07', 'Test data corrupted after testing the payment flow', 'M', 'H', 'High', 'Use transaction rollback or a reset script after each payment test session', 'Restore the test DB from the seed script; log the incident'],
        ['R08', 'Supabase Storage signed URL expires (TTL 1h) mid-test session', 'L', 'M', 'Low', 'Regenerate the signed URL before each video access test; note TTL in the test case', 'Increase TTL in the test env or test video within a < 1h session'],
        ['R09', 'Private video bucket misconfigured → public access', 'L', 'H', 'Medium', 'Verify the RLS policy on the Supabase dashboard before testing', 'Fix the policy immediately; security incident report; re-test all video access'],
    ]),

    ('h1', 'SECTION 16 – APPROVALS'),
    ('p', 'This document becomes effective once signed off by all stakeholders listed below:'),
    ('table', [
        ['Role', 'Full Name', 'Signature', 'Approval Date'],
        ['Test Lead', 'Vo Van Thanh Dat', '', ''],
        ['Project Manager', 'Nguyen Van A', '', ''],
        ['Product Owner', 'Tran Thi B', '', ''],
        ['Dev Lead', 'Le Van C', '', ''],
    ]),
    ('p', 'Note: This document is in Draft status. Signatures will be collected after the first review round is complete (Review Date: 2026-06-19).'),
]


# =====================================================================
# ============== PHẦN B + C - TIẾNG VIỆT ==============================
# =====================================================================
BC = [
    ('h1', 'PHẦN B – AI INTERACTION LOG'),
    ('p', 'Phần này ghi lại 3 lần tương tác với AI (Claude Sonnet 4.6) khi soạn thảo Test Plan cho dự án Bee Academy, kèm đánh giá chất lượng kết quả và những điểm đã được chỉnh sửa so với output gốc của AI.'),

    ('h2', 'Tương tác 1 – Mục 1 & 2: Test Plan Identifier + Introduction'),
    ('table', [
        ['Thông tin', 'Chi tiết'],
        ['Prompt đã dùng', 'Create Test Plan Identifier and Introduction sections for Bee Academy - an online learning platform for Vietnamese THCS students (grade 6-9). Tech stack: React 19 + Spring Boot 3.2 + PostgreSQL via Supabase (JWT ES256/ECDSA P-256). Users: students, teachers, parents, admins. Payment: VNPay/MoMo. Follow IEEE 829-2008 format.'],
    ]),
    ('p', 'Bảng đánh giá kết quả AI:'),
    ('table', [
        ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'],
        ['Đúng chuẩn IEEE 829-2008', '4/5', 'AI tạo đúng các trường Document ID, Author, Status nhưng thiếu Review Date và document history'],
        ['Phù hợp với context Bee Academy', '3/5', 'AI không biết về JWT ES256/ECDSA P-256 của Supabase, đề xuất HS256 generic'],
        ['Đủ chi tiết, có thể dùng ngay', '3/5', 'Introduction quá ngắn, thiếu phần tài liệu tham chiếu đặc thù dự án'],
        ['Cần chỉnh sửa nhiều hay ít', '3/5', 'Phải chỉnh sửa vừa phải: thêm JWT specifics, Supabase endpoints, VNPay context'],
        ['Tổng', '13/20', ''],
    ]),
    ('p', 'Những gì đã chỉnh sửa so với output AI:'),
    ('bullet', 'Thêm thông tin JWT ES256 (ECDSA P-256) thay vì HS256 mà AI đề xuất'),
    ('bullet', 'Bổ sung Supabase JWKS endpoint cụ thể vào References'),
    ('bullet', 'Thêm UseCase v6.5 vào danh sách tài liệu tham chiếu'),
    ('bullet', 'Điều chỉnh scope: loại bỏ Next.js (AI tưởng dự án dùng Next.js) → Vite 6'),

    ('h2', 'Tương tác 2 – Mục 6: Approach (Test Strategy)'),
    ('table', [
        ['Thông tin', 'Chi tiết'],
        ['Prompt đã dùng', 'Write the Test Approach section for Bee Academy following IEEE 829-2008. The system includes: Auth (JWT ES256 ECDSA, Google OAuth, OTP), Course management with video private Supabase bucket, VNPay payment integration, Teacher portal with video upload, Quiz system with JSONB snapshot grading. Include test levels (unit/integration/system/acceptance), test types, design techniques (EP, BVA, decision table, state transition), entry/exit criteria for each level, and tools.'],
    ]),
    ('p', 'Bảng đánh giá kết quả AI:'),
    ('table', [
        ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'],
        ['Đúng chuẩn IEEE 829-2008', '5/5', 'Sau khi được cung cấp context đầy đủ, AI tạo đúng cấu trúc 4 test levels với entry/exit criteria'],
        ['Phù hợp với context Bee Academy', '4/5', 'AI áp dụng EP và BVA đúng cho các trường hợp của Bee Academy (giá tiền VND, cấp lớp 6-9)'],
        ['Đủ chi tiết, có thể dùng ngay', '4/5', 'Khá đầy đủ; entry/exit criteria cụ thể, test types phân loại rõ ràng'],
        ['Cần chỉnh sửa nhiều hay ít', '4/5', 'Chỉ cần chỉnh sửa nhỏ: thay Selenium → Playwright, thêm State Transition cho course lifecycle'],
        ['Tổng', '17/20', ''],
    ]),
    ('p', 'Những gì đã chỉnh sửa so với output AI:'),
    ('bullet', 'Thay Selenium 4.x bằng Playwright 1.40+ (phù hợp hơn với React SPA)'),
    ('bullet', 'Thêm State Transition Testing cho vòng đời khóa học: DRAFT→PENDING→PUBLISHED/REJECTED'),
    ('bullet', 'Bổ sung test type Security Testing riêng cho module Auth (AI gộp chung vào Functional)'),
    ('bullet', 'Điều chỉnh pass rate threshold: từ 90% (AI đề xuất) lên 95% theo requirement của project'),

    ('h2', 'Tương tác 3 – Mục 15: Risks and Contingencies'),
    ('table', [
        ['Thông tin', 'Chi tiết'],
        ['Prompt đã dùng', 'Identify and analyze testing risks for Bee Academy online learning platform (React 19 + Spring Boot 3.2 + Supabase PostgreSQL). Specific technical concerns: VNPay sandbox vs production behavioral differences, Supabase private bucket signed URL TTL 1 hour, JWT ES256 key dependencies, Supabase free tier rate limits during testing. Follow IEEE 829-2008 format with Risk ID, description, category, likelihood (H/M/L), impact (H/M/L), risk level, mitigation strategy, and contingency plan.'],
    ]),
    ('p', 'Bảng đánh giá kết quả AI:'),
    ('table', [
        ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'],
        ['Đúng chuẩn IEEE 829-2008', '4/5', 'Format Risk ID, L, I, Level, Mitigation, Contingency đúng chuẩn; thiếu cột Risk Category'],
        ['Phù hợp với context Bee Academy', '4/5', 'AI nhận diện được rủi ro kỹ thuật đặc thù (signed URL, VNPay sandbox) sau khi được cung cấp context'],
        ['Đủ chi tiết, có thể dùng ngay', '3/5', 'Thiếu rủi ro về business (legal compliance payment, PO sign-off delay) và rủi ro test data'],
        ['Cần chỉnh sửa nhiều hay ít', '3/5', 'Chỉnh sửa vừa phải: thêm 3 rủi ro mới, xóa 2 rủi ro generic không phù hợp'],
        ['Tổng', '14/20', ''],
    ]),
    ('p', 'Những gì đã chỉnh sửa so với output AI:'),
    ('bullet', 'Thêm R07 (Test data corrupt sau payment flow) – AI bỏ sót hoàn toàn rủi ro này'),
    ('bullet', 'Thêm R08 (Signed URL TTL expire mid-session) – AI đề cập nhưng không có contingency plan'),
    ('bullet', 'Thêm R09 (Supabase RLS misconfiguration → video public) – rủi ro security quan trọng'),
    ('bullet', 'Xóa "Database backup failure" – không phù hợp với context dùng Supabase managed service'),
    ('bullet', 'Điều chỉnh mức L của R03 (VNPay sandbox) từ M→H vì đã gặp issue thực tế khi dev'),

    ('p', 'Lưu ý quan trọng: AI có thể tạo ra nội dung nghe có vẻ hợp lý nhưng thiếu chi tiết kỹ thuật cụ thể hoặc không phù hợp với project. Nhiệm vụ của tester là review và refine output, không copy-paste mù quáng.'),

    ('pagebreak', ''),
    ('h1', 'PHẦN C – CÂU HỎI THẢO LUẬN'),
    ('h2', 'CÂU HỎI 1'),
    ('p', 'Khi sử dụng AI để tạo Test Plan, AI thường mạnh ở những mục nào và yếu ở những mục nào trong chuẩn IEEE 829-2008? Giải thích tại sao có sự khác biệt đó. Một tester chuyên nghiệp cần bổ sung gì mà AI không thể tự làm được?'),
    ('p', 'Trả lời:'),
    ('p', 'Qua trải nghiệm thực tế khi làm bài, AI (Claude Sonnet 4.6) thể hiện rõ điểm mạnh và điểm yếu khác nhau tùy từng mục của IEEE 829-2008.'),
    ('p', 'Những mục AI làm tốt: AI mạnh nhất ở các mục có cấu trúc cố định, theo template rõ ràng: Mục 1 (Test Plan Identifier), Mục 9 (Test Deliverables), Mục 12 (Responsibilities – RACI matrix), Mục 14 (Schedule – Gantt chart), Mục 16 (Approvals). Lý do là những mục này về cơ bản chỉ cần điền thông tin vào template có sẵn; AI được huấn luyện trên hàng nghìn tài liệu IEEE, nên nó nắm vững cấu trúc bắt buộc. Ví dụ, khi yêu cầu tạo RACI matrix, AI trả về đúng format R/A/C/I với 5 roles chuẩn ngay lần đầu.'),
    ('p', 'AI cũng làm tốt Mục 6 (Approach) khi được cung cấp context đầy đủ về tech stack và business domain. Sau khi tôi mô tả chi tiết về JWT ES256, Supabase Storage, VNPay payment flow, AI tạo được test levels và entry/exit criteria khá chính xác, chỉ cần điều chỉnh nhỏ.'),
    ('p', 'Những mục AI còn yếu: AI yếu nhất ở Mục 4 (Features to be Tested) và Mục 7 (Pass/Fail Criteria). Đối với Mục 4, AI có xu hướng liệt kê các feature theo tên module một cách generic (Login, Register, View Course...) mà không hiểu các ràng buộc business thực tế. Ví dụ, AI không biết rằng nút "Thêm vào giỏ hàng" trong Bee Academy chỉ khả dụng khi đã login, hoặc việc giáo viên submit khóa học phải qua flow DRAFT→PENDING→Admin duyệt. Những business rule này chỉ có trong CLAUDE.md và UseCase v6.5 – tài liệu nội bộ mà AI không có access.'),
    ('p', 'Mục 15 (Risks) là điểm yếu đáng chú ý nhất. AI thường đề xuất rủi ro generic (build failure, resource shortage, environment instability) – những rủi ro mà mọi dự án web đều gặp. AI bỏ sót hoàn toàn các rủi ro đặc thù của Bee Academy: JWT ES256 key rotation có thể làm tất cả user bị logout đồng loạt, Supabase signed URL TTL 1 giờ có thể expire mid-test-session, hay việc private video bucket bị misconfigure RLS sẽ gây lộ nội dung có phí. Đây là những rủi ro "chỉ người trong dự án mới biết".'),
    ('p', 'Điều tester chuyên nghiệp cần bổ sung:'),
    ('bullet', 'Hiểu domain & business context: AI không thể đọc giữa các dòng của UseCase, biết stakeholder nào quan tâm điều gì, hay đâu là "killer feature" cần test kỹ nhất.'),
    ('bullet', 'Judgment call về priority và risk: Quyết định pass rate threshold 95% hay 90%, hay việc một lỗi Minor trên payment page thực chất là Critical – AI không có context để ra quyết định này.'),
    ('bullet', 'Thương lượng với stakeholders: Test Plan cần phản ánh thỏa thuận thực sự giữa QA, Dev và Business. AI tạo ra tài liệu một chiều, không qua negotiation.'),
    ('bullet', 'Verify feasibility: AI đề xuất lịch test 1 tháng với 3 tester nhưng không biết team còn đang làm việc khác, hay môi trường Supabase free tier có rate limit.'),
    ('bullet', 'Cập nhật liên tục: Test Plan phải thay đổi theo project evolution. AI chỉ tạo snapshot tại một thời điểm, không tự cập nhật khi requirement thay đổi.'),

    ('h2', 'CÂU HỎI 2'),
    ('p', 'Mục 15 (Risks & Contingencies) yêu cầu tester phải suy nghĩ về những gì có thể sai. So sánh danh sách rủi ro do AI đề xuất với những rủi ro bạn tự nghĩ ra dựa trên kiến thức môn học. AI có bỏ sót loại rủi ro nào không? Từ đó, rút ra bài học về giới hạn của AI trong Test Management.'),
    ('p', 'Trả lời:'),
    ('p', 'So sánh rủi ro AI đề xuất vs. rủi ro tự phân tích: Khi tôi yêu cầu AI liệt kê rủi ro cho Bee Academy, danh sách ban đầu AI trả về bao gồm: (1) Unstable test environment, (2) Resource shortage, (3) Requirements change, (4) Build delays, (5) Missing test data. Đây đều là những rủi ro kinh điển trong mọi dự án phần mềm – đúng nhưng chưa đủ.'),
    ('p', 'Khi tự phân tích dựa trên kiến thức môn SWT301 và kiến trúc thực tế của Bee Academy từ CLAUDE.md, tôi phát hiện thêm các loại rủi ro mà AI đã bỏ sót hoàn toàn:'),
    ('p', 'Loại rủi ro AI bỏ sót – Technical Risk đặc thù:'),
    ('bullet', 'R02 – JWT ES256 Key Dependency: Bee Academy dùng ECDSA P-256, khác với HS256 thông thường. Nếu JWKS endpoint của Supabase fetch thất bại khi backend khởi động, ES256 verifier = null và toàn bộ JWT sẽ fail authentication. Đây là rủi ro duy nhất của Supabase JWT implementation, không có trong bất kỳ checklist generic nào.'),
    ('bullet', 'R08 – Signed URL TTL Expiry: Video trong private Supabase bucket được truy cập qua signed URL TTL 1 giờ. Nếu test session kéo dài > 1 giờ, URL hết hạn và tester nhầm tưởng đây là bug video playback. AI không biết điều này vì không có access vào code của Lesson.videoStoragePath.'),
    ('bullet', 'R09 – Supabase RLS Misconfiguration: Row Level Security của Supabase nếu không được set đúng cho bucket course-videos (private), video có thể bị public access mà không cần auth. Đây là security risk nghiêm trọng mà chỉ tester hiểu storage architecture mới nhận ra.'),
    ('p', 'Loại rủi ro AI bỏ sót – Business Risk:'),
    ('bullet', 'Risk về VNPay behavioral difference: Qua thực tế dev, VNPay webhook không reliable trong local environment (ngay cả với ngrok tunneling). AI không biết điều này vì chỉ biết "payment integration risk" theo nghĩa chung chung.'),
    ('bullet', 'Risk về DEV_MODE=true trong production: CLAUDE.md ghi rõ DEV_MODE=true trong backend/.env – nếu quên đổi khi deploy, OTP sẽ log ra console thay vì gửi email thật, đây là security vulnerability nghiêm trọng. AI không phân tích được config file cụ thể này.'),
    ('p', 'Bài học rút ra về giới hạn của AI trong Test Management: AI hoạt động dựa trên pattern matching từ training data. Nó giỏi tạo ra những gì "thường gặp" trong tài liệu kiểm thử, nhưng không thể suy luận từ code base, architecture decision, hay business constraint đặc thù của từng project.'),
    ('p', 'Giới hạn cốt lõi của AI trong Test Management là: AI không có "project context" – nó không đọc được CLAUDE.md, UseCase v6.5, hay biết rằng DEV_MODE=true trong .env là một time bomb chờ phát nổ khi go-live. Một Test Manager giỏi phải kết hợp kiến thức framework (IEEE 829-2008, ISTQB) với deep understanding về hệ thống cụ thể đang test. AI có thể hỗ trợ phần framework, nhưng phần domain knowledge và critical thinking vẫn phải do con người đảm nhận.'),
    ('p', 'Kết luận thực tiễn: Workflow hiệu quả nhất là dùng AI để tạo "80% foundation" (structure, template, generic risks, standard criteria), sau đó tester chuyên nghiệp review, bổ sung project-specific knowledge, và validate feasibility. Không nên copy-paste output AI trực tiếp vào deliverable mà không qua critical review – đây là điều Lab 4 muốn chúng ta học được.'),

    ('h1', 'TÀI LIỆU THAM KHẢO'),
    ('p', '[1] IEEE 829-2008 – IEEE Standard for Software and System Test Documentation'),
    ('p', '[2] ISTQB Foundation Level Syllabus v4.0 – Chapter 5: Test Management'),
    ('p', '[3] Foundations of Software Testing (Craig & Jaskiel) – Chapter 9'),
    ('p', '[4] ISTQB Glossary – https://glossary.istqb.org'),
    ('p', '[5] Bee Academy CLAUDE.md – Bee Academy Architecture & Project Documentation'),
    ('p', '[6] Bee Academy UseCase v6.5 – BEE ACADEMY.md (48 use cases, 9 modules)'),
    ('p', '[7] Supabase Documentation – Storage Buckets, RLS, JWT Authentication'),
    ('p', '[8] Spring Boot Security Documentation – JWT ES256 / ECDSA P-256 Configuration'),
    ('p', '[9] VNPay Developer Documentation – Sandbox Testing Guide'),
]


def build(filename, cover_lines, meta_rows, blocks):
    doc = Document()
    style_doc(doc)
    add_cover(doc, cover_lines, meta_rows)
    doc.add_page_break()
    render(doc, blocks)
    out = os.path.join(BASE, filename)
    doc.save(out)
    print('Saved:', out)


# ---- Build VI Test Plan ----
build(
    'Lab4_TestPlan_TiengViet_VoVanThanhDat_DE190211.docx',
    [
        ('BÁO CÁO THỰC HÀNH', 14, True, None, 4),
        ('LAB 4 – TEST PLAN THEO IEEE 829-2008 (PHẦN A)', 18, True, PRIMARY, 4),
        ('Phiên bản Tiếng Việt', 13, False, DARK, 10),
        ('DỰ ÁN: BEE ACADEMY', 16, True, DARK, 2),
        ('Nền tảng học trực tuyến dành cho học sinh THCS (Lớp 6–9)', 11, False, None, 10),
    ],
    [
        ['Môn học', 'SWT301 – Software Testing'],
        ['Chủ đề', 'Topic 5 – Test Management'],
        ['CLO', 'CLO5 – Xây dựng test plan toàn diện có tích hợp AI'],
        ['Chuẩn áp dụng', 'IEEE 829-2008'],
        ['Dự án thực hành', 'Bee Academy – Online Learning Platform v1.0.0'],
        ['Sinh viên thực hiện', 'Võ Văn Thành Đạt'],
        ['Mã số sinh viên', 'DE190211'],
        ['Ngày thực hiện', '12/06/2026'],
    ],
    VI_A,
)

# ---- Build EN Test Plan ----
build(
    'Lab4_TestPlan_English_VoVanThanhDat_DE190211.docx',
    [
        ('LAB REPORT', 14, True, None, 4),
        ('LAB 4 – TEST PLAN PER IEEE 829-2008 (PART A)', 18, True, PRIMARY, 4),
        ('English Version', 13, False, DARK, 10),
        ('PROJECT: BEE ACADEMY', 16, True, DARK, 2),
        ('Online learning platform for lower-secondary students (Grades 6–9)', 11, False, None, 10),
    ],
    [
        ['Course', 'SWT301 – Software Testing'],
        ['Topic', 'Topic 5 – Test Management'],
        ['CLO', 'CLO5 – Build a comprehensive AI-assisted test plan'],
        ['Standard', 'IEEE 829-2008'],
        ['Practice Project', 'Bee Academy – Online Learning Platform v1.0.0'],
        ['Student', 'Vo Van Thanh Dat'],
        ['Student ID', 'DE190211'],
        ['Date', '2026-06-12'],
    ],
    EN_A,
)

# ---- Build Part B + C ----
build(
    'Lab4_PhanB_C_TiengViet_VoVanThanhDat_DE190211.docx',
    [
        ('BÁO CÁO THỰC HÀNH', 14, True, None, 4),
        ('LAB 4 – PHẦN B & C', 18, True, PRIMARY, 4),
        ('AI Interaction Log + Câu hỏi thảo luận (Tiếng Việt)', 12, False, DARK, 10),
        ('DỰ ÁN: BEE ACADEMY', 16, True, DARK, 2),
        ('Nền tảng học trực tuyến dành cho học sinh THCS (Lớp 6–9)', 11, False, None, 10),
    ],
    [
        ['Môn học', 'SWT301 – Software Testing'],
        ['Chủ đề', 'Topic 5 – Test Management'],
        ['Chuẩn áp dụng', 'IEEE 829-2008'],
        ['Dự án thực hành', 'Bee Academy – Online Learning Platform v1.0.0'],
        ['Sinh viên thực hiện', 'Võ Văn Thành Đạt'],
        ['Mã số sinh viên', 'DE190211'],
        ['AI hỗ trợ', 'Claude Sonnet 4.6'],
        ['Ngày thực hiện', '12/06/2026'],
    ],
    BC,
)

print('DONE - 3 files generated.')
