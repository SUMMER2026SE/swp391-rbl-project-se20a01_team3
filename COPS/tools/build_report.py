from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "reference" / "Group5_COPS_Template.docx"
OUTPUT = ROOT / "output" / "BeeAcademy_COPS_Final_Report.docx"
LOGO_FPT = ROOT / "reference" / "template-package" / "word" / "media" / "image115.png"
LOGO_BEE = ROOT.parent / "frontend" / "public" / "logo-bee.png"


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = display
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_toc(paragraph):
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Mục lục sẽ được cập nhật khi mở tài liệu.")


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def style_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def add_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    style_run(r, bold=True)
    r = p.add_run(text)
    style_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId"); num_id.set(qn("w:val"), "1")
        num_pr.extend([ilvl, num_id]); p_pr.append(num_pr)
        style_run(p.add_run(item))


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, (header, width) in enumerate(zip(headers, widths_cm)):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F4CCCC")
        set_cell_margins(cell)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(header), size=10, bold=True)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, (value, width) in enumerate(zip(row_data, widths_cm)):
            cell = row.cells[idx]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            style_run(p.add_run(str(value)), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


figure_no = 0


def add_step_with_image(doc, step_no, text, image_rel, function_name):
    global figure_no
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(f"Bước {step_no}: "), bold=True)
    style_run(p.add_run(text))
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.keep_with_next = True
    run = image_p.add_run()
    run.add_picture(str(ROOT / "output" / "screenshots" / image_rel), width=Cm(15.7))
    figure_no += 1
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_together = True
    style_run(cap.add_run(f"Hình {figure_no}: Bước {step_no} trong chức năng {function_name}"), size=10, italic=True)


def set_doc_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[name]
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"; h1.font.size = Pt(16); h1.font.bold = True; h1.font.color.rgb = RGBColor(0xC0,0,0)
    h1.paragraph_format.space_before = Pt(12); h1.paragraph_format.space_after = Pt(4)
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"; h2.font.size = Pt(13); h2.font.bold = True; h2.font.color.rgb = RGBColor(0,0,0)
    h2.paragraph_format.space_before = Pt(8); h2.paragraph_format.space_after = Pt(3)
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"; h3.font.size = Pt(12); h3.font.bold = True; h3.font.color.rgb = RGBColor(0,0,0)
    h3.paragraph_format.space_before = Pt(6); h3.paragraph_format.space_after = Pt(3)
    try:
        caption = doc.styles["Caption"]
    except KeyError:
        caption = doc.styles.add_style("Caption", 1)
    caption.font.name = "Times New Roman"; caption.font.size = Pt(10); caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3); caption.paragraph_format.space_after = Pt(8)


def configure_sections(doc):
    for sec in doc.sections:
        sec.page_width = Cm(21.001); sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.498)
        sec.header_distance = Cm(1.249); sec.footer_distance = Cm(1.249)
        sec.different_first_page_header_footer = True
        footer_p = sec.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer_p, "PAGE", "1")
        for run in footer_p.runs: style_run(run, size=10)


def cover(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO_FPT), width=Cm(5.3))
    p.paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO_BEE), width=Cm(2.1))
    p.paragraph_format.space_after = Pt(18)
    for text, size, gap in [
        ("FPT UNIVERSITY", 16, 14),
        ("BEE ACADEMY - NỀN TẢNG HỌC TRỰC TUYẾN", 19, 18),
        ("TÀI LIỆU BÀN GIAO PHIÊN BẢN CUỐI / COPS", 18, 10),
        ("FINAL RELEASE DOCUMENT", 14, 70),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(gap)
        style_run(p.add_run(text), size=size, bold=True, color="C00000" if "COPS" in text else None)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    style_run(p.add_run("Đà Nẵng, tháng 7 năm 2026"), size=12, italic=True)
    doc.add_page_break()


def add_function(doc, number, req, screenshots_by_req):
    req_id, name, module, actor, route, implemented, testable, status, note = req
    doc.add_heading(f"{number}. {name}", level=3)
    add_label_paragraph(doc, "Mã yêu cầu: ", req_id)
    add_label_paragraph(doc, "Mục đích: ", f"Cho phép {actor} thực hiện chức năng {name.lower()} trong hệ thống Bee Academy.")
    add_label_paragraph(doc, "Điều kiện tiên quyết: ", "Truy cập đúng route; đăng nhập đúng vai trò nếu chức năng được bảo vệ.")
    add_label_paragraph(doc, "Trạng thái kiểm thử: ", status)
    add_label_paragraph(doc, "Route/điểm truy cập: ", route)

    if req_id == "REQ-CRS-001":
        add_step_with_image(doc, 1, "Tại trang chủ, chọn [Xem Khóa Học].", "guest/REQ-CRS-001/step-01-open-course-list-highlighted.png", name)
        add_step_with_image(doc, 2, "Tại trang danh sách, nhập từ khóa vào ô [Tìm kiếm khóa học, môn học, giảng viên...].", "guest/REQ-CRS-001/step-02-search-course-highlighted.png", name)
        add_label_paragraph(doc, "Kết quả quan sát: ", "Trang hiển thị 1 khóa học phù hợp, gồm khóa “Công Nghệ 6”.")
    elif req_id == "REQ-CRS-002":
        add_step_with_image(doc, 1, "Chọn thẻ khóa học [Công Nghệ 6].", "guest/REQ-CRS-002/step-01-select-course-highlighted.png", name)
        add_step_with_image(doc, 2, "Quan sát tiêu đề và thông tin tổng quan trên trang chi tiết.", "guest/REQ-CRS-002/step-02-view-course-details-highlighted.png", name)
        add_label_paragraph(doc, "Kết quả quan sát: ", "Trang chi tiết hiển thị đúng tên khóa học, giảng viên, 4 chương, 18 bài và giá khóa học.")
    elif req_id == "REQ-CRS-003":
        add_step_with_image(doc, 1, "Trên trang chi tiết, chọn tab [Nội dung học].", "guest/REQ-CRS-003/step-01-open-course-content-highlighted.png", name)
        add_label_paragraph(doc, "Kết quả quan sát: ", "Mục lục hiển thị nhưng mọi bài trong dữ liệu hiện tại đều có trạng thái [Cần mua khóa]. Không có bài học thử khả dụng để kiểm tra happy path.")
    elif req_id == "REQ-AUTH-001":
        add_step_with_image(doc, 1, "Mở trang [Đăng ký] và quan sát nhóm trường Họ và tên, Email, Mật khẩu, Vai trò.", "guest/REQ-AUTH-001/step-01-open-register-form-highlighted.png", name)
        add_label_paragraph(doc, "Kết quả quan sát: ", "Biểu mẫu đăng ký hiển thị. Không tạo tài khoản hoặc gửi OTP/email thật nên trạng thái được ghi PARTIAL.")
    elif req_id == "REQ-AUTH-002":
        add_step_with_image(doc, 1, "Nhập tài khoản kiểm thử vào biểu mẫu [Đăng Nhập].", "user/REQ-AUTH-002/step-01-enter-credentials-highlighted.png", name)
        add_step_with_image(doc, 2, "Chọn [Đăng Nhập].", "user/REQ-AUTH-002/step-02-submit-login-highlighted.png", name)
        add_label_paragraph(doc, "Kết quả quan sát: ", "Retest ngày 22/07/2026 xác nhận cả 4 tài khoản Student, Parent, Teacher và Admin đăng nhập thành công, điều hướng đúng dashboard theo vai trò.")
    elif req_id == "REQ-AUTH-004":
        add_step_with_image(doc, 1, "Mở trang [Quên mật khẩu] và quan sát biểu mẫu nhận OTP.", "user/REQ-AUTH-004/step-01-open-forgot-password-highlighted.png", name)
        add_label_paragraph(doc, "Kết quả quan sát: ", "Biểu mẫu hiển thị đúng. Không gửi OTP thật và không xác minh được luồng hoàn tất.")
    elif req_id in screenshots_by_req:
        # Retest evidence for the role-specific pages. Keep one image per
        # requirement so the manual remains readable while still documenting
        # the actual UI state observed during the retest.
        image_rel, image_row = screenshots_by_req[req_id][0]
        target_desc = image_row[6] or "Trang chức năng"
        add_step_with_image(doc, 1, f"Mở route [{route}] và quan sát khu vực {target_desc}.", image_rel, name)
        add_label_paragraph(doc, "Kết quả quan sát: ", note)
    else:
        p = doc.add_paragraph()
        status_text = {
            "PASSED": ("Đã xác minh trên UI trong phiên retest.", "008000"),
            "PARTIAL": ("Đã xác minh giao diện hoặc một phần luồng; chưa hoàn tất happy path.", "9C6500"),
            "FAILED": ("Đã tái hiện lỗi trong phiên retest.", "C00000"),
            "BLOCKED": ("Chưa thể hoàn tất do thiếu dữ liệu hoặc điều kiện tiên quyết.", "C00000"),
        }
        message, color = status_text.get(status, ("Chưa xác minh.", "C00000"))
        r = p.add_run(message)
        style_run(r, bold=True, color=color)
        add_label_paragraph(doc, "Bằng chứng: ", note)
        add_label_paragraph(doc, "Quy tắc tài liệu: ", "Không suy diễn Passed khi chưa quan sát đủ happy path; không thực hiện thanh toán, OTP/email, chuyển khoản hoặc khóa tài khoản thật.")


def main():
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    clear_body(doc)
    set_doc_styles(doc)
    configure_sections(doc)
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields"); settings.append(update)
    update.set(qn("w:val"), "true")

    requirements = json.loads((ROOT / "analysis" / "requirements.json").read_text(encoding="utf-8"))
    screenshot_rows = json.loads((ROOT / "analysis" / "screenshots.json").read_text(encoding="utf-8"))
    screenshots_by_req = {}
    for row in screenshot_rows:
        rel = row[8].replace("output/screenshots/", "")
        screenshots_by_req.setdefault(row[1], []).append((rel, row))
    cover(doc)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("MỤC LỤC"), size=16, bold=True, color="C00000")
    toc_p = doc.add_paragraph(); add_toc(toc_p)

    doc.add_heading("I. GÓI BÀN GIAO", level=1)
    deliverables = [
        (1,"README.md","Tổng quan kiến trúc và cách chạy dự án."),
        (2,"Bee_Academy_SRS_4.18_English_Complete_With_Section_7 (1).docx","Đặc tả yêu cầu phần mềm, 44 REQ trong phạm vi COPS."),
        (3,"backend/pom.xml","Cấu hình Java 21, Spring Boot 3.2.5 và Maven."),
        (4,"frontend/package.json","Cấu hình React 19, Vite, TypeScript và Playwright."),
        (5,"backend/db/migrations/","Các script migration PostgreSQL/Supabase."),
        (6,"backend/postman/BeeAcademy.postman_collection.json","Postman collection cho backend API."),
        (7,"SRS_CODE_REVIEW_BEE_ACADEMY.md","Báo cáo đối chiếu SRS và source code."),
        (8,"COPS/output/feature_coverage_matrix.xlsx","Ma trận coverage 44 requirement."),
        (9,"COPS/output/screenshot_manifest.xlsx","Danh mục ảnh chụp và trạng thái kiểm tra highlight."),
        (10,"COPS/automation/playwright/","Bộ khung automation tách theo vai trò."),
        (11,"Source commit 209b0e98a4d04028ed76dcdaa1551ad6be179542","Phiên bản source được dùng để lập báo cáo."),
    ]
    add_table(doc,["STT","Tệp/Gói","Ghi chú"],deliverables,[1.3,6.8,8.1])

    doc.add_heading("II. HƯỚNG DẪN CÀI ĐẶT", level=1)
    sections = [
        ("1. Điều kiện tiên quyết", ["Java Development Kit 21.","Node.js 20 trở lên và npm.","Maven hoặc Maven Wrapper đi kèm repository.","Supabase project cung cấp PostgreSQL, Auth và Storage."]),
        ("2. Lấy source code", ["Clone repository, sau đó checkout commit 209b0e98a4d04028ed76dcdaa1551ad6be179542."]),
        ("3. Cấu hình biến môi trường", ["Tạo backend/.env từ danh sách khóa trong README; không commit secret.","Tạo frontend/.env.local từ frontend/.env.example; đặt VITE_API_BASE_URL=http://localhost:8080."]),
        ("4. Cấu hình database", ["Tạo Supabase PostgreSQL test project.","Chạy lần lượt các migration trong backend/db/migrations và các migration bổ sung ở backend/.","Không chạy migration trên production khi chưa có phê duyệt."]),
        ("5. Chạy backend", ["cd backend", ".\\mvnw.cmd spring-boot:run", "Backend mặc định: http://localhost:8080."]),
        ("6. Chạy frontend", ["cd frontend", "npm ci", "npm run dev", "Frontend mặc định: http://localhost:3000."]),
        ("7. Kiểm tra sau cài đặt", ["Mở http://localhost:3000/courses.","Kiểm tra GET http://localhost:8080/api/courses?page=0&size=1 trả HTTP 200."]),
        ("8. Tài khoản kiểm thử", ["student.cops@beeacademy.test", "parent.cops@beeacademy.test", "teacher.cops@beeacademy.test", "admin.cops@beeacademy.test", "Mật khẩu chỉ đọc từ secret cục bộ; không ghi trong tài liệu."]),
        ("9. Sự cố thường gặp", ["Nếu frontend không gọi được backend, kiểm tra VITE_API_BASE_URL và CORS_ALLOWED_ORIGINS.","Nếu backend không kết nối database, kiểm tra SUPABASE_DB_HOST/USER/PASSWORD và SSL.","Trong lần kiểm thử này, dịch vụ xác thực tạm thời không khả dụng dù frontend/backend local đều trả HTTP 200."]),
    ]
    for title, items in sections:
        doc.add_heading(title, level=2); add_bullets(doc, items)

    doc.add_heading("III. HƯỚNG DẪN SỬ DỤNG", level=1)
    doc.add_heading("1. Tổng quan", level=2)
    doc.add_paragraph("Bee Academy là nền tảng học trực tuyến cho học sinh THCS Việt Nam. Học sinh có thể tìm kiếm, mua và học khóa học; phụ huynh theo dõi tiến độ; giáo viên quản lý nội dung và chấm điểm; quản trị viên duyệt khóa học, tài khoản, khiếu nại và chi trả. Phạm vi tài liệu bao phủ 44 requirement trong yêu cầu COPS.")

    groups = [
        ("2. Hướng dẫn cho Khách", lambda r: r[0] in {"REQ-AUTH-001","REQ-CRS-001","REQ-CRS-002","REQ-CRS-003"}),
        ("3. Hướng dẫn chung cho Người dùng", lambda r: r[0] in {"REQ-AUTH-002","REQ-AUTH-003","REQ-AUTH-004","REQ-AUTH-005"}),
        ("4. Hướng dẫn cho Học sinh", lambda r: r[0].startswith(("REQ-PAY-","REQ-LRN-","REQ-INT-"))),
        ("5. Hướng dẫn cho Phụ huynh", lambda r: r[0].startswith("REQ-PRN-")),
        ("6. Hướng dẫn cho Giáo viên", lambda r: r[0].startswith("REQ-TCH-")),
        ("7. Hướng dẫn cho Quản trị viên", lambda r: r[0].startswith("REQ-ADM-")),
    ]
    for heading, predicate in groups:
        doc.add_heading(heading, level=2)
        selected = [r for r in requirements if predicate(r)]
        for idx, req in enumerate(selected, 1):
            add_function(doc, f"{heading.split('.')[0]}.{idx}", req, screenshots_by_req)

    doc.add_heading("IV. CHỨC NĂNG CHƯA TRIỂN KHAI HOẶC CHƯA THỂ XÁC MINH", level=1)
    blocked = [r for r in requirements if r[7] != "PASSED"]
    def next_action(r):
        if r[0] == "REQ-CRS-003": return "Bổ sung lesson isFree và kiểm thử lại"
        if r[7] == "FAILED": return "Sửa lỗi đã tái hiện và chạy regression test"
        if r[7] == "PARTIAL": return "Chuẩn bị dữ liệu test kiểm soát và hoàn tất happy path"
        return "Chuẩn bị dữ liệu/điều kiện tiên quyết rồi kiểm thử lại"
    rows = [(r[0],r[1],r[7],r[8],next_action(r)) for r in blocked]
    add_table(doc,["REQ ID","Chức năng","Lý do/Trạng thái","Bằng chứng","Hành động đề xuất"],rows,[2.4,3.4,2.2,5.0,3.2])

    doc.add_heading("V. TỔNG KẾT KIỂM THỬ VÀ QA", level=1)
    add_table(doc,["Chỉ số","Kết quả","Ghi chú"],[
        ("Tổng requirement",44,"Theo phạm vi người dùng yêu cầu."),
        ("Có trong tài liệu",44,"Không thiếu REQ."),
        ("Passed",8,"Đăng nhập, tra cứu công khai và một số màn hình đọc/empty state đã xác minh."),
        ("Partial",15,"Đã mở được UI hoặc một phần luồng nhưng không tạo dữ liệu/email/thông báo thật."),
        ("Blocked",19,"Thiếu khóa học đã mua, liên kết con, question bank, bài nộp hoặc dữ liệu nghiệp vụ."),
        ("Failed",2,"Đăng xuất không phản hồi ở một số trang con; xuất Excel chi trả không tạo file."),
        ("Ảnh highlighted",len(screenshot_rows),"Ảnh UI local đã được khoanh vùng bằng viền đỏ và kiểm tra trước khi chèn."),
        ("Frontend lint","Passed","npm run lint."),
        ("Backend compile","Passed","mvn -q -DskipTests compile."),
        ("HTTP frontend/backend","200/200","Kiểm tra local ngày 22/07/2026."),
    ],[4.3,3.0,8.9])
    doc.add_paragraph("Kết luận: dịch vụ xác thực đã hoạt động trở lại và cả 4 vai trò đăng nhập thành công. Retest ghi nhận 8 Passed, 15 Partial, 19 Blocked và 2 Failed. Không có chức năng bị đánh dấu Passed khi chưa quan sát đủ kết quả; không sử dụng ảnh giả, thanh toán thật, OTP/email thật, chuyển khoản hoặc khóa tài khoản thật.")

    doc.core_properties.title = "Bee Academy - Báo cáo COPS cuối cùng"
    doc.core_properties.subject = "Final Release Document"
    doc.core_properties.author = "Bee Academy QA Team"
    doc.core_properties.keywords = "Bee Academy, COPS, QA, User Manual, Final Release"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
