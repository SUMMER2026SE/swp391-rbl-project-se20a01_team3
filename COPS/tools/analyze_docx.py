from __future__ import annotations

import json
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def twips_to_cm(value):
    return None if value is None else round(value / 360000, 3)


def main(path_s: str, out_s: str) -> None:
    path = Path(path_s)
    doc = Document(path)
    sections = []
    for s in doc.sections:
        sections.append({
            "page_width_cm": twips_to_cm(s.page_width),
            "page_height_cm": twips_to_cm(s.page_height),
            "top_margin_cm": twips_to_cm(s.top_margin),
            "bottom_margin_cm": twips_to_cm(s.bottom_margin),
            "left_margin_cm": twips_to_cm(s.left_margin),
            "right_margin_cm": twips_to_cm(s.right_margin),
            "header_distance_cm": twips_to_cm(s.header_distance),
            "footer_distance_cm": twips_to_cm(s.footer_distance),
        })
    styles = {}
    for name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Caption"]:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        font = st.font
        pf = st.paragraph_format
        styles[name] = {
            "font": font.name,
            "size_pt": None if font.size is None else round(font.size.pt, 2),
            "bold": font.bold,
            "italic": font.italic,
            "color": None if font.color is None or font.color.rgb is None else str(font.color.rgb),
            "alignment": None if pf.alignment is None else str(pf.alignment),
            "space_before_pt": None if pf.space_before is None else round(pf.space_before.pt, 2),
            "space_after_pt": None if pf.space_after is None else round(pf.space_after.pt, 2),
            "line_spacing": str(pf.line_spacing),
            "keep_with_next": pf.keep_with_next,
            "page_break_before": pf.page_break_before,
        }
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            paragraphs.append({"index": i, "style": p.style.name if p.style else None, "text": text})
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows[:8]:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": ti, "rows": len(table.rows), "cols": len(table.columns), "sample": rows})
    headers = []
    footers = []
    for si, sec in enumerate(doc.sections):
        headers.append({"section": si, "text": " | ".join(p.text.strip() for p in sec.header.paragraphs if p.text.strip())})
        footers.append({"section": si, "text": " | ".join(p.text.strip() for p in sec.footer.paragraphs if p.text.strip())})
    payload = {
        "source": str(path.resolve()),
        "sections": sections,
        "styles": styles,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "paragraphs": paragraphs,
        "tables": tables,
        "headers": headers,
        "footers": footers,
        "settings_update_fields": doc.settings.element.find(qn("w:updateFields")) is not None,
    }
    Path(out_s).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
