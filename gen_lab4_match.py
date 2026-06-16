# -*- coding: utf-8 -*-
"""Sinh lại Lab4_English (Test Plan tiếng Anh) và Lab4_PhanB_C (tiếng Việt)
theo ĐÚNG style của bản tiếng Việt gen_lab4.py: Times New Roman, trang bìa FPT,
mục lục, header bảng 1F3864, hàng xen kẽ EBF3FB."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'Times New Roman'

# ── helpers (giống hệt gen_lab4.py) ──────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def ct(cell, text, bold=False, center=False, sz=11, color=None, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


def th(table, headers, bg='1F3864', sz=11):
    row = table.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        set_cell_bg(c, bg)
        ct(c, h, bold=True, center=True, sz=sz, color='FFFFFF')


def td(table, data, alt_bg='EBF3FB', sz=11):
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            cell = table.cell(r + 1, c)
            if r % 2 == 1:
                set_cell_bg(cell, alt_bg)
            ct(cell, str(val) if val is not None else '', sz=sz)


def set_col_widths(table, widths_cm):
    for i, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)


def mktbl(doc, rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    return t


def table(doc, headers, rows, widths, header_bg='1F3864'):
    doc.add_paragraph()
    t = mktbl(doc, len(rows) + 1, len(headers))
    th(t, headers, bg=header_bg)
    td(t, rows)
    set_col_widths(t, widths)
    return t


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(15)
    run.font.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(13)
    run.font.bold = True
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True
    return p


def body(doc, text, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(13)
    return p


def bul(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'• {text}')
    run.font.name = FONT
    run.font.size = Pt(13)
    return p


def center_line(doc, text, sz, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    return p


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2.5)
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(13)
    return doc


def cover(doc, univ, faculty, big1, big2, std_line, project, subtitle, info_rows):
    for _ in range(2):
        doc.add_paragraph()
    center_line(doc, univ, 14, bold=True)
    center_line(doc, faculty, 13)
    for _ in range(5):
        doc.add_paragraph()
    center_line(doc, big1, 20, bold=True)
    center_line(doc, big2, 17, bold=True)
    doc.add_paragraph()
    center_line(doc, std_line, 13, italic=True)
    for _ in range(2):
        doc.add_paragraph()
    center_line(doc, project, 16, bold=True)
    center_line(doc, subtitle, 13, italic=True)
    for _ in range(3):
        doc.add_paragraph()
    ctbl = mktbl(doc, len(info_rows), 2)
    for i, (k, v) in enumerate(info_rows):
        rw = ctbl.rows[i]
        set_cell_bg(rw.cells[0], 'D6E4F7')
        ct(rw.cells[0], k, bold=True, sz=12)
        ct(rw.cells[1], v, sz=12)
    ctbl.columns[0].width = Cm(5.5)
    ctbl.columns[1].width = Cm(9)
    doc.add_page_break()


def toc(doc, title, entries):
    center_line(doc, title, 16, bold=True)
    doc.add_paragraph()
    for is_sub, label in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if is_sub:
            p.paragraph_format.left_indent = Cm(1.2)
        r = p.add_run(label)
        r.font.name = FONT
        r.font.size = Pt(13)
        r.font.bold = not is_sub
    doc.add_page_break()


# =====================================================================
# ============== FILE 1: TEST PLAN – ENGLISH ==========================
# =====================================================================
def build_english():
    doc = new_doc()
    cover(
        doc,
        'FPT UNIVERSITY',
        'Faculty of Information Technology',
        'LAB REPORT',
        'LAB 4 – CREATING A TEST PLAN WITH AI SUPPORT',
        'Standard: IEEE 829-2008  |  Course: SWT301 – Software Testing',
        'PROJECT: BEE ACADEMY',
        'Online Learning Platform for Lower-Secondary Students (Grades 6–9)',
        [
            ('Course', 'SWT301 – Software Testing'),
            ('Topic', 'Topic 5 – Test Management'),
            ('CLO', 'CLO5 – Build a comprehensive AI-assisted test plan'),
            ('Standard', 'IEEE 829-2008'),
            ('Practice Project', 'Bee Academy – Online Learning Platform v1.0.0'),
            ('Student', 'Vo Van Thanh Dat'),
            ('Student ID', 'DE190211'),
            ('Date', '2026-06-12'),
        ],
    )
    toc(doc, 'TABLE OF CONTENTS', [
        (False, 'PART A – TEST PLAN PER IEEE 829-2008'),
        (True, 'Section 1  –  Test Plan Identifier'),
        (True, 'Section 2  –  Introduction'),
        (True, 'Section 3  –  Test Items'),
        (True, 'Section 4  –  Features to be Tested'),
        (True, 'Section 5  –  Features Not to be Tested'),
        (True, 'Section 6  –  Approach (Test Strategy)'),
        (True, 'Section 7  –  Item Pass/Fail Criteria'),
        (True, 'Section 8  –  Suspension & Resumption Requirements'),
        (True, 'Section 9  –  Test Deliverables'),
        (True, 'Section 10 –  Testing Tasks'),
        (True, 'Section 11 –  Environmental Needs'),
        (True, 'Section 12 –  Responsibilities'),
        (True, 'Section 13 –  Staffing and Training Needs'),
        (True, 'Section 14 –  Schedule'),
        (True, 'Section 15 –  Risks and Contingencies'),
        (True, 'Section 16 –  Approvals'),
    ])

    # Project overview
    p = doc.add_paragraph()
    r = p.add_run('Practice Project – Bee Academy')
    r.font.name = FONT; r.font.size = Pt(16); r.font.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F3864')
    pBdr.append(bot); pPr.append(pBdr)

    h2(doc, 'System Description')
    body(doc, 'Bee Academy is an online learning platform for Vietnamese lower-secondary students '
              '(grades 6–9), offering video courses by subject. The system serves four user groups:')
    for item in [
        'Student: Browse and purchase courses, take per-chapter quizzes, track progress, and receive completion certificates',
        'Teacher: Create and manage courses, upload lecture videos, configure quizzes, and view revenue',
        'Parent: Link to a student account, monitor learning progress, and contact teachers',
        'Admin: Approve courses, manage user accounts, confirm teacher payouts, and generate reports',
    ]:
        bul(doc, item)

    h2(doc, 'Version Information')
    table(doc, ['Information', 'Detail'], [
        ('System Name', 'Bee Academy – Online Learning Platform'),
        ('Test Version', 'v1.0.0'),
        ('Environment', 'Web application (React 19 + Vite 6 frontend / Spring Boot 3.2 backend / PostgreSQL – Supabase)'),
        ('Test Start Date', '2026-06-12'),
        ('Test End Date', '2026-07-15'),
        ('Team Size', '3 testers + 1 test lead'),
    ], [4.5, 10])

    h2(doc, 'Core Modules')
    modules = [
        ('Module 1: Authentication & Authorization',
         ['- Login / logout (email/password + Google OAuth)',
          '- Account registration with email OTP',
          '- JWT ES256 (ECDSA P-256) via Supabase GoTrue',
          '- Role-based access: STUDENT / TEACHER / PARENT / ADMIN']),
        ('Module 2: Course Catalog & Enrollment',
         ['- Browse and search courses, filter by category / grade / price',
          '- View course detail, watch free preview lessons (isFree=true)',
          '- Enroll in a course after successful payment']),
        ('Module 3: Payment (VNPay / MoMo)',
         ['- Add to cart (login required)',
          '- Pay via VNPay / MoMo – funds go to the company account',
          '- System records revenue_splits; view order history']),
        ('Module 4: Teacher Portal',
         ['- CRUD courses / chapters / lessons',
          '- Upload lecture videos (Supabase private bucket)',
          '- Submit course for admin review (DRAFT → PENDING)']),
        ('Module 5: Quiz & Assessment',
         ['- Question bank (CRUD, classified Easy/Medium/Hard)',
          '- Per-chapter quiz config (count, time, passing score)',
          '- Students take quizzes – graded via JSONB snapshot']),
        ('Module 6: Admin Portal',
         ['- Approve / reject / request revision of courses',
          '- Manage user accounts',
          '- Export payout Excel and confirm teacher transfers']),
        ('Module 7: Parent Portal',
         ['- Send link invitation via student email',
          '- Monitor learning progress (only when link is ACTIVE)',
          '- Contact teacher, view payment history for the child']),
        ('Module 8: Student Dashboard',
         ['- Personal profile, change password, upload avatar',
          '- Favorites list, order history',
          '- Handle link invitations from parents']),
    ]
    for title_txt, lines in modules:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title_txt)
        r.font.name = FONT; r.font.size = Pt(13); r.font.bold = True
        for line in lines:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.2)
            p2.paragraph_format.space_after = Pt(1)
            r2 = p2.add_run(line)
            r2.font.name = FONT; r2.font.size = Pt(12)

    doc.add_page_break()

    # ── PART A ──
    center_line(doc, 'PART A – TEST PLAN PER IEEE 829-2008', 16, bold=True)
    doc.add_paragraph()

    h1(doc, 'SECTION 1 – TEST PLAN IDENTIFIER')
    body(doc, 'This document identifies the official Master Test Plan (MTP) for the Bee Academy project, '
              'version v1.0.0, prepared in accordance with the IEEE 829-2008 standard.')
    table(doc, ['Field', 'Value'], [
        ('Document ID', 'TP-BEEACADEMY-2026-001'),
        ('Project Name', 'Bee Academy – Online Learning Platform'),
        ('System Version', 'v1.0.0'),
        ('Test Plan Version', '1.0'),
        ('Author', 'Vo Van Thanh Dat'),
        ('Created Date', '2026-06-12'),
        ('Review Date', '2026-06-19'),
        ('Status', 'Draft'),
    ], [5, 9.5])

    h1(doc, 'SECTION 2 – INTRODUCTION')
    h2(doc, '2.1  Purpose')
    body(doc, 'This Test Plan defines the scope, approach, resources, and schedule of the testing activities '
              'for Bee Academy v1.0.0 – an online learning platform for Vietnamese lower-secondary students '
              '(grades 6–9). The objective is to ensure the system meets its functional and non-functional '
              'requirements before official release.')
    h2(doc, '2.2  Scope')
    body(doc, 'This Test Plan applies to the System Testing and UAT phases of Bee Academy v1.0.0, covering eight '
              'core modules: Authentication & Authorization, Course Catalog, Enrollment & Payment, Student '
              'Dashboard, Teacher Portal, Quiz & Assessment, Admin Portal, and Parent Portal. The test '
              'environment uses React 19 + Vite 6 (frontend) together with Spring Boot 3.2 (backend) and a '
              'PostgreSQL database provided through Supabase.')
    h2(doc, '2.3  Intended Audience')
    for item in ['Test Lead and QA Engineers – execute the testing activities',
                 'Development Team – fix defects and support the environment',
                 'Project Manager – track progress and make decisions',
                 'Product Owner / Stakeholders – approve UAT and sign off the release']:
        bul(doc, item)
    h2(doc, '2.4  References')
    for item in ['IEEE 829-2008 – IEEE Standard for Software and System Test Documentation',
                 'CLAUDE.md – Bee Academy Architecture & Project Documentation',
                 'BEE ACADEMY.md – Use Case v6.5 (48 use cases, 9 modules, 7 actors)',
                 'ISTQB Foundation Level Syllabus – Chapter 5: Test Management',
                 'Spring Boot 3.2 Security Documentation – JWT ES256 / ECDSA P-256',
                 'Supabase Documentation – Storage Buckets & RLS Policies']:
        bul(doc, item)

    h1(doc, 'SECTION 3 – TEST ITEMS')
    body(doc, 'The following table lists the modules/components of Bee Academy v1.0.0 included in the testing '
              'scope, together with their version and readiness status.')
    table(doc, ['Item ID', 'Module', 'Version', 'Brief Description', 'Status'], [
        ('TI-01', 'Authentication & Authorization', 'v1.0.0', 'Email/password login, OTP, Google OAuth (JWT ES256 + Supabase GoTrue)', 'Ready'),
        ('TI-02', 'Course Catalog & Detail', 'v1.0.0', 'Browse, search, filter courses; view detail and preview lessons', 'Ready'),
        ('TI-03', 'Enrollment & Payment', 'v1.0.0', 'Cart, VNPay/MoMo payment, view order history', 'In Dev'),
        ('TI-04', 'Student Dashboard', 'v1.0.0', 'Profile, account, avatar upload, favorites', 'Ready'),
        ('TI-05', 'Teacher Portal', 'v1.0.0', 'CRUD courses/chapters/lessons; upload video (private) + docs; submit for review', 'Partial'),
        ('TI-06', 'Quiz & Assessment System', 'v1.0.0', 'Per-chapter quiz config, question bank, attempt + grading via JSONB snapshot', 'Ready'),
        ('TI-07', 'Admin Portal', 'v1.0.0', 'Course approval, user management, revenue reports, payout confirmation', 'Partial'),
        ('TI-08', 'Parent Portal', 'v1.0.0', 'Link to student, monitor progress, contact teacher', 'Ready'),
    ], [1.5, 3.2, 1.8, 6.5, 1.5])

    h1(doc, 'SECTION 4 – FEATURES TO BE TESTED')
    body(doc, 'The table below lists at least 18 features within the testing scope, evenly distributed across '
              'the 8 core modules.')
    table(doc, ['Feature ID', 'Feature Name', 'Module', 'Priority', 'Brief Description'], [
        ('F-01', 'Email/password login', 'Auth', 'High', 'Valid/invalid credentials; verify the returned JWT'),
        ('F-02', 'Account registration with email OTP', 'Auth', 'High', 'Flow: send OTP → verify → create STUDENT profile'),
        ('F-03', 'Google OAuth login (ES256 JWT)', 'Auth', 'High', 'Google consent → Supabase callback → sync profile'),
        ('F-04', 'Role-based access control (RBAC)', 'Auth', 'High', 'Student/Teacher/Parent/Admin correct route guards'),
        ('F-05', 'Refresh token & session timeout', 'Auth', 'Medium', 'Expired token auto-refresh; logout revokes token'),
        ('F-06', 'Browse course list + filters', 'Course', 'High', 'Filter by category, price, grade; pagination'),
        ('F-07', 'View detail and free preview lesson', 'Course', 'High', 'isFree=true allows viewing; private video requires enrollment'),
        ('F-08', 'Add to cart – login required', 'Payment', 'High', 'Not logged in → redirect to /login with state.from'),
        ('F-09', 'VNPay/MoMo payment (sandbox)', 'Payment', 'High', 'Create order → redirect payment → webhook → enroll'),
        ('F-10', 'View order history', 'Dashboard', 'Medium', 'Order list, status, purchase date'),
        ('F-11', 'Update profile and upload avatar', 'Dashboard', 'Medium', 'Edit name, bio; upload image to Supabase Storage'),
        ('F-12', 'Teacher create/edit/delete course', 'Teacher', 'High', 'CRUD with chapter/lesson; status DRAFT→PENDING'),
        ('F-13', 'Upload lecture video (private bucket)', 'Teacher', 'High', 'Upload to Supabase private; store storagePath'),
        ('F-14', 'Submit course for admin review', 'Teacher', 'High', 'Change status DRAFT→PENDING; admin receives notification'),
        ('F-15', 'Configure chapter quiz + question bank', 'Quiz', 'High', 'CRUD questions; set count, time, passing score'),
        ('F-16', 'Student takes quiz + views result', 'Quiz', 'High', 'Random pick from bank; JSONB snapshot; auto-grading'),
        ('F-17', 'Admin approve / reject / request revision', 'Admin', 'High', 'Review content; approve→PUBLISHED; reject with reason'),
        ('F-18', 'Parent links and monitors progress', 'Parent', 'Medium', 'Send email invite; after ACTIVE, view progress'),
    ], [1.4, 4, 2, 1.8, 5.3])

    h1(doc, 'SECTION 5 – FEATURES NOT TO BE TESTED')
    body(doc, 'The following features are excluded from the v1.0.0 testing scope, with specific justifications:')
    table(doc, ['Excluded Feature', 'Justification'], [
        ('Penetration Testing / Security Audit', 'Requires a dedicated security specialist; to be performed before production go-live'),
        ('Load Testing / Performance Testing', 'Requires a staging environment configured like production; not feasible in local dev'),
        ('Cross-browser (Safari, Edge)', 'v1.0 scope supports only Chrome + Firefox; Safari/Edge tested in a later sprint'),
        ('Mobile responsiveness (< 768px)', 'Mobile version will be tested separately in Sprint 3 once responsive design is complete'),
        ('Real MoMo integration (production API)', 'Only VNPay sandbox is available; MoMo production API requires a contract and separate testing'),
        ('Certificate feature (UC42-43)', 'Not developed in v1.0.0; planned for Sprint 4'),
        ('Messaging / AI chat with teacher (UC20-21)', 'Chat module not developed; requires WebSocket or third-party integration'),
        ('Teacher bank account feature (UC45-46)', 'Backend not complete; manual admin payout will be tested when the feature is complete'),
    ], [5.5, 9])

    h1(doc, 'SECTION 6 – APPROACH (TEST STRATEGY)')
    body(doc, 'The testing strategy for Bee Academy v1.0.0 applies multi-level testing, combining black-box and '
              'white-box techniques to ensure comprehensive quality assurance before release.')
    h2(doc, '6.1  Test Levels')
    table(doc, ['Level', 'Owner', 'Technique', 'Tools'], [
        ('Unit Testing', 'Dev Team', 'White-box, Statement Coverage (≥80%)', 'JUnit 5 / Jest'),
        ('Integration Testing', 'Dev + QA', 'API Testing, Black-box, Contract Test', 'Postman / REST Assured'),
        ('System Testing', 'QA Team', 'Black-box: EP, BVA; Exploratory', 'Playwright / JIRA'),
        ('Acceptance Testing', 'PO + Stakeholders', 'Scenario-based, User Acceptance', 'Manual / TestRail'),
    ], [3.5, 3, 4, 4])
    h2(doc, '6.2  Test Types')
    table(doc, ['Type', 'Objective', 'Applied Scope'], [
        ('Functional Testing', 'Verify business logic against Use Case v6.5', 'All 18 features in Section 4'),
        ('Security Testing', 'Test JWT authentication, authorization bypass, SQL injection', 'Auth, Teacher Portal, Admin Portal modules'),
        ('Usability Testing', 'Test intuitive UX/UI, desktop responsiveness', 'Landing Page, Course Detail, Student Dashboard'),
        ('Regression Testing', 'Ensure bug fixes do not break existing features', 'After each Critical/Major defect fix'),
    ], [3.5, 6, 5])
    h2(doc, '6.3  Test Design Techniques')
    for item in ['Equivalence Partitioning (EP) – for input validation (email, password, price)',
                 'Boundary Value Analysis (BVA) – for price range, grade filter, upload file size',
                 'Decision Table – for the payment flow (VNPay callback: success/fail/cancel/timeout)',
                 'State Transition – for the course lifecycle (DRAFT→PENDING→PUBLISHED/REJECTED)',
                 'Error Guessing – expired JWT, invalid OTP, expired Supabase signed URL',
                 'Exploratory Testing – find defects not covered by test cases via session-based exploration']:
        bul(doc, item)
    h2(doc, '6.4  Entry & Exit Criteria')
    table(doc, ['Level', 'Entry Criteria', 'Exit Criteria'], [
        ('System Testing', 'Build deployed successfully to test env; unit test pass ≥80%; test cases & data prepared', 'Test case pass rate ≥95%; no open Critical defect; Test Summary Report produced'),
        ('Acceptance Testing', 'System Testing passed its exit criteria; regression testing complete; UAT environment ready', 'All UAT scenarios PASS; PO signs off; release checklist complete'),
    ], [2.5, 6.5, 5.5])

    h1(doc, 'SECTION 7 – ITEM PASS/FAIL CRITERIA')
    h2(doc, '7.1  Pass Criteria')
    for item in ['Test case PASS rate ≥ 95% of all executed test cases',
                 'No Critical or Blocker defect in Open status',
                 'No Major defect remaining Open after D-3 before release',
                 'All Major defects have a documented workaround in JIRA',
                 'Authentication module: no Security defect of any kind (JWT bypass, SQL injection)',
                 'Payment module: the entire VNPay sandbox payment flow completes successfully',
                 'Test Summary Report approved by the Test Lead']:
        bul(doc, item)
    h2(doc, '7.2  Fail Criteria (stop release)')
    for item in ['Test case PASS rate < 90%',
                 'At least one unfixed Critical/Blocker defect exists',
                 'The Authentication module has any open security defect',
                 'The VNPay payment flow cannot complete (payment flow broken)',
                 'User data is lost or incorrectly overwritten after a transaction',
                 'The system crashes or fails to start in the test environment']:
        bul(doc, item)
    h2(doc, '7.3  Defect Severity Definition')
    table(doc, ['Severity', 'Definition', 'Example in Bee Academy'], [
        ('Critical / Blocker', 'System unusable; no workaround', 'Cannot log in; payment crash; data loss'),
        ('Major', 'Important function broken; temporary workaround exists', 'Quiz grades incorrectly; video fails to play'),
        ('Minor', 'Function works but has a small defect; low impact', 'Typo in text; minor UI misalignment'),
        ('Cosmetic', 'UI defect not affecting functionality', 'Icon shows wrong color; uneven spacing'),
    ], [2.5, 5.5, 6.5])

    h1(doc, 'SECTION 8 – SUSPENSION CRITERIA AND RESUMPTION REQUIREMENTS')
    body(doc, 'The table below describes situations that require suspending testing and the conditions required '
              'to resume.')
    table(doc, ['Situation', 'Suspension Condition', 'Resumption Condition'], [
        ('Build/Deploy failure', 'Test env cannot deploy the new version; error rate > 50%', 'Build passes; smoke test of 5 core features succeeds; Dev confirms'),
        ('Critical Blocker appears', '≥1 Critical defect found affecting the module under test', 'Dev fixes and merges; QA verifies the fix; regression test of that module passes'),
        ('Test environment down', 'Backend or Supabase unreachable for > 30 continuous minutes', 'Environment restored; health check passes; Test Lead confirms'),
        ('Test data corrupted', 'Test data is wrong, missing, or overwritten due to a script/migration error', 'Restore test data from backup; QA confirms data integrity'),
    ], [3.5, 5.5, 5.5])
    body(doc, 'When suspended, the Test Lead must record the cause in the Daily Report, update the status in '
              'JIRA, and notify the Project Manager within 1 hour.')

    h1(doc, 'SECTION 9 – TEST DELIVERABLES')
    body(doc, 'List of deliverables produced throughout the testing process:')
    table(doc, ['Phase', 'Deliverable', 'Responsible', 'Deadline'], [
        ('Pre-testing', 'Master Test Plan (this document)', 'Test Lead', '2026-06-12'),
        ('Pre-testing', 'Test Cases (Excel / TestRail)', 'QA Team', '2026-06-17'),
        ('Pre-testing', 'Test Data Script (SQL + seed data)', 'QA Team', '2026-06-16'),
        ('Pre-testing', 'Test Environment Setup Guide', 'Dev + QA', '2026-06-15'),
        ('During testing', 'Daily Test Progress Report', 'Test Lead', 'Daily'),
        ('During testing', 'Defect Report (JIRA tickets)', 'QA Team', 'When found'),
        ('During testing', 'Test Execution Log', 'QA Team', 'Daily'),
        ('During testing', 'Change Request Log (if requirements change)', 'Test Lead', 'As needed'),
        ('Post-testing', 'Test Summary Report', 'Test Lead', '2026-07-14'),
        ('Post-testing', 'Defect Metrics & Quality Report', 'Test Lead', '2026-07-14'),
    ], [2.8, 5, 3.5, 3.2])

    h1(doc, 'SECTION 10 – TESTING TASKS')
    body(doc, 'The table below decomposes the testing activities into specific tasks, with execution order and '
              'clear dependencies.')
    table(doc, ['Task ID', 'Task Name', 'Predecessor', 'Effort', 'Role', 'Start', 'End'], [
        ('T01', 'Analyze requirements & Use Case v6.5', '–', '8h', 'Test Lead', '06/12', '06/13'),
        ('T02', 'Design test cases (18 features)', 'T01', '24h', 'QA Team', '06/14', '06/17'),
        ('T03', 'Prepare test data & seed scripts', 'T01', '8h', 'QA Team', '06/14', '06/16'),
        ('T04', 'Set up and configure test environment', 'T01', '4h', 'Dev + QA', '06/12', '06/14'),
        ('T05', 'System Test – Auth + Course + Dashboard', 'T02,T03,T04', '20h', 'QA Team', '06/18', '06/23'),
        ('T06', 'System Test – Payment + Teacher Portal', 'T05', '20h', 'QA Team', '06/24', '06/30'),
        ('T07', 'System Test – Quiz + Admin + Parent', 'T06', '20h', 'QA Team', '07/01', '07/07'),
        ('T08', 'Regression Testing (after fixing defects)', 'T07', '16h', 'QA Team', '07/08', '07/10'),
        ('T09', 'User Acceptance Testing (UAT)', 'T08', '16h', 'PO + Stakeholders', '07/11', '07/13'),
        ('T10', 'Write Test Summary Report & Sign-off', 'T09', '4h', 'Test Lead', '07/14', '07/15'),
    ], [1.3, 5, 2.5, 1.5, 2.5, 1.5, 1.5])

    h1(doc, 'SECTION 11 – ENVIRONMENTAL NEEDS')
    h2(doc, '11.1  Hardware Requirements')
    table(doc, ['Component', 'Minimum Configuration'], [
        ('Test Server (Backend)', 'CPU: 4 cores, RAM: 8GB, Storage: 100GB SSD; or Cloud: Supabase + Railway'),
        ('Tester Workstation', 'CPU: 2 cores, RAM: 8GB, OS: Windows 11 / Ubuntu 22.04'),
        ('CI/CD Pipeline', 'GitHub Actions (2 workers, 4GB RAM per runner)'),
    ], [4.5, 10])
    h2(doc, '11.2  Software Requirements')
    table(doc, ['Software', 'Version', 'Purpose'], [
        ('OS', 'Windows 11 / Ubuntu 22.04', 'Test execution environment'),
        ('Node.js', '20 LTS', 'Build and run React 19 + Vite frontend'),
        ('Java JDK', '17', 'Spring Boot 3.2 backend runtime'),
        ('Maven', '3.9+', 'Backend build tool'),
        ('PostgreSQL (Supabase)', '15', 'Database; uses a Supabase test project'),
        ('Chrome', 'Latest', 'Primary browser (System Test)'),
        ('Firefox', 'Latest', 'Secondary browser (Regression)'),
        ('Postman', 'v10+', 'API testing – Auth endpoints, JWT validation'),
        ('Playwright', '1.40+', 'UI automation – E2E test scenarios'),
        ('JIRA', 'Cloud', 'Defect tracking and test management'),
        ('Git / GitHub', 'Latest', 'Version control; CI/CD integration'),
    ], [3.5, 3.5, 7.5])
    h2(doc, '11.3  Test Environment Configuration')
    for item in ['Frontend: http://localhost:3000 (Vite dev server, port fixed in vite.config.ts)',
                 'Backend: http://localhost:8080 (Spring Boot, port configured in backend/.env)',
                 'Database: Supabase PostgreSQL (separate test project, isolated from production)',
                 'Storage: Supabase Storage – 2 buckets: course-videos (private) + course-docs (public)',
                 'JWT: ES256 (ECDSA P-256) verified via the Supabase JWKS endpoint',
                 'CORS: allow origin http://localhost:3000',
                 'DEV_MODE=true: OTP logged to console instead of sending a real email']:
        bul(doc, item)
    h2(doc, '11.4  Test Data Requirements')
    for item in ['5 Student accounts (valid credentials + edge cases: wrong password, locked account)',
                 '3 Teacher accounts (1 with an approved course, 1 pending, 1 newly created)',
                 '2 Parent accounts (1 linked to a student, 1 not linked)',
                 '1 Admin account',
                 '10 sample courses (DRAFT: 3, PENDING: 3, PUBLISHED: 4, REJECTED: 1)',
                 'VNPay sandbox test payment cards: success + fail + timeout scenarios',
                 '50 sample question-bank items (10 per module, distributed Easy/Medium/Hard)',
                 'Sample video file (≤100MB, MP4) and PDF document (≤10MB) for upload testing']:
        bul(doc, item)

    h1(doc, 'SECTION 12 – RESPONSIBILITIES')
    body(doc, 'RACI matrix assigning responsibility for each testing activity '
              '(R=Responsible, A=Accountable, C=Consulted, I=Informed):')
    table(doc, ['Activity', 'Test Lead', 'QA Engineer', 'Developer', 'PM', 'PO'], [
        ('Write Test Plan', 'R/A', 'C', 'C', 'I', 'I'),
        ('Design Test Cases', 'A', 'R', 'C', 'I', 'C'),
        ('Prepare Test Data', 'A', 'R', 'C', 'I', '–'),
        ('Set up environment', 'A', 'R', 'R/A', 'I', '–'),
        ('Execute Testing', 'A', 'R', 'I', 'I', 'I'),
        ('Bug Report & Tracking', 'A', 'R', 'I', 'I', 'I'),
        ('Fix Defects', 'I', 'C', 'R/A', 'I', '–'),
        ('Test Summary Report', 'R/A', 'C', 'I', 'I', 'I'),
    ], [4.5, 2, 2.2, 2.2, 1.5, 1.5])

    h1(doc, 'SECTION 13 – STAFFING AND TRAINING NEEDS')
    h2(doc, '13.1  Team Composition')
    table(doc, ['Role', 'Headcount', 'Allocation', 'Required Skills', 'Training Needed'], [
        ('Test Lead', '1', '100%', 'IEEE 829-2008; ISTQB Foundation; API testing; JIRA', 'None'),
        ('QA Engineer', '2', '100%', 'Manual testing; basic SQL; Chrome DevTools; Postman', 'Playwright workshop (2 days); Supabase API basics (1 day)'),
        ('Dev (support)', '2', '20%', 'JUnit 5 unit test; code review; environment setup', 'None'),
    ], [2.5, 1.8, 2, 5, 3.2])
    h2(doc, '13.2  Training Plan')
    table(doc, ['Training Course', 'Audience', 'Duration', 'Timing'], [
        ('Playwright E2E Automation Workshop', 'QA Engineers', '2 days', 'Week 1 (2026-06-12 to 06-13)'),
        ('Supabase Storage & JWT ES256 Overview', 'QA Team', '3 hours', '2026-06-14'),
    ], [5.5, 2.5, 2.5, 4])

    h1(doc, 'SECTION 14 – SCHEDULE')
    h2(doc, '14.1  Milestones')
    table(doc, ['Milestone', 'Date', 'Description'], [
        ('MS-01', '2026-06-12', 'Kickoff, Test Plan approved'),
        ('MS-02', '2026-06-17', 'Test Cases & Test Data ready; test environment ready'),
        ('MS-03', '2026-07-07', 'System Testing complete (all 8 modules)'),
        ('MS-04', '2026-07-10', 'Regression Testing complete'),
        ('MS-05', '2026-07-13', 'UAT complete; PO sign-off'),
        ('MS-06', '2026-07-15', 'Test Summary Report delivered'),
    ], [2.5, 3.5, 8.5])
    h2(doc, '14.2  Weekly Schedule (Gantt Chart)')
    table(doc, ['Phase / Activity', 'W1 (6/12-16)', 'W2 (6/17-23)', 'W3 (6/24-30)', 'W4 (7/1-10)', 'W5 (7/11-15)'], [
        ('Planning & Test Design', 'XX', 'XX', '', '', ''),
        ('Environment Setup', 'XX', '', '', '', ''),
        ('System Test – Auth & Course', '', 'XX', '', '', ''),
        ('System Test – Pay & Teacher', '', '', 'XX', '', ''),
        ('System Test – Quiz,Admin,Parent', '', '', '', 'XX', ''),
        ('Regression Testing', '', '', '', 'XX', ''),
        ('UAT & Test Closure', '', '', '', '', 'XX'),
    ], [4.5, 2.2, 2.2, 2.2, 2.2, 2.2])

    h1(doc, 'SECTION 15 – RISKS AND CONTINGENCIES')
    body(doc, 'The table below identifies testing risks, rates their level (L×I), and defines mitigation / '
              'contingency plans. (L=Likelihood: H/M/L  |  I=Impact: H/M/L)')
    table(doc, ['Risk ID', 'Risk', 'L', 'I', 'Level', 'Mitigation', 'Contingency'], [
        ('R01', 'Unstable test environment (Supabase local dev, CORS errors)', 'M', 'H', 'High', 'Prepare a separate staging environment; docker-compose backup', 'Use a hosted Supabase test project instead of localhost'),
        ('R02', 'JWT ES256 key rotation logs out all users unexpectedly', 'L', 'H', 'Medium', 'Do not rotate the JWKS key during testing; note it in env setup', 'Restart backend and re-login; log the incident'),
        ('R03', 'VNPay sandbox does not reflect production behavior', 'H', 'M', 'High', 'Use both success + fail test cards; document any discrepancies', 'Test the payment flow with a mock server if the sandbox is unstable'),
        ('R04', 'Requirements change during testing (Use Case v6.5 not frozen)', 'M', 'H', 'High', 'Freeze requirements from W1; all changes go through a Change Request', 'Re-scope test cases; adjust the schedule if needed'),
        ('R05', 'Staff shortage due to sick leave or personal reasons', 'L', 'H', 'Medium', 'Cross-training among QA Engineers; clear test case documentation', 'Test Lead temporarily supports execution; prioritize High features'),
        ('R06', 'Build delays from the Dev Team affect the schedule', 'M', 'M', 'Medium', 'Daily stand-up; automated CI/CD; early integration merge', 'Test Ready modules first; delay testing In-Dev modules'),
        ('R07', 'Test data corrupted after testing the payment flow', 'M', 'H', 'High', 'Use transaction rollback or a reset script after each payment test session', 'Restore the test DB from the seed script; log the incident'),
        ('R08', 'Supabase Storage signed URL expires (TTL 1h) mid-test session', 'L', 'M', 'Low', 'Regenerate the signed URL before each video access test; note TTL in the test case', 'Increase TTL in the test env or test video within a < 1h session'),
        ('R09', 'Private video bucket misconfigured → public access', 'L', 'H', 'Medium', 'Verify the RLS policy on the Supabase dashboard before testing', 'Fix the policy immediately; security incident report; re-test all video access'),
    ], [1.3, 4.7, 0.7, 0.7, 1.5, 3.8, 3.8])

    h1(doc, 'SECTION 16 – APPROVALS')
    body(doc, 'This document becomes effective once signed off by all stakeholders listed below:')
    table(doc, ['Role', 'Full Name', 'Signature', 'Approval Date'], [
        ('Test Lead', 'Vo Van Thanh Dat', '', ''),
        ('Project Manager', 'Nguyen Van A', '', ''),
        ('Product Owner', 'Tran Thi B', '', ''),
        ('Dev Lead', 'Le Van C', '', ''),
    ], [3.5, 3.5, 4, 3.5])
    doc.add_paragraph()
    body(doc, 'Note: This document is in Draft status. Signatures will be collected after the first review round '
              'is complete (Review Date: 2026-06-19).')

    out = r'd:\Vo_Van_Thanh_Dat\HocAI\Lab4_TestPlan_English_VoVanThanhDat_DE190211.docx'
    doc.save(out)
    print('Saved:', out)


# =====================================================================
# ============== FILE 2: PHẦN B + C – TIẾNG VIỆT ======================
# =====================================================================
def screenshot_ph(doc, label):
    doc.add_paragraph()
    body(doc, '[Ảnh chụp màn hình cuộc hội thoại với AI – đính kèm bên dưới]')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.name = FONT; r.font.size = Pt(12); r.font.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def build_bc():
    doc = new_doc()
    cover(
        doc,
        'TRƯỜNG ĐẠI HỌC FPT – FPT UNIVERSITY',
        'Khoa Công nghệ Thông tin',
        'BÁO CÁO THỰC HÀNH',
        'LAB 4 – PHẦN B & C',
        'Chuẩn áp dụng: IEEE 829-2008  |  Môn học: SWT301 – Software Testing',
        'DỰ ÁN: BEE ACADEMY',
        'AI Interaction Log + Câu hỏi thảo luận (Tiếng Việt)',
        [
            ('Môn học', 'SWT301 – Software Testing'),
            ('Chủ đề', 'Topic 5 – Test Management'),
            ('Chuẩn áp dụng', 'IEEE 829-2008'),
            ('Dự án thực hành', 'Bee Academy – Online Learning Platform v1.0.0'),
            ('Sinh viên thực hiện', 'Võ Văn Thành Đạt'),
            ('Mã số sinh viên', 'DE190211'),
            ('AI hỗ trợ', 'Claude Sonnet 4.6'),
            ('Ngày thực hiện', '12/06/2026'),
        ],
    )
    toc(doc, 'MỤC LỤC', [
        (False, 'PHẦN B – AI INTERACTION LOG'),
        (True, 'Tương tác 1 – Mục 1 & 2: Test Plan Identifier + Introduction'),
        (True, 'Tương tác 2 – Mục 6: Approach (Test Strategy)'),
        (True, 'Tương tác 3 – Mục 15: Risks and Contingencies'),
        (False, 'PHẦN C – CÂU HỎI THẢO LUẬN'),
        (True, 'Câu hỏi 1'),
        (True, 'Câu hỏi 2'),
        (False, 'TÀI LIỆU THAM KHẢO'),
    ])

    # ── PHẦN B ──
    center_line(doc, 'PHẦN B – AI INTERACTION LOG', 16, bold=True)
    doc.add_paragraph()
    body(doc, 'Phần này ghi lại 3 lần tương tác với AI (Claude Sonnet 4.6) khi soạn thảo Test Plan cho dự án '
              'Bee Academy, kèm đánh giá chất lượng kết quả và những điểm đã được chỉnh sửa so với output gốc của AI.')

    # Tương tác 1
    h2(doc, 'Tương tác 1 – Mục 1 & 2: Test Plan Identifier + Introduction')
    table(doc, ['Thông tin', 'Chi tiết'], [
        ('Prompt đã dùng',
         'Create Test Plan Identifier and Introduction sections for Bee Academy - an online learning platform for '
         'Vietnamese THCS students (grade 6-9). Tech stack: React 19 + Spring Boot 3.2 + PostgreSQL via Supabase '
         '(JWT ES256/ECDSA P-256). Users: students, teachers, parents, admins. Payment: VNPay/MoMo. Follow IEEE '
         '829-2008 format.'),
    ], [3.5, 11], header_bg='2E4057')
    screenshot_ph(doc, '[ Screenshot 1: Prompt và Response từ AI cho Mục 1 & 2 ]')
    doc.add_paragraph()
    h3(doc, 'Bảng đánh giá kết quả AI:')
    table(doc, ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'], [
        ('Đúng chuẩn IEEE 829-2008', '4/5', 'AI tạo đúng các trường Document ID, Author, Status nhưng thiếu Review Date và document history'),
        ('Phù hợp với context Bee Academy', '3/5', 'AI không biết về JWT ES256/ECDSA P-256 của Supabase, đề xuất HS256 generic'),
        ('Đủ chi tiết, có thể dùng ngay', '3/5', 'Introduction quá ngắn, thiếu phần tài liệu tham chiếu đặc thù dự án'),
        ('Cần chỉnh sửa nhiều hay ít', '3/5', 'Phải chỉnh sửa vừa phải: thêm JWT specifics, Supabase endpoints, VNPay context'),
        ('Tổng', '13/20', ''),
    ], [5, 2, 7.5])
    h3(doc, 'Những gì đã chỉnh sửa so với output AI:')
    for item in ['Thêm thông tin JWT ES256 (ECDSA P-256) thay vì HS256 mà AI đề xuất',
                 'Bổ sung Supabase JWKS endpoint cụ thể vào References',
                 'Thêm UseCase v6.5 vào danh sách tài liệu tham chiếu',
                 'Điều chỉnh scope: loại bỏ Next.js (AI tưởng dự án dùng Next.js) → Vite 6']:
        bul(doc, item)

    # Tương tác 2
    h2(doc, 'Tương tác 2 – Mục 6: Approach (Test Strategy)')
    table(doc, ['Thông tin', 'Chi tiết'], [
        ('Prompt đã dùng',
         'Write the Test Approach section for Bee Academy following IEEE 829-2008. The system includes: Auth '
         '(JWT ES256 ECDSA, Google OAuth, OTP), Course management with video private Supabase bucket, VNPay payment '
         'integration, Teacher portal with video upload, Quiz system with JSONB snapshot grading. Include test levels '
         '(unit/integration/system/acceptance), test types, design techniques (EP, BVA, decision table, state '
         'transition), entry/exit criteria for each level, and tools.'),
    ], [3.5, 11], header_bg='2E4057')
    screenshot_ph(doc, '[ Screenshot 2: Prompt và Response từ AI cho Mục 6 ]')
    doc.add_paragraph()
    h3(doc, 'Bảng đánh giá kết quả AI:')
    table(doc, ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'], [
        ('Đúng chuẩn IEEE 829-2008', '5/5', 'Sau khi được cung cấp context đầy đủ, AI tạo đúng cấu trúc 4 test levels với entry/exit criteria'),
        ('Phù hợp với context Bee Academy', '4/5', 'AI áp dụng EP và BVA đúng cho các trường hợp của Bee Academy (giá tiền VND, cấp lớp 6-9)'),
        ('Đủ chi tiết, có thể dùng ngay', '4/5', 'Khá đầy đủ; entry/exit criteria cụ thể, test types phân loại rõ ràng'),
        ('Cần chỉnh sửa nhiều hay ít', '4/5', 'Chỉ cần chỉnh sửa nhỏ: thay Selenium → Playwright, thêm State Transition cho course lifecycle'),
        ('Tổng', '17/20', ''),
    ], [5, 2, 7.5])
    h3(doc, 'Những gì đã chỉnh sửa so với output AI:')
    for item in ['Thay Selenium 4.x bằng Playwright 1.40+ (phù hợp hơn với React SPA)',
                 'Thêm State Transition Testing cho vòng đời khóa học: DRAFT→PENDING→PUBLISHED/REJECTED',
                 'Bổ sung test type Security Testing riêng cho module Auth (AI gộp chung vào Functional)',
                 'Điều chỉnh pass rate threshold: từ 90% (AI đề xuất) lên 95% theo requirement của project']:
        bul(doc, item)

    # Tương tác 3
    h2(doc, 'Tương tác 3 – Mục 15: Risks and Contingencies')
    table(doc, ['Thông tin', 'Chi tiết'], [
        ('Prompt đã dùng',
         'Identify and analyze testing risks for Bee Academy online learning platform (React 19 + Spring Boot 3.2 '
         '+ Supabase PostgreSQL). Specific technical concerns: VNPay sandbox vs production behavioral differences, '
         'Supabase private bucket signed URL TTL 1 hour, JWT ES256 key dependencies, Supabase free tier rate limits '
         'during testing. Follow IEEE 829-2008 format with Risk ID, description, category, likelihood (H/M/L), '
         'impact (H/M/L), risk level, mitigation strategy, and contingency plan.'),
    ], [3.5, 11], header_bg='2E4057')
    screenshot_ph(doc, '[ Screenshot 3: Prompt và Response từ AI cho Mục 15 ]')
    doc.add_paragraph()
    h3(doc, 'Bảng đánh giá kết quả AI:')
    table(doc, ['Tiêu chí', 'Điểm (1–5)', 'Nhận xét'], [
        ('Đúng chuẩn IEEE 829-2008', '4/5', 'Format Risk ID, L, I, Level, Mitigation, Contingency đúng chuẩn; thiếu cột Risk Category'),
        ('Phù hợp với context Bee Academy', '4/5', 'AI nhận diện được rủi ro kỹ thuật đặc thù (signed URL, VNPay sandbox) sau khi được cung cấp context'),
        ('Đủ chi tiết, có thể dùng ngay', '3/5', 'Thiếu rủi ro về business (legal compliance payment, PO sign-off delay) và rủi ro test data'),
        ('Cần chỉnh sửa nhiều hay ít', '3/5', 'Chỉnh sửa vừa phải: thêm 3 rủi ro mới, xóa 2 rủi ro generic không phù hợp'),
        ('Tổng', '14/20', ''),
    ], [5, 2, 7.5])
    h3(doc, 'Những gì đã chỉnh sửa so với output AI:')
    for item in ['Thêm R07 (Test data corrupt sau payment flow) – AI bỏ sót hoàn toàn rủi ro này',
                 'Thêm R08 (Signed URL TTL expire mid-session) – AI đề cập nhưng không có contingency plan',
                 'Thêm R09 (Supabase RLS misconfiguration → video public) – rủi ro security quan trọng',
                 'Xóa "Database backup failure" – không phù hợp với context dùng Supabase managed service',
                 'Điều chỉnh mức L của R03 (VNPay sandbox) từ M→H vì đã gặp issue thực tế khi dev']:
        bul(doc, item)
    doc.add_paragraph()
    body(doc, 'Lưu ý quan trọng: AI có thể tạo ra nội dung nghe có vẻ hợp lý nhưng thiếu chi tiết kỹ thuật cụ thể '
              'hoặc không phù hợp với project. Nhiệm vụ của tester là review và refine output, không copy-paste mù quáng.')

    doc.add_page_break()

    # ── PHẦN C ──
    center_line(doc, 'PHẦN C – CÂU HỎI THẢO LUẬN', 16, bold=True)
    doc.add_paragraph()

    h1(doc, 'CÂU HỎI 1')
    body(doc, 'Khi sử dụng AI để tạo Test Plan, AI thường mạnh ở những mục nào và yếu ở những mục nào trong chuẩn '
              'IEEE 829-2008? Giải thích tại sao có sự khác biệt đó. Một tester chuyên nghiệp cần bổ sung gì mà AI '
              'không thể tự làm được?', justify=False)
    doc.add_paragraph()
    h2(doc, 'Trả lời:')
    body(doc, 'Qua trải nghiệm thực tế khi làm bài, AI (Claude Sonnet 4.6) thể hiện rõ điểm mạnh và điểm yếu khác '
              'nhau tùy từng mục của IEEE 829-2008.')
    h3(doc, 'Những mục AI làm tốt:')
    body(doc, 'AI mạnh nhất ở các mục có cấu trúc cố định, theo template rõ ràng: Mục 1 (Test Plan Identifier), '
              'Mục 9 (Test Deliverables), Mục 12 (Responsibilities – RACI matrix), Mục 14 (Schedule – Gantt chart), '
              'Mục 16 (Approvals). Lý do là những mục này về cơ bản chỉ cần điền thông tin vào template có sẵn; AI '
              'được huấn luyện trên hàng nghìn tài liệu IEEE, nên nó nắm vững cấu trúc bắt buộc. Ví dụ, khi yêu cầu '
              'tạo RACI matrix, AI trả về đúng format R/A/C/I với 5 roles chuẩn ngay lần đầu.')
    body(doc, 'AI cũng làm tốt Mục 6 (Approach) khi được cung cấp context đầy đủ về tech stack và business domain. '
              'Sau khi tôi mô tả chi tiết về JWT ES256, Supabase Storage, VNPay payment flow, AI tạo được test levels '
              'và entry/exit criteria khá chính xác, chỉ cần điều chỉnh nhỏ.')
    h3(doc, 'Những mục AI còn yếu:')
    body(doc, 'AI yếu nhất ở Mục 4 (Features to be Tested) và Mục 7 (Pass/Fail Criteria). Đối với Mục 4, AI có xu '
              'hướng liệt kê các feature theo tên module một cách generic (Login, Register, View Course...) mà không '
              'hiểu các ràng buộc business thực tế. Ví dụ, AI không biết rằng nút "Thêm vào giỏ hàng" trong Bee '
              'Academy chỉ khả dụng khi đã login, hoặc việc giáo viên submit khóa học phải qua flow '
              'DRAFT→PENDING→Admin duyệt. Những business rule này chỉ có trong CLAUDE.md và UseCase v6.5 – tài liệu '
              'nội bộ mà AI không có access.')
    body(doc, 'Mục 15 (Risks) là điểm yếu đáng chú ý nhất. AI thường đề xuất rủi ro generic (build failure, resource '
              'shortage, environment instability) – những rủi ro mà mọi dự án web đều gặp. AI bỏ sót hoàn toàn các '
              'rủi ro đặc thù của Bee Academy: JWT ES256 key rotation có thể làm tất cả user bị logout đồng loạt, '
              'Supabase signed URL TTL 1 giờ có thể expire mid-test-session, hay việc private video bucket bị '
              'misconfigure RLS sẽ gây lộ nội dung có phí. Đây là những rủi ro "chỉ người trong dự án mới biết".')
    h3(doc, 'Điều tester chuyên nghiệp cần bổ sung:')
    for item in ['Hiểu domain & business context: AI không thể đọc giữa các dòng của UseCase, biết stakeholder nào '
                 'quan tâm điều gì, hay đâu là "killer feature" cần test kỹ nhất.',
                 'Judgment call về priority và risk: Quyết định pass rate threshold 95% hay 90%, hay việc một lỗi '
                 'Minor trên payment page thực chất là Critical – AI không có context để ra quyết định này.',
                 'Thương lượng với stakeholders: Test Plan cần phản ánh thỏa thuận thực sự giữa QA, Dev và Business. '
                 'AI tạo ra tài liệu một chiều, không qua negotiation.',
                 'Verify feasibility: AI đề xuất lịch test 1 tháng với 3 tester nhưng không biết team còn đang làm '
                 'việc khác, hay môi trường Supabase free tier có rate limit.',
                 'Cập nhật liên tục: Test Plan phải thay đổi theo project evolution. AI chỉ tạo snapshot tại một thời '
                 'điểm, không tự cập nhật khi requirement thay đổi.']:
        bul(doc, item)

    h1(doc, 'CÂU HỎI 2')
    body(doc, 'Mục 15 (Risks & Contingencies) yêu cầu tester phải suy nghĩ về những gì có thể sai. So sánh danh sách '
              'rủi ro do AI đề xuất với những rủi ro bạn tự nghĩ ra dựa trên kiến thức môn học. AI có bỏ sót loại rủi '
              'ro nào không? Từ đó, rút ra bài học về giới hạn của AI trong Test Management.', justify=False)
    doc.add_paragraph()
    h2(doc, 'Trả lời:')
    h3(doc, 'So sánh rủi ro AI đề xuất vs. rủi ro tự phân tích:')
    body(doc, 'Khi tôi yêu cầu AI liệt kê rủi ro cho Bee Academy, danh sách ban đầu AI trả về bao gồm: (1) Unstable '
              'test environment, (2) Resource shortage, (3) Requirements change, (4) Build delays, (5) Missing test '
              'data. Đây đều là những rủi ro kinh điển trong mọi dự án phần mềm – đúng nhưng chưa đủ.')
    body(doc, 'Khi tự phân tích dựa trên kiến thức môn SWT301 và kiến trúc thực tế của Bee Academy từ CLAUDE.md, tôi '
              'phát hiện thêm các loại rủi ro mà AI đã bỏ sót hoàn toàn:')
    h3(doc, 'Loại rủi ro AI bỏ sót – Technical Risk đặc thù:')
    for item in ['R02 – JWT ES256 Key Dependency: Bee Academy dùng ECDSA P-256, khác với HS256 thông thường. Nếu '
                 'JWKS endpoint của Supabase fetch thất bại khi backend khởi động, ES256 verifier = null và toàn bộ '
                 'JWT sẽ fail authentication. Đây là rủi ro duy nhất của Supabase JWT implementation, không có trong '
                 'bất kỳ checklist generic nào.',
                 'R08 – Signed URL TTL Expiry: Video trong private Supabase bucket được truy cập qua signed URL TTL '
                 '1 giờ. Nếu test session kéo dài > 1 giờ, URL hết hạn và tester nhầm tưởng đây là bug video '
                 'playback. AI không biết điều này vì không có access vào code của Lesson.videoStoragePath.',
                 'R09 – Supabase RLS Misconfiguration: Row Level Security của Supabase nếu không được set đúng cho '
                 'bucket course-videos (private), video có thể bị public access mà không cần auth. Đây là security '
                 'risk nghiêm trọng mà chỉ tester hiểu storage architecture mới nhận ra.']:
        bul(doc, item)
    h3(doc, 'Loại rủi ro AI bỏ sót – Business Risk:')
    for item in ['Risk về VNPay behavioral difference: Qua thực tế dev, VNPay webhook không reliable trong local '
                 'environment (ngay cả với ngrok tunneling). AI không biết điều này vì chỉ biết "payment integration '
                 'risk" theo nghĩa chung chung.',
                 'Risk về DEV_MODE=true trong production: CLAUDE.md ghi rõ DEV_MODE=true trong backend/.env – nếu '
                 'quên đổi khi deploy, OTP sẽ log ra console thay vì gửi email thật, đây là security vulnerability '
                 'nghiêm trọng. AI không phân tích được config file cụ thể này.']:
        bul(doc, item)
    h3(doc, 'Bài học rút ra về giới hạn của AI trong Test Management:')
    body(doc, 'AI hoạt động dựa trên pattern matching từ training data. Nó giỏi tạo ra những gì "thường gặp" trong '
              'tài liệu kiểm thử, nhưng không thể suy luận từ code base, architecture decision, hay business '
              'constraint đặc thù của từng project.')
    body(doc, 'Giới hạn cốt lõi của AI trong Test Management là: AI không có "project context" – nó không đọc được '
              'CLAUDE.md, UseCase v6.5, hay biết rằng DEV_MODE=true trong .env là một time bomb chờ phát nổ khi '
              'go-live. Một Test Manager giỏi phải kết hợp kiến thức framework (IEEE 829-2008, ISTQB) với deep '
              'understanding về hệ thống cụ thể đang test. AI có thể hỗ trợ phần framework, nhưng phần domain '
              'knowledge và critical thinking vẫn phải do con người đảm nhận.')
    body(doc, 'Kết luận thực tiễn: Workflow hiệu quả nhất là dùng AI để tạo "80% foundation" (structure, template, '
              'generic risks, standard criteria), sau đó tester chuyên nghiệp review, bổ sung project-specific '
              'knowledge, và validate feasibility. Không nên copy-paste output AI trực tiếp vào deliverable mà không '
              'qua critical review – đây là điều Lab 4 muốn chúng ta học được.')

    doc.add_page_break()
    h1(doc, 'TÀI LIỆU THAM KHẢO')
    refs_list = [
        'IEEE 829-2008 – IEEE Standard for Software and System Test Documentation',
        'ISTQB Foundation Level Syllabus v4.0 – Chapter 5: Test Management',
        'Foundations of Software Testing (Craig & Jaskiel) – Chapter 9',
        'ISTQB Glossary – https://glossary.istqb.org',
        'Bee Academy CLAUDE.md – Bee Academy Architecture & Project Documentation',
        'Bee Academy UseCase v6.5 – BEE ACADEMY.md (48 use cases, 9 modules)',
        'Supabase Documentation – Storage Buckets, RLS, JWT Authentication',
        'Spring Boot Security Documentation – JWT ES256 / ECDSA P-256 Configuration',
        'VNPay Developer Documentation – Sandbox Testing Guide',
    ]
    for i, rtext in enumerate(refs_list, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'[{i}] {rtext}')
        run.font.name = FONT
        run.font.size = Pt(13)

    out = r'd:\Vo_Van_Thanh_Dat\HocAI\Lab4_PhanB_C_TiengViet_VoVanThanhDat_DE190211.docx'
    doc.save(out)
    print('Saved:', out)


build_english()
build_bc()
print('DONE')
