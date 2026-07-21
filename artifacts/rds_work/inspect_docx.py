from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def clean(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def iter_block_items(parent):
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, parent)


def main() -> None:
    src = Path(sys.argv[1])
    out = Path(sys.argv[2])
    doc = Document(src)
    report = {
        "source": str(src.resolve()),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "sections": [],
        "styles": [],
        "blocks": [],
        "headers_footers": [],
    }
    for i, section in enumerate(doc.sections, 1):
        report["sections"].append({
            "index": i,
            "page_width_emu": section.page_width,
            "page_height_emu": section.page_height,
            "top_margin_emu": section.top_margin,
            "bottom_margin_emu": section.bottom_margin,
            "left_margin_emu": section.left_margin,
            "right_margin_emu": section.right_margin,
            "header_distance_emu": section.header_distance,
            "footer_distance_emu": section.footer_distance,
            "start_type": str(section.start_type),
        })
        for kind, part in (("header", section.header), ("footer", section.footer)):
            report["headers_footers"].append({
                "section": i,
                "kind": kind,
                "paragraphs": [clean(p.text) for p in part.paragraphs],
                "tables": [[clean(c.text) for row in t.rows for c in row.cells] for t in part.tables],
            })
    for style in doc.styles:
        if style.type == 1:
            report["styles"].append({
                "name": style.name,
                "font": style.font.name,
                "size_pt": style.font.size.pt if style.font.size else None,
                "bold": style.font.bold,
                "italic": style.font.italic,
            })
    p_idx = 0
    t_idx = 0
    for block_no, (kind, obj) in enumerate(iter_block_items(doc), 1):
        if kind == "paragraph":
            p_idx += 1
            text = clean(obj.text)
            report["blocks"].append({
                "block": block_no,
                "type": "paragraph",
                "paragraph": p_idx,
                "style": obj.style.name if obj.style else None,
                "text": text,
                "has_drawing": bool(obj._p.xpath(".//w:drawing | .//w:pict")),
            })
        else:
            t_idx += 1
            rows = []
            for r_idx, row in enumerate(obj.rows, 1):
                rows.append({
                    "row": r_idx,
                    "cells": [clean(cell.text) for cell in row.cells],
                })
            report["blocks"].append({
                "block": block_no,
                "type": "table",
                "table": t_idx,
                "rows": len(obj.rows),
                "cols": len(obj.columns),
                "data": rows,
            })
    out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({k: report[k] for k in ("paragraph_count", "table_count", "inline_shape_count", "section_count")}, ensure_ascii=False))


if __name__ == "__main__":
    main()
